"""
Word Export Manager - Sistema simplificado de exportación a Word
Maneja la exportación de consultas con análisis IA a documentos Word profesionales
"""

import os
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tkinter as tk
from tkinter import messagebox
import re


class WordExportManager:
    """Clase para manejar exportación de consultas a documentos Word"""
    
    def __init__(self):
        self.prospects_folder = "Prospectos"
        self.ensure_prospects_folder()
    
    def ensure_prospects_folder(self):
        """Crea la carpeta Prospectos si no existe"""
        if not os.path.exists(self.prospects_folder):
            os.makedirs(self.prospects_folder)
            print(f"Carpeta '{self.prospects_folder}' creada exitosamente")
    
    def sanitize_filename(self, name):
        """Limpia el nombre para uso en archivos"""
        # Remover caracteres no válidos para nombres de archivo
        sanitized = re.sub(r'[<>:"/\\|?*]', '', name)
        # Reemplazar espacios con guiones bajos
        sanitized = sanitized.replace(' ', '_')
        # Limitar longitud
        return sanitized[:50] if len(sanitized) > 50 else sanitized
    
    def generate_filename(self, prospect_name, consultation_date=None):
        """Genera nombre de archivo con formato: [Nombre_Consultante]_Analisis_[Fecha].docx"""
        clean_name = self.sanitize_filename(prospect_name)
        
        if consultation_date:
            if isinstance(consultation_date, str):
                date_str = consultation_date.replace('-', '')
            else:
                date_str = consultation_date.strftime("%Y%m%d")
        else:
            date_str = datetime.datetime.now().strftime("%Y%m%d")
        
        timestamp = datetime.datetime.now().strftime("%H%M%S")
        filename = f"{clean_name}_Analisis_{date_str}_{timestamp}.docx"
        
        return os.path.join(self.prospects_folder, filename)
    
    def export_consultation_to_word(self, prospect_data, consultation_data, parent_window=None, show_open_dialog=True):
        """
        Exporta una consulta con análisis IA a documento Word
        
        Args:
            prospect_data (dict): Datos del prospecto
            consultation_data (dict): Datos de la consulta con análisis IA
            parent_window: Ventana padre para mostrar mensajes
            
        Returns:
            str: Ruta del archivo generado o None si hubo error
        """
        try:
            # Crear documento Word
            doc = Document()
            
            # Configurar estilos
            self.setup_document_styles(doc)
            
            # Título del documento
            title = doc.add_heading('ANÁLISIS DE CONSULTA LEGAL', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del prospecto
            doc.add_heading('DATOS DEL CONSULTANTE', level=1)
            
            prospect_table = doc.add_table(rows=4, cols=2)
            prospect_table.style = 'Table Grid'
            
            # Llenar tabla de datos del prospecto
            cells = prospect_table.rows[0].cells
            cells[0].text = 'Nombre:'
            cells[1].text = prospect_data.get('nombre', 'N/A')
            
            cells = prospect_table.rows[1].cells
            cells[0].text = 'Contacto:'
            cells[1].text = prospect_data.get('contacto', 'N/A')
            
            cells = prospect_table.rows[2].cells
            cells[0].text = 'Estado:'
            cells[1].text = prospect_data.get('estado', 'N/A')
            
            cells = prospect_table.rows[3].cells
            cells[0].text = 'Fecha de Consulta:'
            fecha_consulta = consultation_data.get('fecha_consulta')
            if fecha_consulta:
                if isinstance(fecha_consulta, str):
                    fecha_str = fecha_consulta
                else:
                    fecha_str = fecha_consulta.strftime("%d/%m/%Y")
            else:
                fecha_str = 'N/A'
            cells[1].text = fecha_str
            
            # Espacio
            doc.add_paragraph()
            
            # Relato original del cliente
            doc.add_heading('RELATO ORIGINAL DEL CLIENTE', level=1)
            relato = consultation_data.get('relato_original_cliente', 'No registrado')
            doc.add_paragraph(relato)
            
            # Espacio
            doc.add_paragraph()
            
            # Análisis de IA
            doc.add_heading('ANÁLISIS LEGAL PROFESIONAL', level=1)
            analisis_ia = consultation_data.get('hechos_reformulados_ia', 'No disponible')
            
            if analisis_ia and analisis_ia != 'No disponible':
                # Dividir el análisis en párrafos para mejor legibilidad
                paragraphs = analisis_ia.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            else:
                doc.add_paragraph('Análisis no disponible o no completado.')
            
            # Pie de página con información de generación
            doc.add_paragraph()
            doc.add_paragraph('_' * 50)
            footer_info = f"Documento generado automáticamente el {datetime.datetime.now().strftime('%d/%m/%Y a las %H:%M:%S')}"
            footer_para = doc.add_paragraph(footer_info)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Generar nombre de archivo y guardar
            filepath = self.generate_filename(
                prospect_data['nombre'], 
                consultation_data.get('fecha_consulta')
            )
            
            doc.save(filepath)
            
            # Mostrar confirmación
            if parent_window and show_open_dialog:
                message = f"Análisis exportado exitosamente:\n\n{os.path.abspath(filepath)}"
                messagebox.showinfo("Exportación Exitosa", message, parent=parent_window)
                
                # Preguntar si quiere abrir el archivo
                if messagebox.askyesno("Abrir Documento", 
                                     "¿Desea abrir el documento generado?", 
                                     parent=parent_window):
                    try:
                        os.startfile(os.path.abspath(filepath))  # Windows
                    except:
                        try:
                            import subprocess
                            subprocess.run(['xdg-open', os.path.abspath(filepath)])  # Linux
                        except:
                            try:
                                subprocess.run(['open', os.path.abspath(filepath)])  # macOS
                            except:
                                messagebox.showinfo("Información", 
                                                   f"No se pudo abrir automáticamente.\nArchivo guardado en:\n{os.path.abspath(filepath)}", 
                                                   parent=parent_window)
            
            return filepath
            
        except Exception as e:
            error_msg = f"Error al exportar a Word: {str(e)}"
            print(error_msg)
            if parent_window:
                messagebox.showerror("Error de Exportación", error_msg, parent=parent_window)
            return None
    
    def setup_document_styles(self, doc):
        """Configura estilos personalizados para el documento"""
        try:
            # Estilo para párrafos normales
            styles = doc.styles
            
            # Verificar si el estilo ya existe
            try:
                normal_style = styles['Normal']
                normal_style.font.name = 'Calibri'
                normal_style.font.size = Inches(0.12)  # 12pt aproximadamente
            except:
                pass  # Si no se puede modificar, usar el estilo por defecto
                
        except Exception as e:
            print(f"Advertencia: No se pudieron configurar estilos personalizados: {e}")
    
    def export_multiple_consultations(self, prospect_data, consultations_list, parent_window=None):
        """
        Exporta múltiples consultas de un prospecto a un solo documento Word
        
        Args:
            prospect_data (dict): Datos del prospecto
            consultations_list (list): Lista de consultas con análisis IA
            parent_window: Ventana padre para mostrar mensajes
            
        Returns:
            str: Ruta del archivo generado o None si hubo error
        """
        try:
            if not consultations_list:
                if parent_window:
                    messagebox.showwarning("Sin Consultas", 
                                         "No hay consultas para exportar.", 
                                         parent=parent_window)
                return None
            
            # Crear documento Word
            doc = Document()
            self.setup_document_styles(doc)
            
            # Título del documento
            title = doc.add_heading('HISTORIAL DE CONSULTAS LEGALES', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del prospecto
            doc.add_heading('DATOS DEL CONSULTANTE', level=1)
            
            prospect_info = f"""
Nombre: {prospect_data.get('nombre', 'N/A')}
Contacto: {prospect_data.get('contacto', 'N/A')}
Estado: {prospect_data.get('estado', 'N/A')}
Total de Consultas: {len(consultations_list)}
Fecha de Reporte: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}
"""
            doc.add_paragraph(prospect_info)
            
            # Separador
            doc.add_paragraph('=' * 60)
            
            # Procesar cada consulta
            for i, consultation in enumerate(consultations_list, 1):
                doc.add_heading(f'CONSULTA #{i}', level=1)
                
                # Fecha de la consulta
                fecha_consulta = consultation.get('fecha_consulta')
                if fecha_consulta:
                    if isinstance(fecha_consulta, str):
                        fecha_str = fecha_consulta
                    else:
                        fecha_str = fecha_consulta.strftime("%d/%m/%Y")
                    doc.add_paragraph(f"Fecha: {fecha_str}")
                
                # Relato original
                doc.add_heading('Relato Original del Cliente:', level=2)
                relato = consultation.get('relato_original_cliente', 'No registrado')
                doc.add_paragraph(relato)
                
                # Análisis de IA
                doc.add_heading('Análisis Legal Profesional:', level=2)
                analisis_ia = consultation.get('hechos_reformulados_ia', 'No disponible')
                
                if analisis_ia and analisis_ia != 'No disponible':
                    paragraphs = analisis_ia.split('\n\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
                else:
                    doc.add_paragraph('Análisis no disponible o no completado.')
                
                # Separador entre consultas
                if i < len(consultations_list):
                    doc.add_paragraph('-' * 60)
                    doc.add_paragraph()
            
            # Pie de página
            doc.add_paragraph()
            doc.add_paragraph('_' * 50)
            footer_info = f"Documento generado automáticamente el {datetime.datetime.now().strftime('%d/%m/%Y a las %H:%M:%S')}"
            footer_para = doc.add_paragraph(footer_info)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Generar nombre de archivo
            clean_name = self.sanitize_filename(prospect_data['nombre'])
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{clean_name}_Historial_Consultas_{timestamp}.docx"
            filepath = os.path.join(self.prospects_folder, filename)
            
            doc.save(filepath)
            
            # Mostrar confirmación
            if parent_window:
                message = f"Historial de consultas exportado exitosamente:\n\n{os.path.abspath(filepath)}"
                messagebox.showinfo("Exportación Exitosa", message, parent=parent_window)
                
                # Preguntar si quiere abrir el archivo
                if messagebox.askyesno("Abrir Documento", 
                                     "¿Desea abrir el documento generado?", 
                                     parent=parent_window):
                    try:
                        os.startfile(os.path.abspath(filepath))
                    except:
                        messagebox.showinfo("Información", 
                                           f"Archivo guardado en:\n{os.path.abspath(filepath)}", 
                                           parent=parent_window)
            
            return filepath
            
        except Exception as e:
            error_msg = f"Error al exportar historial: {str(e)}"
            print(error_msg)
            if parent_window:
                messagebox.showerror("Error de Exportación", error_msg, parent=parent_window)
            return None