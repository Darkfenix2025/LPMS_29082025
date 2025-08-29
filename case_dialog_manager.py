"""
Case Dialog Manager - Refactorizado desde main_app.py
Maneja toda la lógica relacionada con la interfaz de casos
"""
from docxtpl import DocxTemplate
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import crm_database as db
import os
import datetime 
import locale
import logging
import time
import traceback

# Import for generic document generation
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from escrito_generico_dialog import EscritoGenericoDialog

from num2words import num2words

class ErrorMessageManager:
    """
    Clase para manejar mensajes de error de manera centralizada y consistente.
    """
    
    @staticmethod
    def get_error_message(error_type, details=None):
        """
        Obtiene un mensaje de error formateado basado en el tipo de error.
        
        Args:
            error_type (str): Tipo de error
            details (dict): Detalles adicionales del error
            
        Returns:
            dict: Diccionario con título, mensaje y sugerencias
        """
        details = details or {}
        
        error_messages = {
            'missing_case': {
                'title': 'Caso No Encontrado',
                'message': f'No se encontraron datos del caso con ID: {details.get("case_id", "desconocido")}',
                'suggestions': [
                    'Verifique que el caso existe en la base de datos',
                    'Intente refrescar la lista de casos',
                    'Contacte al administrador si el problema persiste'
                ]
            },
            'empty_caratula': {
                'title': 'Datos Incompletos',
                'message': 'El caso no tiene carátula definida',
                'suggestions': [
                    'Complete los datos del caso antes de generar el acuerdo',
                    'Vaya a la pestaña de casos y edite la información',
                    'La carátula es obligatoria para generar documentos'
                ]
            },
            'no_parties': {
                'title': 'Partes Faltantes',
                'message': 'No se encontraron partes asociadas al caso',
                'suggestions': [
                    'Agregue las partes del caso antes de generar el acuerdo',
                    'Vaya a la pestaña de partes y agregue actores y demandados',
                    'Un acuerdo de mediación requiere al menos un actor y un demandado'
                ]
            },
            'no_actors': {
                'title': 'Actores Faltantes',
                'message': 'No se encontraron actores en el caso',
                'suggestions': [
                    'Agregue al menos un actor al caso',
                    'Verifique que las partes estén correctamente configuradas como "Actor"',
                    'Un acuerdo de mediación requiere al menos una parte actora'
                ]
            },
            'no_defendants': {
                'title': 'Demandados Faltantes',
                'message': 'No se encontraron demandados en el caso',
                'suggestions': [
                    'Agregue al menos un demandado al caso',
                    'Verifique que las partes estén correctamente configuradas como "Demandado"',
                    'Un acuerdo de mediación requiere al menos una parte demandada'
                ]
            },
            'missing_dependencies': {
                'title': 'Dependencias Faltantes',
                'message': 'Faltan librerías o archivos necesarios para generar el acuerdo',
                'suggestions': [
                    'Instale las librerías faltantes usando pip',
                    'Verifique que las plantillas estén en la ubicación correcta',
                    'Contacte al administrador del sistema para asistencia'
                ]
            },
            'template_error': {
                'title': 'Error de Plantilla',
                'message': f'Problema con la plantilla: {details.get("error", "error desconocido")}',
                'suggestions': [
                    'Verifique que la plantilla no esté corrupta',
                    'Reemplace la plantilla con una copia válida',
                    'Contacte al soporte técnico si el problema persiste'
                ]
            },
            'validation_error': {
                'title': 'Error de Validación',
                'message': f'Los datos ingresados no son válidos: {details.get("field", "campo desconocido")}',
                'suggestions': [
                    'Revise los datos ingresados',
                    'Asegúrese de que todos los campos obligatorios estén completos',
                    'Verifique el formato de los números y fechas'
                ]
            },
            'file_permission': {
                'title': 'Error de Permisos',
                'message': f'No se puede guardar el archivo: {details.get("path", "ubicación desconocida")}',
                'suggestions': [
                    'Verifique que tiene permisos de escritura en la ubicación',
                    'Intente guardar en una ubicación diferente',
                    'Cierre otros programas que puedan estar usando el archivo'
                ]
            },
            'unexpected_error': {
                'title': 'Error Inesperado',
                'message': f'Error inesperado: {details.get("error", "error desconocido")}',
                'suggestions': [
                    'Intente la operación nuevamente',
                    'Reinicie la aplicación si el problema persiste',
                    'Contacte al soporte técnico con los detalles del error'
                ]
            }
        }
        
        return error_messages.get(error_type, error_messages['unexpected_error'])
    
    @staticmethod
    def show_error_dialog(parent, error_type, details=None):
        """
        Muestra un diálogo de error formateado.
        
        Args:
            parent: Widget padre para el diálogo
            error_type (str): Tipo de error
            details (dict): Detalles adicionales del error
        """
        error_info = ErrorMessageManager.get_error_message(error_type, details)
        
        message = error_info['message']
        if error_info['suggestions']:
            message += "\n\nSoluciones sugeridas:\n"
            for i, suggestion in enumerate(error_info['suggestions'], 1):
                message += f"{i}. {suggestion}\n"
        
        messagebox.showerror(error_info['title'], message, parent=parent)
    
    @staticmethod
    def log_error(error_type, details=None, technical_details=None):
        """
        Registra un error en el log del sistema.
        
        Args:
            error_type (str): Tipo de error
            details (dict): Detalles del error
            technical_details (str): Detalles técnicos adicionales
        """
        import datetime
        
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        error_info = ErrorMessageManager.get_error_message(error_type, details)
        
        log_message = f"[{timestamp}] ERROR: {error_type}"
        log_message += f"\nMensaje: {error_info['message']}"
        
        if details:
            log_message += f"\nDetalles: {details}"
        
        if technical_details:
            log_message += f"\nDetalles técnicos: {technical_details}"
        
        print(log_message)
        
        # Aquí se podría agregar escritura a archivo de log si se desea
        # with open('error.log', 'a', encoding='utf-8') as f:
        #     f.write(log_message + '\n\n')

def convertir_numero_a_letras(numero_str):
    """
    Convierte un número en formato string a su representación en letras en español.
    Implementa manejo robusto de errores y validación exhaustiva de casos edge.
    
    Args:
        numero_str (str): Número en formato string (puede incluir comas, puntos, espacios)
    
    Returns:
        str: Número convertido a letras en mayúsculas, o mensaje de error descriptivo
    """
    # Validación inicial de entrada
    validation_result = _validate_number_input(numero_str)
    if not validation_result['valid']:
        return validation_result['error_message']
    
    try:
        # Limpiar y normalizar el string de entrada
        numero_limpio = _clean_and_normalize_number_string(numero_str)
        if not numero_limpio['success']:
            return numero_limpio['error_message']
        
        # Convertir a número y validar rango
        numero_procesado = _convert_and_validate_number(numero_limpio['cleaned_number'])
        if not numero_procesado['success']:
            return numero_procesado['error_message']
        
        # Realizar conversión a letras
        conversion_result = _perform_number_to_words_conversion(numero_procesado['number'])
        return conversion_result
        
    except Exception as e:
        # Log del error para debugging con más detalles
        error_msg = f"Error crítico en convertir_numero_a_letras con entrada '{numero_str}': {e}"
        print(f"[ERROR] {error_msg}")
        
        # Registrar stack trace para debugging
        import traceback
        print(f"[ERROR] Stack trace: {traceback.format_exc()}")
        
        return "ERROR EN CONVERSIÓN"

def _validate_number_input(numero_str):
    """
    Valida la entrada para la conversión de números.
    
    Args:
        numero_str: String a validar
        
    Returns:
        dict: Resultado de la validación
    """
    try:
        # Verificar que no sea None o vacío
        if numero_str is None:
            return {'valid': False, 'error_message': 'ENTRADA NULA'}
        
        # Verificar que sea string
        if not isinstance(numero_str, str):
            return {'valid': False, 'error_message': 'ENTRADA NO ES TEXTO'}
        
        # Verificar que no esté vacío después de limpiar espacios
        if not numero_str.strip():
            return {'valid': False, 'error_message': 'ENTRADA VACÍA'}
        
        # Verificar longitud razonable
        if len(numero_str.strip()) > 50:
            return {'valid': False, 'error_message': 'ENTRADA DEMASIADO LARGA'}
        
        # Verificar que contenga al menos un dígito
        if not any(c.isdigit() for c in numero_str):
            return {'valid': False, 'error_message': 'NO CONTIENE DÍGITOS'}
        
        return {'valid': True, 'error_message': None}
        
    except Exception as e:
        print(f"[ERROR] Error validando entrada de número: {e}")
        return {'valid': False, 'error_message': 'ERROR DE VALIDACIÓN'}

def _clean_and_normalize_number_string(numero_str):
    """
    Limpia y normaliza el string de número.
    
    Args:
        numero_str: String a limpiar
        
    Returns:
        dict: Resultado de la limpieza
    """
    try:
        numero_limpio = numero_str.strip()
        
        # Remover caracteres de moneda comunes
        currency_symbols = ['$', '€', '£', '¥', '₹', '₽', 'USD', 'EUR', 'ARS', 'CLP', 'MXN']
        for symbol in currency_symbols:
            numero_limpio = numero_limpio.replace(symbol, '')
        
        # Remover espacios internos
        numero_limpio = numero_limpio.replace(' ', '')
        
        # Manejar diferentes formatos de separadores decimales
        numero_limpio = _handle_decimal_separators(numero_limpio)
        
        # Remover caracteres no válidos excepto dígitos, punto y signo negativo
        numero_limpio = ''.join(c for c in numero_limpio if c.isdigit() or c in '.-')
        
        # Validar resultado de la limpieza
        if not numero_limpio or numero_limpio in ['.', '-', '-.', '--']:
            return {'success': False, 'error_message': 'NÚMERO INVÁLIDO DESPUÉS DE LIMPIEZA'}
        
        # Validar formato de número
        if numero_limpio.count('.') > 1 or numero_limpio.count('-') > 1:
            return {'success': False, 'error_message': 'FORMATO DE NÚMERO INVÁLIDO'}
        
        # Validar posición del signo negativo
        if '-' in numero_limpio and not numero_limpio.startswith('-'):
            return {'success': False, 'error_message': 'SIGNO NEGATIVO EN POSICIÓN INCORRECTA'}
        
        return {'success': True, 'cleaned_number': numero_limpio}
        
    except Exception as e:
        print(f"[ERROR] Error limpiando string de número: {e}")
        return {'success': False, 'error_message': 'ERROR EN LIMPIEZA'}

def _handle_decimal_separators(numero_str):
    """
    Maneja diferentes formatos de separadores decimales.
    
    Args:
        numero_str: String con posibles separadores
        
    Returns:
        str: String normalizado
    """
    try:
        # Si contiene tanto coma como punto, determinar el formato
        if ',' in numero_str and '.' in numero_str:
            pos_coma = numero_str.rfind(',')
            pos_punto = numero_str.rfind('.')
            
            if pos_coma > pos_punto:
                # Formato argentino: 1.000,50
                return numero_str.replace('.', '').replace(',', '.')
            else:
                # Formato internacional: 1,000.50
                return numero_str.replace(',', '')
                
        elif ',' in numero_str:
            # Solo coma - determinar si es separador de miles o decimal
            partes = numero_str.split(',')
            if len(partes) == 2 and len(partes[1]) <= 2 and partes[1].isdigit():
                # Probablemente separador decimal (ej: 1000,50)
                return numero_str.replace(',', '.')
            else:
                # Probablemente separador de miles (ej: 1,000 o 1,000,000)
                return numero_str.replace(',', '')
                
        elif '.' in numero_str:
            # Solo punto - determinar si es separador de miles o decimal
            partes = numero_str.split('.')
            if (len(partes) == 2 and 
                len(partes[1]) <= 2 and 
                len(partes[0]) <= 3 and 
                partes[1].isdigit()):
                # Probablemente separador decimal (ej: 100.50)
                return numero_str
            else:
                # Probablemente separador de miles (ej: 1.000 o 1.000.000)
                return numero_str.replace('.', '')
        
        return numero_str
        
    except Exception as e:
        print(f"[ERROR] Error manejando separadores decimales: {e}")
        return numero_str

def _convert_and_validate_number(numero_str):
    """
    Convierte el string a número y valida el rango.
    
    Args:
        numero_str: String limpio de número
        
    Returns:
        dict: Resultado de la conversión
    """
    try:
        # Intentar conversión a float
        try:
            numero_float = float(numero_str)
        except ValueError as e:
            return {'success': False, 'error_message': f'FORMATO NUMÉRICO INVÁLIDO: {e}'}
        except OverflowError:
            return {'success': False, 'error_message': 'NÚMERO DEMASIADO GRANDE PARA PROCESAR'}
        
        # Verificar que no sea infinito o NaN
        if not _is_finite_number(numero_float):
            return {'success': False, 'error_message': 'NÚMERO NO FINITO'}
        
        # Truncar decimales para la conversión a letras
        numero_int = int(numero_float)
        
        # Validar rango razonable
        if abs(numero_int) > 999999999999:  # Más de 999 mil millones
            return {'success': False, 'error_message': 'NÚMERO DEMASIADO GRANDE'}
        
        # Validar que no sea demasiado pequeño (pero permitir cero)
        if numero_int != 0 and abs(numero_int) < 1:
            return {'success': False, 'error_message': 'NÚMERO DEMASIADO PEQUEÑO'}
        
        return {'success': True, 'number': numero_int}
        
    except Exception as e:
        print(f"[ERROR] Error convirtiendo y validando número: {e}")
        return {'success': False, 'error_message': 'ERROR EN CONVERSIÓN NUMÉRICA'}

def _is_finite_number(number):
    """Verifica si un número es finito."""
    try:
        import math
        return math.isfinite(number)
    except:
        # Fallback para versiones de Python sin math.isfinite
        return not (str(number).lower() in ['inf', '-inf', 'nan'])

def _perform_number_to_words_conversion(numero_int):
    """
    Realiza la conversión del número a palabras.
    
    Args:
        numero_int: Número entero a convertir
        
    Returns:
        str: Número convertido a letras
    """
    try:
        # Caso especial para cero
        if numero_int == 0:
            return "CERO"
        
        # Verificar disponibilidad de num2words
        try:
            from num2words import num2words
        except ImportError as e:
            print(f"[ERROR] Librería num2words no disponible: {e}")
            return "ERROR: LIBRERÍA NO DISPONIBLE"
        
        # Realizar conversión con manejo de errores específicos
        try:
            resultado = num2words(numero_int, to='cardinal', lang='es')
        except ValueError as e:
            print(f"[ERROR] Error de valor en num2words: {e}")
            return "ERROR: VALOR NO SOPORTADO"
        except Exception as e:
            print(f"[ERROR] Error en num2words: {e}")
            return "ERROR EN CONVERSIÓN"
        
        # Validar resultado de num2words
        if not resultado or not isinstance(resultado, str):
            return "ERROR: RESULTADO INVÁLIDO"
        
        # Limpiar y formatear el resultado
        resultado = resultado.strip()
        if not resultado:
            return "ERROR: RESULTADO VACÍO"
        
        # Convertir a mayúsculas
        resultado = resultado.upper()
        
        # Manejar números negativos de manera consistente
        if numero_int < 0:
            if not resultado.startswith("MENOS"):
                resultado = "MENOS " + resultado.replace("MENOS ", "")
        
        # Validar longitud del resultado
        if len(resultado) > 500:  # Evitar resultados extremadamente largos
            return "NÚMERO DEMASIADO COMPLEJO"
        
        return resultado
        
    except Exception as e:
        print(f"[ERROR] Error crítico en conversión a palabras: {e}")
        return "ERROR CRÍTICO EN CONVERSIÓN"

def convertir_plazo_a_letras(plazo_str):
    """
    Convierte un plazo en días a su representación en letras.
    Función específica para plazos que siempre deben ser números enteros positivos.
    Implementa validación robusta y manejo de casos edge.
    
    Args:
        plazo_str (str): Plazo en días como string
    
    Returns:
        str: Plazo convertido a letras en mayúsculas
    """
    # Validación inicial de entrada para plazos
    validation_result = _validate_period_input(plazo_str)
    if not validation_result['valid']:
        return validation_result['default_value']
    
    try:
        # Limpiar y extraer el número del plazo
        plazo_procesado = _clean_and_extract_period_number(plazo_str)
        if not plazo_procesado['success']:
            return plazo_procesado['default_value']
        
        # Validar rango específico para plazos
        plazo_validado = _validate_period_range(plazo_procesado['period_number'])
        if not plazo_validado['valid']:
            return plazo_validado['result_message']
        
        # Realizar conversión a letras para plazos
        conversion_result = _perform_period_to_words_conversion(plazo_validado['period'])
        return conversion_result
        
    except Exception as e:
        # Log del error para debugging con contexto específico de plazos
        error_msg = f"Error crítico en convertir_plazo_a_letras con entrada '{plazo_str}': {e}"
        print(f"[ERROR] {error_msg}")
        
        # Registrar stack trace para debugging
        import traceback
        print(f"[ERROR] Stack trace en conversión de plazo: {traceback.format_exc()}")
        
        return "CERO"

def _validate_period_input(plazo_str):
    """
    Valida la entrada para la conversión de plazos.
    
    Args:
        plazo_str: String a validar
        
    Returns:
        dict: Resultado de la validación
    """
    try:
        # Verificar que no sea None
        if plazo_str is None:
            print("[DEBUG] Plazo es None, usando valor por defecto")
            return {'valid': False, 'default_value': 'CERO'}
        
        # Verificar que sea string
        if not isinstance(plazo_str, str):
            print(f"[WARNING] Plazo no es string: {type(plazo_str)}, convirtiendo")
            try:
                plazo_str = str(plazo_str)
            except:
                return {'valid': False, 'default_value': 'CERO'}
        
        # Verificar que no esté vacío después de limpiar espacios
        if not plazo_str.strip():
            print("[DEBUG] Plazo vacío, usando valor por defecto")
            return {'valid': False, 'default_value': 'CERO'}
        
        # Verificar longitud razonable
        if len(plazo_str.strip()) > 20:
            print(f"[WARNING] Plazo demasiado largo: {len(plazo_str)} caracteres")
            return {'valid': False, 'default_value': 'PLAZO INVÁLIDO'}
        
        return {'valid': True, 'default_value': None}
        
    except Exception as e:
        print(f"[ERROR] Error validando entrada de plazo: {e}")
        return {'valid': False, 'default_value': 'CERO'}

def _clean_and_extract_period_number(plazo_str):
    """
    Limpia y extrae el número del plazo.
    
    Args:
        plazo_str: String del plazo
        
    Returns:
        dict: Resultado de la extracción
    """
    try:
        plazo_limpio = plazo_str.strip()
        
        # Remover palabras comunes relacionadas con tiempo
        time_words = ['días', 'dia', 'day', 'days', 'hábiles', 'corridos', 'calendario']
        for word in time_words:
            plazo_limpio = plazo_limpio.lower().replace(word.lower(), '')
        
        # Remover espacios y caracteres especiales, mantener solo dígitos
        plazo_limpio = ''.join(filter(str.isdigit, plazo_limpio))
        
        # Verificar que quedó algo después de la limpieza
        if not plazo_limpio:
            print("[DEBUG] No se encontraron dígitos en el plazo")
            return {'success': False, 'default_value': 'CERO'}
        
        # Verificar longitud razonable para un plazo
        if len(plazo_limpio) > 10:  # Más de 10 dígitos es poco realista para días
            print(f"[WARNING] Plazo con demasiados dígitos: {len(plazo_limpio)}")
            return {'success': False, 'default_value': 'PLAZO DEMASIADO GRANDE'}
        
        # Convertir a entero
        try:
            plazo_int = int(plazo_limpio)
        except ValueError as e:
            print(f"[ERROR] Error convirtiendo plazo a entero: {e}")
            return {'success': False, 'default_value': 'PLAZO INVÁLIDO'}
        
        return {'success': True, 'period_number': plazo_int}
        
    except Exception as e:
        print(f"[ERROR] Error limpiando y extrayendo número de plazo: {e}")
        return {'success': False, 'default_value': 'CERO'}

def _validate_period_range(plazo_int):
    """
    Valida el rango del plazo.
    
    Args:
        plazo_int: Número entero del plazo
        
    Returns:
        dict: Resultado de la validación
    """
    try:
        # Validar que sea un número positivo
        if plazo_int < 0:
            print(f"[WARNING] Plazo negativo: {plazo_int}")
            return {'valid': False, 'result_message': 'PLAZO INVÁLIDO'}
        
        # Caso especial para cero
        if plazo_int == 0:
            print("[DEBUG] Plazo es cero")
            return {'valid': True, 'period': plazo_int, 'result_message': 'CERO'}
        
        # Validar rangos específicos con mensajes descriptivos
        if plazo_int > 3650:  # Más de 10 años
            print(f"[WARNING] Plazo extremadamente largo: {plazo_int} días")
            return {'valid': False, 'result_message': 'MÁS DE DIEZ AÑOS'}
        elif plazo_int > 1825:  # Más de 5 años
            print(f"[WARNING] Plazo muy largo: {plazo_int} días")
            return {'valid': False, 'result_message': 'MÁS DE CINCO AÑOS'}
        elif plazo_int > 730:  # Más de 2 años
            print(f"[INFO] Plazo largo: {plazo_int} días")
            return {'valid': False, 'result_message': 'MÁS DE DOS AÑOS'}
        elif plazo_int > 365:  # Más de 1 año
            print(f"[INFO] Plazo mayor a un año: {plazo_int} días")
            return {'valid': False, 'result_message': 'MÁS DE UN AÑO'}
        
        # Plazo en rango normal (1-365 días)
        return {'valid': True, 'period': plazo_int}
        
    except Exception as e:
        print(f"[ERROR] Error validando rango de plazo: {e}")
        return {'valid': False, 'result_message': 'ERROR EN VALIDACIÓN'}

def _perform_period_to_words_conversion(plazo_int):
    """
    Realiza la conversión del plazo a palabras.
    
    Args:
        plazo_int: Número entero del plazo
        
    Returns:
        str: Plazo convertido a letras
    """
    try:
        # Caso especial para cero
        if plazo_int == 0:
            return "CERO"
        
        # Verificar disponibilidad de num2words
        try:
            from num2words import num2words
        except ImportError as e:
            print(f"[ERROR] Librería num2words no disponible para conversión de plazo: {e}")
            return f"ERROR: {plazo_int} DÍAS"
        
        # Realizar conversión con manejo de errores específicos
        try:
            resultado = num2words(plazo_int, to='cardinal', lang='es')
        except ValueError as e:
            print(f"[ERROR] Error de valor en num2words para plazo: {e}")
            return f"ERROR: {plazo_int} DÍAS"
        except Exception as e:
            print(f"[ERROR] Error en num2words para plazo: {e}")
            return f"ERROR: {plazo_int} DÍAS"
        
        # Validar resultado de num2words
        if not resultado or not isinstance(resultado, str):
            print(f"[ERROR] Resultado inválido de num2words para plazo: {resultado}")
            return f"ERROR: {plazo_int} DÍAS"
        
        # Limpiar y formatear el resultado
        resultado = resultado.strip().upper()
        
        if not resultado:
            print("[ERROR] Resultado vacío después de limpiar")
            return f"ERROR: {plazo_int} DÍAS"
        
        # Validar longitud del resultado para plazos
        if len(resultado) > 200:  # Los plazos no deberían generar textos muy largos
            print(f"[WARNING] Resultado de conversión de plazo muy largo: {len(resultado)} caracteres")
            return "PLAZO DEMASIADO COMPLEJO"
        
        # Agregar contexto específico para ciertos rangos
        if plazo_int == 1:
            resultado = "UN"  # Más natural que "UNO" para días
        elif plazo_int in [30, 60, 90, 120, 180]:
            # Casos comunes en plazos legales, agregar nota si es necesario
            pass
        
        return resultado
        
    except Exception as e:
        print(f"[ERROR] Error crítico en conversión de plazo a palabras: {e}")
        return f"ERROR: {plazo_int} DÍAS"
    
class CaseManager:
    """Clase que maneja toda la lógica de casos"""

    def __init__(self, app_controller=None):
        self.app_controller = app_controller
        self.db = db
        
        # Configure logging for CaseManager
        self.logger = logging.getLogger(f"{__name__}.CaseManager")
        if not self.logger.handlers:
            # Create console handler with formatting
            console_handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            console_handler.setFormatter(formatter)
            self.logger.addHandler(console_handler)
            self.logger.setLevel(logging.INFO)
        
        self.logger.info(f"CaseManager initialized with app_controller: {app_controller is not None}")

    def _validate_mediation_dependencies(self):
        """
        Valida que todas las dependencias necesarias para la generación de acuerdos de mediación estén disponibles.
        Retorna un diccionario con el resultado de la validación.
        """
        validation_result = {
            'success': True,
            'errors': [],
            'warnings': []
        }
        
        # Verificar librería num2words
        try:
            import num2words
            # Test básico de funcionalidad
            num2words.num2words(1, lang='es')
        except ImportError:
            validation_result['success'] = False
            validation_result['errors'].append({
                'type': 'missing_library',
                'message': 'La librería num2words no está instalada.',
                'suggestion': 'Instale num2words ejecutando: pip install num2words'
            })
        except Exception as e:
            validation_result['success'] = False
            validation_result['errors'].append({
                'type': 'library_error',
                'message': f'Error al usar num2words: {str(e)}',
                'suggestion': 'Reinstale num2words o verifique la configuración del idioma español'
            })
        
        # Verificar librería docxtpl
        try:
            from docxtpl import DocxTemplate
        except ImportError:
            validation_result['success'] = False
            validation_result['errors'].append({
                'type': 'missing_library',
                'message': 'La librería docxtpl no está instalada.',
                'suggestion': 'Instale docxtpl ejecutando: pip install docxtpl'
            })
        
        # Verificar existencia de plantilla (usar la original por defecto)
        template_path = 'plantillas/mediacion/acuerdo_base.docx'
        print(f"[INFO] Usando plantilla original: {template_path}")

        if not os.path.exists(template_path):
            validation_result['success'] = False
            validation_result['errors'].append({
                'type': 'missing_template',
                'message': f'No se encuentra la plantilla en: {template_path}',
                'suggestion': 'Verifique que el archivo de plantilla existe y tiene los permisos correctos'
            })
        else:
            # Verificar que el archivo no esté corrupto
            try:
                from docxtpl import DocxTemplate
                DocxTemplate(template_path)
            except Exception as e:
                validation_result['success'] = False
                validation_result['errors'].append({
                    'type': 'corrupted_template',
                    'message': f'La plantilla parece estar corrupta: {str(e)}',
                    'suggestion': 'Reemplace el archivo de plantilla con una copia válida'
                })
        
        # Verificar permisos de escritura en directorio actual
        try:
            test_file = 'test_write_permission.tmp'
            with open(test_file, 'w') as f:
                f.write('test')
            os.remove(test_file)
        except Exception as e:
            validation_result['warnings'].append({
                'type': 'write_permission',
                'message': f'Posibles problemas de permisos de escritura: {str(e)}',
                'suggestion': 'Verifique los permisos del directorio de trabajo'
            })
        
        return validation_result

    def _validate_template_and_context(self, template_path, context):
        """
        Valida que la plantilla existe y que el contexto contiene todas las variables necesarias.
        Implementa validación exhaustiva con múltiples niveles de verificación.
        
        Args:
            template_path (str): Ruta a la plantilla
            context (dict): Contexto con variables para la plantilla
            
        Returns:
            dict: Resultado de la validación con éxito/errores/advertencias
        """
        print(f"[DEBUG] Iniciando validación completa de plantilla: {template_path}")
        
        validation_result = {
            'success': True,
            'errors': [],
            'warnings': [],
            'info': []
        }
        
        try:
            # 1. Validación inicial de parámetros
            param_validation = self._validate_template_parameters(template_path, context)
            if not param_validation['success']:
                validation_result.update(param_validation)
                return validation_result
            
            # 2. Validación de existencia y accesibilidad del archivo
            file_validation = self._validate_template_file_access(template_path)
            if not file_validation['success']:
                validation_result.update(file_validation)
                return validation_result
            
            validation_result['warnings'].extend(file_validation.get('warnings', []))
            
            # 3. Validación de carga y estructura de la plantilla
            template_validation = self._validate_template_structure(template_path)
            if not template_validation['success']:
                validation_result.update(template_validation)
                return validation_result
            
            validation_result['warnings'].extend(template_validation.get('warnings', []))
            doc = template_validation['template_object']
            
            # 4. Análisis de variables de la plantilla
            variables_analysis = self._analyze_template_variables(doc, context)
            validation_result['warnings'].extend(variables_analysis.get('warnings', []))
            validation_result['info'].extend(variables_analysis.get('info', []))
            
            # 5. Validación profunda del contexto
            context_validation = self._validate_context_comprehensive(context)
            validation_result['warnings'].extend(context_validation.get('warnings', []))
            validation_result['info'].extend(context_validation.get('info', []))
            
            # 6. Validación de compatibilidad plantilla-contexto
            compatibility_validation = self._validate_template_context_compatibility(doc, context)
            validation_result['warnings'].extend(compatibility_validation.get('warnings', []))
            
            print(f"[DEBUG] Validación completada - Errores: {len(validation_result['errors'])}, Advertencias: {len(validation_result['warnings'])}")
            
            return validation_result
            
        except Exception as e:
            print(f"[ERROR] Error crítico durante validación de plantilla: {e}")
            validation_result['success'] = False
            validation_result['errors'].append({
                'type': 'critical_validation_error',
                'message': f'Error crítico durante la validación: {str(e)}',
                'suggestion': 'Verifique la plantilla y el contexto manualmente'
            })
            return validation_result

    def _validate_template_parameters(self, template_path, context):
        """
        Valida los parámetros de entrada para la validación de plantilla.
        
        Args:
            template_path: Ruta de la plantilla
            context: Contexto del documento
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            result = {'success': True, 'errors': [], 'warnings': []}
            
            # Validar template_path
            if not template_path:
                result['success'] = False
                result['errors'].append({
                    'type': 'invalid_template_path',
                    'message': 'Ruta de plantilla no proporcionada',
                    'suggestion': 'Especifique una ruta válida para la plantilla'
                })
                return result
            
            if not isinstance(template_path, str):
                result['success'] = False
                result['errors'].append({
                    'type': 'invalid_template_path_type',
                    'message': f'Ruta de plantilla debe ser string, recibido: {type(template_path)}',
                    'suggestion': 'Proporcione la ruta como string'
                })
                return result
            
            # Validar context
            if not isinstance(context, dict):
                result['success'] = False
                result['errors'].append({
                    'type': 'invalid_context_type',
                    'message': f'Contexto debe ser diccionario, recibido: {type(context)}',
                    'suggestion': 'Proporcione el contexto como diccionario'
                })
                return result
            
            if not context:
                result['warnings'].append({
                    'type': 'empty_context',
                    'message': 'Contexto está vacío',
                    'suggestion': 'Verifique que el contexto contenga las variables necesarias'
                })
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error validando parámetros de plantilla: {e}")
            return {
                'success': False,
                'errors': [{
                    'type': 'parameter_validation_error',
                    'message': f'Error validando parámetros: {str(e)}',
                    'suggestion': 'Verifique los parámetros de entrada'
                }]
            }

    def _validate_template_file_access(self, template_path):
        """
        Valida el acceso al archivo de plantilla.
        
        Args:
            template_path: Ruta de la plantilla
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            result = {'success': True, 'errors': [], 'warnings': []}
            
            # Verificar existencia del archivo
            if not os.path.exists(template_path):
                result['success'] = False
                result['errors'].append({
                    'type': 'missing_template',
                    'message': f'La plantilla no existe: {template_path}',
                    'suggestion': 'Verifique que el archivo de plantilla esté en la ubicación correcta'
                })
                return result
            
            # Verificar que es un archivo (no directorio)
            if not os.path.isfile(template_path):
                result['success'] = False
                result['errors'].append({
                    'type': 'template_not_file',
                    'message': f'La ruta no apunta a un archivo: {template_path}',
                    'suggestion': 'Verifique que la ruta apunte a un archivo válido'
                })
                return result
            
            # Verificar permisos de lectura
            if not os.access(template_path, os.R_OK):
                result['success'] = False
                result['errors'].append({
                    'type': 'template_no_read_permission',
                    'message': f'Sin permisos de lectura para la plantilla: {template_path}',
                    'suggestion': 'Verifique los permisos del archivo de plantilla'
                })
                return result
            
            # Verificar tamaño del archivo
            try:
                file_size = os.path.getsize(template_path)
                if file_size == 0:
                    result['success'] = False
                    result['errors'].append({
                        'type': 'empty_template_file',
                        'message': f'El archivo de plantilla está vacío: {template_path}',
                        'suggestion': 'Reemplace con una plantilla válida'
                    })
                    return result
                elif file_size > 50 * 1024 * 1024:  # 50MB
                    result['warnings'].append({
                        'type': 'large_template_file',
                        'message': f'Archivo de plantilla muy grande: {file_size:,} bytes',
                        'suggestion': 'Considere optimizar la plantilla'
                    })
                
                print(f"[DEBUG] Plantilla accesible: {file_size:,} bytes")
                
            except OSError as e:
                result['warnings'].append({
                    'type': 'template_size_check_failed',
                    'message': f'No se pudo verificar el tamaño del archivo: {str(e)}',
                    'suggestion': 'Verifique manualmente el archivo'
                })
            
            # Verificar extensión del archivo
            if not template_path.lower().endswith('.docx'):
                result['warnings'].append({
                    'type': 'unexpected_template_extension',
                    'message': f'Extensión de archivo inesperada: {os.path.splitext(template_path)[1]}',
                    'suggestion': 'Se esperaba un archivo .docx'
                })
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error validando acceso a archivo de plantilla: {e}")
            return {
                'success': False,
                'errors': [{
                    'type': 'file_access_validation_error',
                    'message': f'Error validando acceso al archivo: {str(e)}',
                    'suggestion': 'Verifique manualmente el archivo de plantilla'
                }]
            }

    def _validate_template_structure(self, template_path):
        """
        Valida la estructura y carga de la plantilla.
        
        Args:
            template_path: Ruta de la plantilla
            
        Returns:
            dict: Resultado de la validación con objeto de plantilla
        """
        try:
            result = {'success': True, 'errors': [], 'warnings': []}
            
            # Verificar disponibilidad de docxtpl
            try:
                from docxtpl import DocxTemplate
            except ImportError as e:
                result['success'] = False
                result['errors'].append({
                    'type': 'missing_docxtpl_library',
                    'message': f'Librería docxtpl no disponible: {str(e)}',
                    'suggestion': 'Instale docxtpl ejecutando: pip install docxtpl'
                })
                return result
            
            # Intentar cargar la plantilla
            try:
                print(f"[DEBUG] Cargando plantilla: {template_path}")
                doc = DocxTemplate(template_path)
                result['template_object'] = doc
                print("[DEBUG] Plantilla cargada exitosamente")
                
            except Exception as e:
                error_msg = str(e).lower()
                
                if 'corrupt' in error_msg or 'damaged' in error_msg:
                    result['success'] = False
                    result['errors'].append({
                        'type': 'corrupted_template',
                        'message': f'La plantilla está corrupta o dañada: {str(e)}',
                        'suggestion': 'Reemplace la plantilla con una copia válida'
                    })
                elif 'permission' in error_msg or 'access' in error_msg:
                    result['success'] = False
                    result['errors'].append({
                        'type': 'template_access_denied',
                        'message': f'Acceso denegado a la plantilla: {str(e)}',
                        'suggestion': 'Verifique los permisos del archivo'
                    })
                elif 'format' in error_msg or 'invalid' in error_msg:
                    result['success'] = False
                    result['errors'].append({
                        'type': 'invalid_template_format',
                        'message': f'Formato de plantilla inválido: {str(e)}',
                        'suggestion': 'Verifique que sea un archivo .docx válido'
                    })
                else:
                    result['success'] = False
                    result['errors'].append({
                        'type': 'template_load_error',
                        'message': f'Error cargando la plantilla: {str(e)}',
                        'suggestion': 'Verifique que la plantilla sea válida'
                    })
                
                return result
            
            # Validaciones adicionales de la plantilla cargada
            try:
                # Verificar que la plantilla tiene contenido
                if hasattr(doc, 'docx') and hasattr(doc.docx, 'part'):
                    print("[DEBUG] Estructura de plantilla verificada")
                else:
                    result['warnings'].append({
                        'type': 'template_structure_warning',
                        'message': 'No se pudo verificar completamente la estructura de la plantilla',
                        'suggestion': 'La plantilla puede funcionar, pero verifique manualmente'
                    })
                
            except Exception as e:
                result['warnings'].append({
                    'type': 'template_structure_check_failed',
                    'message': f'No se pudo verificar la estructura: {str(e)}',
                    'suggestion': 'La plantilla puede funcionar, pero verifique manualmente'
                })
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error validando estructura de plantilla: {e}")
            return {
                'success': False,
                'errors': [{
                    'type': 'structure_validation_error',
                    'message': f'Error validando estructura: {str(e)}',
                    'suggestion': 'Verifique manualmente la plantilla'
                }]
            }

    def _analyze_template_variables(self, doc, context):
        """
        Analiza las variables de la plantilla y su compatibilidad con el contexto.
        
        Args:
            doc: Objeto DocxTemplate
            context: Contexto del documento
            
        Returns:
            dict: Resultado del análisis
        """
        try:
            result = {'warnings': [], 'info': []}
            
            # Obtener variables requeridas por la plantilla
            try:
                required_vars = doc.get_undeclared_template_variables()
                print(f"[DEBUG] Variables requeridas por la plantilla: {len(required_vars)}")
                
                if required_vars:
                    print(f"[DEBUG] Variables encontradas: {sorted(required_vars)}")
                    result['info'].append({
                        'type': 'template_variables_found',
                        'message': f'Plantilla requiere {len(required_vars)} variables',
                        'details': sorted(required_vars)
                    })
                else:
                    result['info'].append({
                        'type': 'no_template_variables',
                        'message': 'La plantilla no requiere variables dinámicas',
                        'details': []
                    })
                
                # Analizar variables faltantes
                missing_vars = [var for var in required_vars if var not in context]
                if missing_vars:
                    result['warnings'].append({
                        'type': 'missing_context_vars',
                        'message': f'Variables faltantes en el contexto: {", ".join(sorted(missing_vars))}',
                        'suggestion': 'Se usarán valores por defecto para las variables faltantes',
                        'details': sorted(missing_vars)
                    })
                    print(f"[WARNING] Variables faltantes: {sorted(missing_vars)}")
                
                # Analizar variables extra en el contexto
                extra_vars = [var for var in context.keys() if var not in required_vars and required_vars]
                if extra_vars and required_vars:  # Solo reportar si hay variables requeridas
                    result['info'].append({
                        'type': 'extra_context_vars',
                        'message': f'Variables adicionales en el contexto: {len(extra_vars)}',
                        'details': sorted(extra_vars)
                    })
                    print(f"[DEBUG] Variables extra en contexto: {len(extra_vars)}")
                
            except Exception as e:
                result['warnings'].append({
                    'type': 'template_analysis_error',
                    'message': f'No se pudieron analizar las variables de la plantilla: {str(e)}',
                    'suggestion': 'Se continuará con el contexto proporcionado'
                })
                print(f"[WARNING] Error analizando variables de plantilla: {e}")
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error en análisis de variables de plantilla: {e}")
            return {
                'warnings': [{
                    'type': 'variable_analysis_error',
                    'message': f'Error analizando variables: {str(e)}',
                    'suggestion': 'Continúe con precaución'
                }]
            }

    def _validate_context_comprehensive(self, context):
        """
        Realiza validación exhaustiva del contexto.
        
        Args:
            context: Contexto del documento
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            result = {'warnings': [], 'info': []}
            
            # Validar estructura básica del contexto
            basic_validation = self._validate_context_data_types(context)
            result['warnings'].extend(basic_validation)
            
            # Validar campos críticos específicos
            critical_validation = self._validate_critical_context_fields(context)
            result['warnings'].extend(critical_validation.get('warnings', []))
            result['info'].extend(critical_validation.get('info', []))
            
            # Validar listas de partes en detalle
            parties_validation = self._validate_parties_in_context(context)
            result['warnings'].extend(parties_validation.get('warnings', []))
            result['info'].extend(parties_validation.get('info', []))
            
            # Validar datos monetarios y numéricos
            numeric_validation = self._validate_numeric_context_fields(context)
            result['warnings'].extend(numeric_validation.get('warnings', []))
            
            print(f"[DEBUG] Validación de contexto completada - {len(result['warnings'])} advertencias")
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error en validación exhaustiva de contexto: {e}")
            return {
                'warnings': [{
                    'type': 'comprehensive_context_validation_error',
                    'message': f'Error en validación de contexto: {str(e)}',
                    'suggestion': 'Revise manualmente el contexto'
                }]
            }

    def _validate_critical_context_fields(self, context):
        """
        Valida campos críticos del contexto.
        
        Args:
            context: Contexto del documento
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            result = {'warnings': [], 'info': []}
            
            # Campos críticos con sus validaciones
            critical_fields = {
                'NUMERO_EXPEDIENTE': {'type': str, 'required': True, 'min_length': 1},
                'CARATULA': {'type': str, 'required': True, 'min_length': 5},
                'DIA_ACUERDO': {'type': int, 'required': True, 'min_value': 1, 'max_value': 31},
                'MES_ACUERDO': {'type': str, 'required': True, 'min_length': 3},
                'AÑO_ACUERDO': {'type': int, 'required': True, 'min_value': 2000, 'max_value': 2100},
                'ACTORES': {'type': list, 'required': True, 'min_length': 1},
                'DEMANDADOS': {'type': list, 'required': True, 'min_length': 1},
                'MONTO_COMPENSACION_NUMEROS': {'type': str, 'required': True, 'min_length': 1}
            }
            
            for field, validation in critical_fields.items():
                if field in context:
                    value = context[field]
                    
                    # Validar tipo
                    if not isinstance(value, validation['type']):
                        result['warnings'].append({
                            'type': 'wrong_critical_field_type',
                            'message': f'Campo crítico {field} tiene tipo incorrecto: {type(value).__name__} (esperado: {validation["type"].__name__})',
                            'suggestion': f'Corrija el tipo de {field}'
                        })
                        continue
                    
                    # Validaciones específicas por tipo
                    if validation['type'] == str:
                        if 'min_length' in validation and len(value.strip()) < validation['min_length']:
                            result['warnings'].append({
                                'type': 'critical_field_too_short',
                                'message': f'Campo {field} demasiado corto: {len(value)} caracteres (mínimo: {validation["min_length"]})',
                                'suggestion': f'Proporcione un valor más completo para {field}'
                            })
                    
                    elif validation['type'] == int:
                        if 'min_value' in validation and value < validation['min_value']:
                            result['warnings'].append({
                                'type': 'critical_field_value_too_low',
                                'message': f'Campo {field} valor demasiado bajo: {value} (mínimo: {validation["min_value"]})',
                                'suggestion': f'Verifique el valor de {field}'
                            })
                        if 'max_value' in validation and value > validation['max_value']:
                            result['warnings'].append({
                                'type': 'critical_field_value_too_high',
                                'message': f'Campo {field} valor demasiado alto: {value} (máximo: {validation["max_value"]})',
                                'suggestion': f'Verifique el valor de {field}'
                            })
                    
                    elif validation['type'] == list:
                        if 'min_length' in validation and len(value) < validation['min_length']:
                            result['warnings'].append({
                                'type': 'critical_list_too_short',
                                'message': f'Lista {field} demasiado corta: {len(value)} elementos (mínimo: {validation["min_length"]})',
                                'suggestion': f'Agregue más elementos a {field}'
                            })
                
                elif validation.get('required', False):
                    result['warnings'].append({
                        'type': 'missing_critical_field',
                        'message': f'Campo crítico faltante: {field}',
                        'suggestion': f'Agregue el campo {field} al contexto'
                    })
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error validando campos críticos: {e}")
            return {'warnings': [], 'info': []}

    def _validate_parties_in_context(self, context):
        """
        Valida las listas de partes en el contexto.
        
        Args:
            context: Contexto del documento
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            result = {'warnings': [], 'info': []}
            
            for lista_field in ['ACTORES', 'DEMANDADOS']:
                if lista_field in context and isinstance(context[lista_field], list):
                    partes = context[lista_field]
                    
                    result['info'].append({
                        'type': 'parties_count',
                        'message': f'{lista_field}: {len(partes)} partes encontradas'
                    })
                    
                    for i, parte in enumerate(partes):
                        if not isinstance(parte, dict):
                            result['warnings'].append({
                                'type': 'invalid_party_structure',
                                'message': f'{lista_field}[{i}] no es un diccionario válido',
                                'suggestion': f'Verifique la estructura de datos de {lista_field}'
                            })
                            continue
                        
                        # Validar campos requeridos de la parte
                        required_party_fields = ['nombre_completo', 'rol_id']
                        for field in required_party_fields:
                            if field not in parte:
                                result['warnings'].append({
                                    'type': 'missing_party_field',
                                    'message': f'{lista_field}[{i}] no tiene campo {field}',
                                    'suggestion': f'Agregue el campo {field} a todas las partes'
                                })
                            elif not parte[field]:
                                result['warnings'].append({
                                    'type': 'empty_party_field',
                                    'message': f'{lista_field}[{i}] tiene campo {field} vacío',
                                    'suggestion': f'Proporcione un valor para {field}'
                                })
                        
                        # Validar representantes si existen
                        if 'representantes' in parte:
                            representantes = parte['representantes']
                            if isinstance(representantes, list):
                                result['info'].append({
                                    'type': 'representatives_count',
                                    'message': f'{lista_field}[{i}] tiene {len(representantes)} representantes'
                                })
                            else:
                                result['warnings'].append({
                                    'type': 'invalid_representatives_structure',
                                    'message': f'{lista_field}[{i}] representantes no es una lista',
                                    'suggestion': 'Los representantes deben ser una lista'
                                })
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error validando partes en contexto: {e}")
            return {'warnings': [], 'info': []}

    def _validate_numeric_context_fields(self, context):
        """
        Valida campos numéricos y monetarios en el contexto.
        
        Args:
            context: Contexto del documento
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            result = {'warnings': []}
            
            # Campos numéricos a validar
            numeric_fields = {
                'MONTO_COMPENSACION_NUMEROS': {'type': 'monetary', 'required': True},
                'PLAZO_PAGO_DIAS': {'type': 'period', 'required': True},
                'DIA_ACUERDO': {'type': 'day', 'required': True},
                'AÑO_ACUERDO': {'type': 'year', 'required': True}
            }
            
            for field, validation in numeric_fields.items():
                if field in context:
                    value = context[field]
                    
                    if validation['type'] == 'monetary':
                        # Validar formato monetario
                        if isinstance(value, str):
                            if not value.strip():
                                result['warnings'].append({
                                    'type': 'empty_monetary_field',
                                    'message': f'Campo monetario {field} está vacío',
                                    'suggestion': 'Proporcione un monto válido'
                                })
                            elif not any(c.isdigit() for c in value):
                                result['warnings'].append({
                                    'type': 'invalid_monetary_format',
                                    'message': f'Campo monetario {field} no contiene dígitos: {value}',
                                    'suggestion': 'Proporcione un monto numérico válido'
                                })
                    
                    elif validation['type'] == 'period':
                        # Validar formato de período
                        if isinstance(value, str):
                            if not value.strip():
                                result['warnings'].append({
                                    'type': 'empty_period_field',
                                    'message': f'Campo de plazo {field} está vacío',
                                    'suggestion': 'Proporcione un plazo válido'
                                })
                    
                    elif validation['type'] == 'day':
                        # Validar día
                        if isinstance(value, int):
                            if not (1 <= value <= 31):
                                result['warnings'].append({
                                    'type': 'invalid_day_value',
                                    'message': f'Día inválido: {value}',
                                    'suggestion': 'El día debe estar entre 1 y 31'
                                })
                    
                    elif validation['type'] == 'year':
                        # Validar año
                        if isinstance(value, int):
                            current_year = datetime.date.today().year
                            if not (current_year - 10 <= value <= current_year + 10):
                                result['warnings'].append({
                                    'type': 'suspicious_year_value',
                                    'message': f'Año sospechoso: {value}',
                                    'suggestion': 'Verifique que el año sea correcto'
                                })
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error validando campos numéricos: {e}")
            return {'warnings': []}

    def _validate_template_context_compatibility(self, doc, context):
        """
        Valida la compatibilidad entre la plantilla y el contexto.
        
        Args:
            doc: Objeto DocxTemplate
            context: Contexto del documento
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            result = {'warnings': []}
            
            # Intentar una renderización de prueba (sin guardar)
            try:
                # Crear una copia del contexto para la prueba
                test_context = context.copy()
                
                # Preparar contexto seguro para la prueba
                required_vars = doc.get_undeclared_template_variables()
                safe_test_context = self._prepare_safe_context(test_context, required_vars)
                
                # Intentar renderización de prueba
                print("[DEBUG] Realizando prueba de renderización...")
                
                # Crear una copia temporal del documento para la prueba
                import tempfile
                import shutil
                
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_path = temp_file.name
                
                try:
                    # Copiar la plantilla al archivo temporal
                    shutil.copy2(doc.docx.part.package.part_related_by(doc.docx.part.package.main_document_part.rels['rId1']).blob, temp_path)
                    
                    # Crear nuevo objeto DocxTemplate con el archivo temporal
                    from docxtpl import DocxTemplate
                    test_doc = DocxTemplate(temp_path)
                    
                    # Intentar renderización
                    test_doc.render(safe_test_context)
                    
                    print("[DEBUG] Prueba de renderización exitosa")
                    
                except Exception as render_error:
                    result['warnings'].append({
                        'type': 'render_test_failed',
                        'message': f'Prueba de renderización falló: {str(render_error)}',
                        'suggestion': 'Puede haber problemas de compatibilidad entre plantilla y contexto'
                    })
                    print(f"[WARNING] Prueba de renderización falló: {render_error}")
                
                finally:
                    # Limpiar archivo temporal
                    try:
                        os.unlink(temp_path)
                    except:
                        pass
                
            except Exception as e:
                result['warnings'].append({
                    'type': 'compatibility_test_error',
                    'message': f'No se pudo realizar prueba de compatibilidad: {str(e)}',
                    'suggestion': 'Proceda con precaución'
                })
                print(f"[WARNING] Error en prueba de compatibilidad: {e}")
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error validando compatibilidad: {e}")
            return {'warnings': []}

    def _validate_context_data_types(self, context):
        """
        Valida los tipos de datos en el contexto para detectar posibles problemas.
        
        Args:
            context (dict): Contexto a validar
            
        Returns:
            list: Lista de issues encontrados
        """
        issues = []
        
        try:
            # Validaciones específicas para campos críticos
            critical_fields = {
                'NUMERO_EXPEDIENTE': str,
                'CARATULA': str,
                'DIA_ACUERDO': int,
                'MES_ACUERDO': str,
                'AÑO_ACUERDO': int,
                'ACTORES': list,
                'DEMANDADOS': list
            }
            
            for field, expected_type in critical_fields.items():
                if field in context:
                    value = context[field]
                    if not isinstance(value, expected_type):
                        issues.append({
                            'type': 'wrong_data_type',
                            'message': f'Campo {field} tiene tipo incorrecto: {type(value).__name__} (esperado: {expected_type.__name__})',
                            'suggestion': f'Verifique que {field} sea del tipo correcto'
                        })
            
            # Validar listas de partes
            for lista_field in ['ACTORES', 'DEMANDADOS']:
                if lista_field in context and isinstance(context[lista_field], list):
                    for i, parte in enumerate(context[lista_field]):
                        if not isinstance(parte, dict):
                            issues.append({
                                'type': 'invalid_party_data',
                                'message': f'{lista_field}[{i}] no es un diccionario válido',
                                'suggestion': f'Verifique la estructura de datos de {lista_field}'
                            })
                        elif 'nombre_completo' not in parte:
                            issues.append({
                                'type': 'missing_party_name',
                                'message': f'{lista_field}[{i}] no tiene nombre_completo',
                                'suggestion': 'Todas las partes deben tener un nombre completo'
                            })
            
            return issues
            
        except Exception as e:
            return [{
                'type': 'context_validation_error',
                'message': f'Error validando contexto: {str(e)}',
                'suggestion': 'Revise manualmente la estructura del contexto'
            }]

    def _prepare_safe_context(self, context, template_required_vars=None):
        """
        Prepara un contexto seguro agregando valores por defecto para variables faltantes.
        Implementa validación exhaustiva y corrección de tipos de datos.
        
        Args:
            context (dict): Contexto original
            template_required_vars (list): Variables requeridas por la plantilla
            
        Returns:
            dict: Contexto seguro con valores por defecto y tipos corregidos
        """
        print("[DEBUG] Preparando contexto seguro para renderización")
        
        try:
            # Validar entrada
            if not isinstance(context, dict):
                print(f"[ERROR] Contexto no es un diccionario: {type(context)}")
                context = {}
            
            safe_context = context.copy()
            
            # Obtener valores por defecto completos
            default_values = self._get_comprehensive_default_values()
            
            # Procesar variables requeridas por la plantilla
            if template_required_vars:
                safe_context = self._add_missing_template_variables(safe_context, template_required_vars, default_values)
            
            # Validar y corregir tipos de datos
            safe_context = self._validate_and_correct_data_types(safe_context, default_values)
            
            # Validar y limpiar listas de partes
            safe_context = self._validate_and_clean_parties_lists(safe_context)
            
            # Validar y corregir campos numéricos
            safe_context = self._validate_and_correct_numeric_fields(safe_context)
            
            # Validar y corregir campos de fecha
            safe_context = self._validate_and_correct_date_fields(safe_context)
            
            # Agregar variables auxiliares necesarias
            safe_context = self._add_auxiliary_variables(safe_context)
            
            print(f"[DEBUG] Contexto seguro preparado con {len(safe_context)} variables")
            return safe_context
            
        except Exception as e:
            print(f"[ERROR] Error crítico preparando contexto seguro: {e}")
            # Fallback: retornar contexto básico mínimo
            return self._create_minimal_safe_context()

    def _get_comprehensive_default_values(self):
        """
        Obtiene valores por defecto completos para todas las variables conocidas.
        
        Returns:
            dict: Diccionario con valores por defecto
        """
        try:
            current_date = datetime.date.today()
            
            # Obtener mes en español
            meses_espanol = [
                'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
            ]
            mes_actual = meses_espanol[current_date.month - 1]
            
            return {
                # Información del caso
                'NUMERO_EXPEDIENTE': 'SIN NÚMERO',
                'CARATULA': 'SIN CARÁTULA',
                
                # Fecha del acuerdo
                'DIA_ACUERDO': current_date.day,
                'MES_ACUERDO': mes_actual,
                'AÑO_ACUERDO': current_date.year,
                
                # Partes del caso
                'ACTORES': [],
                'DEMANDADOS': [],
                'NOMBRE_ACTOR': 'SIN ACTOR',
                
                # Montos y compensaciones
                'MONTO_COMPENSACION_NUMEROS': '0',
                'MONTO_COMPENSACION_LETRAS': 'CERO',
                'PLAZO_PAGO_DIAS': '0',
                'PLAZO_PAGO_LETRAS': 'CERO',
                
                # Datos bancarios
                'BANCO_ACTOR': '',
                'CBU_ACTOR': '',
                'ALIAS_ACTOR': '',
                'CUIT_ACTOR': '',
                
                # Honorarios (compatibilidad completa con plantilla)
                'MONTO_HONORARIOS_NUMEROS': '0',
                'MONTO_HONORARIOS_LETRAS': 'CERO',
                'HONORARIOS_LETRAS': 'CERO',  # Variable específica de la plantilla existente
                'PLAZO_PAGO_HONORARIOS_DIAS': '0',
                
                # Variables auxiliares
                'rep': {},
                
                # Variables adicionales comunes
                'LUGAR_ACUERDO': 'Ciudad Autónoma de Buenos Aires',
                'MEDIADOR': 'Mediador designado',
                'CENTRO_MEDIACION': 'Centro de Mediación',
                'OBSERVACIONES': '',
                'CLAUSULAS_ADICIONALES': '',
                
                # Variables de formato
                'FECHA_COMPLETA': f"{current_date.day} de {mes_actual} de {current_date.year}",
                'FECHA_CORTA': current_date.strftime('%d/%m/%Y')
            }
            
        except Exception as e:
            print(f"[ERROR] Error obteniendo valores por defecto: {e}")
            return {
                'NUMERO_EXPEDIENTE': 'ERROR',
                'CARATULA': 'ERROR',
                'ACTORES': [],
                'DEMANDADOS': []
            }

    def _add_missing_template_variables(self, safe_context, template_required_vars, default_values):
        """
        Agrega variables faltantes requeridas por la plantilla.
        
        Args:
            safe_context: Contexto actual
            template_required_vars: Variables requeridas
            default_values: Valores por defecto
            
        Returns:
            dict: Contexto con variables agregadas
        """
        try:
            variables_agregadas = 0
            
            for var in template_required_vars:
                if var not in safe_context:
                    if var in default_values:
                        safe_context[var] = default_values[var]
                        print(f"[DEBUG] Agregado valor por defecto para {var}: {default_values[var]}")
                        variables_agregadas += 1
                    else:
                        # Intentar inferir el tipo de variable por su nombre
                        inferred_value = self._infer_variable_default_value(var)
                        safe_context[var] = inferred_value
                        print(f"[DEBUG] Agregado valor inferido para variable desconocida {var}: {inferred_value}")
                        variables_agregadas += 1
            
            if variables_agregadas > 0:
                print(f"[DEBUG] Se agregaron {variables_agregadas} variables faltantes")
            
            return safe_context
            
        except Exception as e:
            print(f"[ERROR] Error agregando variables faltantes: {e}")
            return safe_context

    def _infer_variable_default_value(self, variable_name):
        """
        Infiere un valor por defecto basado en el nombre de la variable.
        
        Args:
            variable_name: Nombre de la variable
            
        Returns:
            Valor por defecto inferido
        """
        try:
            var_lower = variable_name.lower()
            
            # Patrones para diferentes tipos de variables
            if any(pattern in var_lower for pattern in ['fecha', 'date', 'dia', 'mes', 'año', 'year']):
                if 'dia' in var_lower or 'day' in var_lower:
                    return datetime.date.today().day
                elif 'mes' in var_lower or 'month' in var_lower:
                    meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                            'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
                    return meses[datetime.date.today().month - 1]
                elif 'año' in var_lower or 'year' in var_lower:
                    return datetime.date.today().year
                else:
                    return datetime.date.today().strftime('%d/%m/%Y')
            
            elif any(pattern in var_lower for pattern in ['monto', 'amount', 'precio', 'valor', 'suma']):
                return '0'
            
            elif any(pattern in var_lower for pattern in ['letras', 'words', 'texto']):
                return 'CERO'
            
            elif any(pattern in var_lower for pattern in ['plazo', 'dias', 'period', 'days']):
                return '0'
            
            elif any(pattern in var_lower for pattern in ['nombre', 'name']):
                return 'SIN ESPECIFICAR'
            
            elif any(pattern in var_lower for pattern in ['lista', 'list', 'array']):
                return []
            
            elif any(pattern in var_lower for pattern in ['numero', 'number', 'id']):
                return '0'
            
            elif any(pattern in var_lower for pattern in ['banco', 'bank', 'cbu', 'cuenta']):
                return ''
            
            else:
                # Valor por defecto genérico
                return ''
                
        except Exception as e:
            print(f"[WARNING] Error infiriendo valor por defecto para {variable_name}: {e}")
            return ''

    def _validate_and_correct_data_types(self, safe_context, default_values):
        """
        Valida y corrige tipos de datos en el contexto.
        
        Args:
            safe_context: Contexto a validar
            default_values: Valores por defecto con tipos correctos
            
        Returns:
            dict: Contexto con tipos corregidos
        """
        try:
            correcciones = 0
            
            for key, value in safe_context.items():
                if key in default_values:
                    expected_type = type(default_values[key])
                    
                    if not isinstance(value, expected_type):
                        # Intentar conversión de tipo
                        try:
                            if expected_type == str:
                                safe_context[key] = str(value) if value is not None else ''
                            elif expected_type == int:
                                if isinstance(value, str):
                                    # Extraer números del string
                                    numeric_part = ''.join(filter(str.isdigit, value))
                                    safe_context[key] = int(numeric_part) if numeric_part else default_values[key]
                                else:
                                    safe_context[key] = int(value)
                            elif expected_type == list:
                                if not isinstance(value, list):
                                    safe_context[key] = default_values[key]
                            elif expected_type == dict:
                                if not isinstance(value, dict):
                                    safe_context[key] = default_values[key]
                            
                            correcciones += 1
                            print(f"[DEBUG] Corregido tipo de {key}: {type(value).__name__} -> {expected_type.__name__}")
                            
                        except (ValueError, TypeError) as e:
                            print(f"[WARNING] No se pudo convertir {key}, usando valor por defecto: {e}")
                            safe_context[key] = default_values[key]
                            correcciones += 1
            
            if correcciones > 0:
                print(f"[DEBUG] Se corrigieron {correcciones} tipos de datos")
            
            return safe_context
            
        except Exception as e:
            print(f"[ERROR] Error validando y corrigiendo tipos de datos: {e}")
            return safe_context

    def _validate_and_clean_parties_lists(self, safe_context):
        """
        Valida y limpia las listas de partes.
        
        Args:
            safe_context: Contexto a validar
            
        Returns:
            dict: Contexto con listas de partes validadas
        """
        try:
            for lista_field in ['ACTORES', 'DEMANDADOS']:
                if lista_field in safe_context:
                    partes = safe_context[lista_field]
                    
                    if not isinstance(partes, list):
                        print(f"[WARNING] {lista_field} no es una lista, corrigiendo")
                        safe_context[lista_field] = []
                        continue
                    
                    # Validar cada parte en la lista
                    partes_validas = []
                    for i, parte in enumerate(partes):
                        if isinstance(parte, dict):
                            # Asegurar campos mínimos
                            parte_limpia = {
                                'nombre_completo': parte.get('nombre_completo', f'Parte {i+1}'),
                                'rol_id': parte.get('rol_id', i+1),
                                'rol_principal': parte.get('rol_principal', lista_field[:-1].title()),  # ACTORES -> Actor
                                'representantes': parte.get('representantes', [])
                            }
                            
                            # Validar representantes
                            if not isinstance(parte_limpia['representantes'], list):
                                parte_limpia['representantes'] = []
                            
                            partes_validas.append(parte_limpia)
                        else:
                            print(f"[WARNING] Parte inválida en {lista_field}[{i}], omitiendo")
                    
                    safe_context[lista_field] = partes_validas
                    print(f"[DEBUG] {lista_field}: {len(partes_validas)} partes válidas")
            
            return safe_context
            
        except Exception as e:
            print(f"[ERROR] Error validando listas de partes: {e}")
            return safe_context

    def _validate_and_correct_numeric_fields(self, safe_context):
        """
        Valida y corrige campos numéricos.
        
        Args:
            safe_context: Contexto a validar
            
        Returns:
            dict: Contexto con campos numéricos corregidos
        """
        try:
            numeric_fields = [
                'MONTO_COMPENSACION_NUMEROS', 'PLAZO_PAGO_DIAS', 
                'MONTO_HONORARIOS_NUMEROS', 'PLAZO_PAGO_HONORARIOS_DIAS'
            ]
            
            for field in numeric_fields:
                if field in safe_context:
                    value = safe_context[field]
                    
                    # Asegurar que sea string para campos numéricos de texto
                    if not isinstance(value, str):
                        safe_context[field] = str(value) if value is not None else '0'
                    
                    # Validar que contenga al menos un dígito
                    if not any(c.isdigit() for c in safe_context[field]):
                        safe_context[field] = '0'
                        print(f"[DEBUG] Corregido campo numérico vacío: {field} = '0'")
            
            return safe_context
            
        except Exception as e:
            print(f"[ERROR] Error validando campos numéricos: {e}")
            return safe_context

    def _validate_and_correct_date_fields(self, safe_context):
        """
        Valida y corrige campos de fecha.
        
        Args:
            safe_context: Contexto a validar
            
        Returns:
            dict: Contexto con campos de fecha corregidos
        """
        try:
            current_date = datetime.date.today()
            
            # Validar día
            if 'DIA_ACUERDO' in safe_context:
                dia = safe_context['DIA_ACUERDO']
                if not isinstance(dia, int) or not (1 <= dia <= 31):
                    safe_context['DIA_ACUERDO'] = current_date.day
                    print(f"[DEBUG] Corregido día del acuerdo: {dia} -> {current_date.day}")
            
            # Validar año
            if 'AÑO_ACUERDO' in safe_context:
                año = safe_context['AÑO_ACUERDO']
                if not isinstance(año, int) or not (2000 <= año <= 2100):
                    safe_context['AÑO_ACUERDO'] = current_date.year
                    print(f"[DEBUG] Corregido año del acuerdo: {año} -> {current_date.year}")
            
            # Validar mes
            if 'MES_ACUERDO' in safe_context:
                mes = safe_context['MES_ACUERDO']
                meses_validos = [
                    'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                    'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
                ]
                if not isinstance(mes, str) or mes.lower() not in meses_validos:
                    safe_context['MES_ACUERDO'] = meses_validos[current_date.month - 1]
                    print(f"[DEBUG] Corregido mes del acuerdo: {mes} -> {safe_context['MES_ACUERDO']}")
            
            return safe_context
            
        except Exception as e:
            print(f"[ERROR] Error validando campos de fecha: {e}")
            return safe_context

    def _add_auxiliary_variables(self, safe_context):
        """
        Agrega variables auxiliares necesarias para la plantilla.
        
        Args:
            safe_context: Contexto actual
            
        Returns:
            dict: Contexto con variables auxiliares
        """
        try:
            # Variable auxiliar 'rep' para representantes (requerida por algunas plantillas)
            if 'rep' not in safe_context:
                safe_context['rep'] = {}
            
            # Asegurar que NOMBRE_ACTOR esté sincronizado con la lista de actores
            if 'ACTORES' in safe_context and safe_context['ACTORES']:
                primer_actor = safe_context['ACTORES'][0]
                if isinstance(primer_actor, dict) and 'nombre_completo' in primer_actor:
                    safe_context['NOMBRE_ACTOR'] = primer_actor['nombre_completo']
            
            # Agregar fecha completa formateada si no existe
            if 'FECHA_COMPLETA' not in safe_context:
                dia = safe_context.get('DIA_ACUERDO', datetime.date.today().day)
                mes = safe_context.get('MES_ACUERDO', 'enero')
                año = safe_context.get('AÑO_ACUERDO', datetime.date.today().year)
                safe_context['FECHA_COMPLETA'] = f"{dia} de {mes} de {año}"
            
            return safe_context
            
        except Exception as e:
            print(f"[ERROR] Error agregando variables auxiliares: {e}")
            return safe_context

    def _create_minimal_safe_context(self):
        """
        Crea un contexto mínimo seguro como fallback.
        
        Returns:
            dict: Contexto mínimo funcional
        """
        try:
            current_date = datetime.date.today()
            meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                    'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
            
            return {
                'NUMERO_EXPEDIENTE': 'ERROR - CONTEXTO MÍNIMO',
                'CARATULA': 'ERROR - CONTEXTO MÍNIMO',
                'DIA_ACUERDO': current_date.day,
                'MES_ACUERDO': meses[current_date.month - 1],
                'AÑO_ACUERDO': current_date.year,
                'ACTORES': [],
                'DEMANDADOS': [],
                'NOMBRE_ACTOR': 'ERROR',
                'MONTO_COMPENSACION_NUMEROS': '0',
                'MONTO_COMPENSACION_LETRAS': 'CERO',
                'PLAZO_PAGO_DIAS': '0',
                'PLAZO_PAGO_LETRAS': 'CERO',
                'rep': {}
            }
            
        except Exception as e:
            print(f"[ERROR] Error creando contexto mínimo: {e}")
            return {'ERROR': 'CONTEXTO NO DISPONIBLE'}

    def _show_validation_errors(self, validation_result):
        """
        Muestra los errores de validación al usuario de manera clara y útil.
        """
        if not validation_result['errors'] and not validation_result['warnings']:
            return
        
        error_message = "Se encontraron problemas para generar acuerdos de mediación:\n\n"
        
        # Mostrar errores críticos
        if validation_result['errors']:
            error_message += "ERRORES CRÍTICOS:\n"
            for i, error in enumerate(validation_result['errors'], 1):
                error_message += f"{i}. {error['message']}\n"
                error_message += f"   Solución: {error['suggestion']}\n\n"
        
        # Mostrar advertencias
        if validation_result['warnings']:
            error_message += "ADVERTENCIAS:\n"
            for i, warning in enumerate(validation_result['warnings'], 1):
                error_message += f"{i}. {warning['message']}\n"
                error_message += f"   Recomendación: {warning['suggestion']}\n\n"
        
        if validation_result['errors']:
            messagebox.showerror(
                "Error - Dependencias Faltantes", 
                error_message,
                parent=self.app_controller.root
            )
        else:
            messagebox.showwarning(
                "Advertencia - Problemas Menores", 
                error_message,
                parent=self.app_controller.root
            )

    def open_case_dialog(self, case_id=None):
        """Abre el diálogo para agregar o editar un caso"""
        is_edit = case_id is not None
        client_context_id = None
        client_context_name = "N/A"

        if is_edit:
            case_data = self.db.get_case_by_id(case_id)
            if not case_data:
                messagebox.showerror(
                    "Error",
                    "No se pudieron cargar los datos del caso.",
                    parent=self.app_controller.root,
                )
                return
            dialog_title = f"Editar Caso ID: {case_id}"
            client_context_id = case_data["cliente_id"]
            client_info = self.db.get_client_by_id(client_context_id)
            if client_info:
                client_context_name = client_info.get(
                    "nombre", f"ID {client_context_id}"
                )
        else:  # Nuevo caso
            if not self.app_controller.selected_client:
                messagebox.showwarning(
                    "Advertencia",
                    "Selecciona un cliente para agregarle un caso.",
                    parent=self.app_controller.root,
                )
                return
            client_context_id = self.app_controller.selected_client["id"]
            client_context_name = self.app_controller.selected_client.get(
                "nombre", f"ID {client_context_id}"
            )
            dialog_title = f"Agregar Caso para: {client_context_name}"
            case_data = {}  # Datos iniciales vacíos

        dialog = tk.Toplevel(self.app_controller.root)
        dialog.title(dialog_title)
        dialog.transient(self.app_controller.root)
        dialog.grab_set()
        dialog.resizable(True, True)

        # Geometría y centrado
        dialog_width = 580
        dialog_height = 680
        parent_x = self.app_controller.root.winfo_x()
        parent_y = self.app_controller.root.winfo_y()
        parent_width = self.app_controller.root.winfo_width()
        parent_height = self.app_controller.root.winfo_height()
        x_pos = parent_x + (parent_width - dialog_width) // 2
        y_pos = parent_y + (parent_height - dialog_height) // 2
        dialog.geometry(f"{dialog_width}x{dialog_height}+{x_pos}+{y_pos}")
        dialog.minsize(dialog_width - 80, dialog_height - 200)
        dialog.resizable(True, True)

        frame = ttk.Frame(dialog, padding="15")
        frame.pack(fill=tk.BOTH, expand=True)
        frame.columnconfigure(0, weight=0)
        frame.columnconfigure(
            1,
            weight=1
        )  # Columna para widgets de entrada (expandible)

        # Variables de Tkinter
        caratula_var = tk.StringVar(value=case_data.get("caratula", ""))
        num_exp_var = tk.StringVar(value=case_data.get("numero_expediente", ""))
        anio_car_var = tk.StringVar(value=case_data.get("anio_caratula", ""))
        juzgado_var = tk.StringVar(value=case_data.get("juzgado", ""))
        jurisdiccion_var = tk.StringVar(value=case_data.get("jurisdiccion", ""))
        notas_initial_case = case_data.get("notas", "")
        if notas_initial_case is None:
            notas_initial_case = ""
        ruta_var = tk.StringVar(value=case_data.get("ruta_carpeta", ""))
        inact_days_var = tk.IntVar(value=case_data.get("inactivity_threshold_days", 30))
        inact_enabled_var = tk.IntVar(value=case_data.get("inactivity_enabled", 1))
        etiquetas_caso_var = tk.StringVar()  # Nueva variable para etiquetas del caso

        if is_edit and case_id:  # Solo cargar si estamos editando un caso existente
            etiquetas_actuales_obj = self.db.get_etiquetas_de_caso(case_id)
            etiquetas_actuales_nombres = [
                e["nombre_etiqueta"] for e in etiquetas_actuales_obj
            ]
            etiquetas_caso_var.set(", ".join(etiquetas_actuales_nombres))

        row_idx = 0
        ttk.Label(frame, text="Cliente:").grid(
            row=0, column=0, sticky=tk.W, pady=3, padx=5
        )
        ttk.Label(frame, text=f"{client_context_name} (ID: {client_context_id})").grid(
            row=0, column=1, sticky=tk.W, pady=3, padx=5
        )
        row_idx += 1

        ttk.Label(frame, text="*Carátula:").grid(
            row=1, column=0, sticky=tk.W, pady=3, padx=5
        )
        caratula_entry = ttk.Entry(frame, textvariable=caratula_var)
        caratula_entry.grid(row=1, column=1, sticky=tk.EW, pady=3, padx=5)
        row_idx += 1

        # Nro Expediente y Año Carátula en la misma fila
        exp_anio_frame = ttk.Frame(frame)
        exp_anio_frame.grid(row=row_idx, column=1, sticky=tk.EW, pady=0, padx=0)

        ttk.Label(frame, text="Expediente:").grid(
            row=row_idx, column=0, sticky=tk.W, pady=3, padx=5
        )

        ttk.Label(exp_anio_frame, text="Nro:").pack(side=tk.LEFT, padx=(0, 2), pady=3)
        ttk.Entry(exp_anio_frame, textvariable=num_exp_var, width=15).pack(
            side=tk.LEFT, padx=(0, 10), pady=3
        )
        ttk.Label(exp_anio_frame, text="Año:").pack(side=tk.LEFT, padx=(0, 2), pady=3)
        ttk.Entry(exp_anio_frame, textvariable=anio_car_var, width=8).pack(
            side=tk.LEFT, padx=(0, 5), pady=3
        )
        row_idx += 1

        # Juzgado
        ttk.Label(frame, text="Juzgado:").grid(
            row=row_idx, column=0, sticky=tk.W, pady=3, padx=5
        )
        ttk.Entry(frame, textvariable=juzgado_var).grid(
            row=row_idx, column=1, sticky=tk.EW, pady=3, padx=5
        )
        row_idx += 1

        # Jurisdicción
        ttk.Label(frame, text="Jurisdicción:").grid(
            row=row_idx, column=0, sticky=tk.W, pady=3, padx=5
        )
        ttk.Entry(frame, textvariable=jurisdiccion_var).grid(
            row=row_idx, column=1, sticky=tk.EW, pady=3, padx=5
        )
        row_idx += 1

        row_idx += 1

        # Carpeta de Documentos
        '''ttk.Label(frame, text="Carpeta Documentos:").grid(
            row=row_idx, column=0, sticky=tk.W, pady=3, padx=5
        )
        folder_frame = ttk.Frame(frame)
        folder_frame.grid(row=row_idx, column=1, sticky=tk.EW, pady=3, padx=5)
        folder_frame.columnconfigure(0, weight=1)

        folder_entry = ttk.Entry(folder_frame, textvariable=ruta_var)
        folder_entry.grid(row=0, column=0, sticky=tk.EW, padx=(0, 5))
        browse_btn = ttk.Button(
            folder_frame, text="Examinar", command=lambda: self._browse_folder(ruta_var)
        )
        browse_btn.grid(row=0, column=1)
        row_idx += 1
        '''

        # Etiquetas del caso
        ttk.Label(frame, text="Etiquetas:").grid(
            row=row_idx, column=0, sticky=tk.W, pady=(10, 3), padx=5
        )
        ttk.Label(frame, text="(separadas por coma)").grid(
            row=row_idx, column=1, sticky=tk.E, pady=(10, 3), padx=0
        )
        row_idx += 1
        etiquetas_entry = ttk.Entry(frame, textvariable=etiquetas_caso_var)
        etiquetas_entry.grid(
            row=row_idx, column=0, columnspan=2, sticky=tk.EW, pady=3, padx=5
        )
        row_idx += 1

        # Configuración de inactividad
        inact_frame = ttk.LabelFrame(
            frame, text="Configuración de Inactividad", padding="10"
        )
        inact_frame.grid(
            row=row_idx, column=0, columnspan=2, sticky=tk.EW, pady=10, padx=5
        )
        inact_frame.columnconfigure(1, weight=1)

        ttk.Checkbutton(
            inact_frame,
            text="Habilitar alertas de inactividad",
            variable=inact_enabled_var,
        ).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=3)

        ttk.Label(inact_frame, text="Días umbral:").grid(
            row=1, column=0, sticky=tk.W, pady=3, padx=5
        )
        ttk.Spinbox(
            inact_frame, from_=1, to=365, textvariable=inact_days_var, width=10
        ).grid(row=1, column=1, sticky=tk.W, pady=3, padx=5)
        row_idx += 1

        # Notas
        ttk.Label(frame, text="Notas:").grid(
            row=row_idx, column=0, sticky=tk.NW, pady=3, padx=5
        )
        notas_frame = ttk.Frame(frame)
        notas_frame.grid(row=row_idx, column=1, sticky=tk.NSEW, pady=3, padx=5)
        notas_frame.rowconfigure(0, weight=1)
        notas_frame.columnconfigure(0, weight=1)

        notas_text = tk.Text(notas_frame, height=4, wrap=tk.WORD)
        notas_text.grid(row=0, column=0, sticky="nsew")
        notas_scroll = ttk.Scrollbar(
            notas_frame, orient=tk.VERTICAL, command=notas_text.yview
        )
        notas_scroll.grid(row=0, column=1, sticky="ns")
        notas_text["yscrollcommand"] = notas_scroll.set
        notas_text.insert("1.0", notas_initial_case)

        frame.rowconfigure(row_idx, weight=1)  # Notas expandibles
        row_idx += 1

        # Botones
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=row_idx, column=0, columnspan=2, pady=15)

        save_cmd = lambda: self.save_case(
            case_id,
            client_context_id,
            caratula_var.get(),
            num_exp_var.get(),
            anio_car_var.get(),
            juzgado_var.get(),
            jurisdiccion_var.get(),
            notas_text.get("1.0", tk.END).strip(),
            ruta_var.get(),
            inact_days_var.get(),
            inact_enabled_var.get(),
            etiquetas_caso_var.get(),
            dialog,
        )
        ttk.Button(button_frame, text="Guardar", command=save_cmd).pack(
            side=tk.LEFT, padx=5
        )
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(
            side=tk.LEFT, padx=5
        )

        caratula_entry.focus_set()
        self.app_controller.root.wait_window(dialog)

    def _ask_agreement_details_dialog(self, actor_principal=None):
        """
        Abre un diálogo para solicitar los detalles variables necesarios para el acuerdo de mediación.
        Implementa validación exhaustiva, manejo de errores robusto y experiencia de usuario mejorada.

        Args:
            actor_principal: Datos del actor principal para pre-rellenar información (opcional)

        Returns:
            dict: Diccionario con los datos validados o None si se cancela
        """
        print("[DEBUG] Abriendo diálogo de detalles del acuerdo con validación mejorada")

        try:
            # Validar entrada
            if not self._validate_agreement_dialog_prerequisites(actor_principal):
                return None

            # Crear y configurar el diálogo
            dialog = self._create_agreement_dialog()
            if not dialog:
                return None

            # Preparar datos iniciales
            initial_data = self._prepare_initial_agreement_data(actor_principal)

            # Crear variables de Tkinter con validación
            data_vars = self._create_agreement_data_variables(initial_data)

            # Crear interfaz del diálogo
            self._create_agreement_dialog_interface(dialog, data_vars)

            # Configurar validaciones y eventos
            self._setup_agreement_dialog_validation(dialog, data_vars)

            # Mostrar diálogo y esperar resultado
            result = self._show_agreement_dialog_and_wait(dialog, data_vars)

            return result

        except Exception as e:
            print(f"[ERROR] Error crítico en diálogo de detalles del acuerdo: {e}")
            self._show_agreement_dialog_error(e)
            return None

    def _validate_agreement_dialog_prerequisites(self, actor_principal):
        """
        Valida los prerequisitos para mostrar el diálogo de acuerdo.
        
        Args:
            actor_principal: Datos del actor principal
            
        Returns:
            bool: True si se pueden mostrar el diálogo
        """
        try:
            # Validar que tenemos acceso a la interfaz
            if not (hasattr(self, 'app_controller') and 
                   self.app_controller and 
                   hasattr(self.app_controller, 'root') and
                   self.app_controller.root):
                print("[ERROR] No hay acceso a la interfaz de usuario para el diálogo")
                return False
            
            # Validar actor_principal (puede ser None o dict)
            if actor_principal is not None and not isinstance(actor_principal, dict):
                print(f"[WARNING] Actor principal no es un diccionario válido: {type(actor_principal)}")
                # Continuar con datos vacíos
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Error validando prerequisitos del diálogo: {e}")
            return False

    def _create_agreement_dialog(self):
        """
        Crea y configura la ventana del diálogo.
        
        Returns:
            tk.Toplevel: Ventana del diálogo o None si hay error
        """
        try:
            dialog = tk.Toplevel(self.app_controller.root)
            dialog.title("Completar Datos del Acuerdo de Mediación")
            dialog.transient(self.app_controller.root)
            dialog.grab_set()
            dialog.resizable(True, True)
            
            # Configurar tamaño y posición
            dialog_width = 500
            dialog_height = 450
            
            # Centrar el diálogo
            screen_width = dialog.winfo_screenwidth()
            screen_height = dialog.winfo_screenheight()
            x = (screen_width - dialog_width) // 2
            y = (screen_height - dialog_height) // 2
            
            dialog.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")
            dialog.minsize(400, 350)
            
            # Configurar icono si está disponible
            try:
                dialog.iconbitmap(default=self.app_controller.root.iconbitmap())
            except:
                pass  # Ignorar si no hay icono
            
            return dialog
            
        except Exception as e:
            print(f"[ERROR] Error creando diálogo de acuerdo: {e}")
            return None

    def _prepare_initial_agreement_data(self, actor_principal=None):
        """
        Prepara los datos iniciales completos para el diálogo de acuerdo.

        Args:
            actor_principal: Datos del actor principal (opcional)

        Returns:
            dict: Datos iniciales preparados
        """
        try:
            # Datos por defecto completos
            default_data = {
                # Compensación
                "monto_compensacion_numeros": "",
                "plazo_pago_dias": "30",

                # Honorarios letrado del actor
                "monto_honorarios_letrado_actor": "",
                "plazo_pago_honorarios_dias": "30",

                # Honorarios mediador
                "monto_honorarios_mediador": "",
                "plazo_pago_mediador_dias": "15",

                # Datos bancarios del actor
                "banco_actor": "",
                "cbu_actor": "",
                "alias_actor": "",
                "cuit_actor": "",

                # Datos bancarios del letrado del actor
                "banco_letrado_actor": "",
                "cbu_letrado_actor": "",
                "alias_letrado_actor": "",
                "cuit_letrado_actor": ""
            }

            # Si tenemos datos del actor principal, usarlos para pre-rellenar
            if isinstance(actor_principal, dict):
                # Mapear campos del actor a campos del diálogo
                actor_field_mapping = {
                    'banco': 'banco_actor',
                    'cbu': 'cbu_actor',
                    'alias_cbu': 'alias_actor',
                    'cuit': 'cuit_actor'
                }

                for actor_field, dialog_field in actor_field_mapping.items():
                    if actor_field in actor_principal and actor_principal[actor_field]:
                        default_data[dialog_field] = str(actor_principal[actor_field]).strip()
                        print(f"[DEBUG] Pre-rellenado {dialog_field} con: {default_data[dialog_field]}")

            return default_data

        except Exception as e:
            print(f"[ERROR] Error preparando datos iniciales: {e}")
            return {
                "monto_compensacion_numeros": "",
                "plazo_pago_dias": "30",
                "monto_honorarios_letrado_actor": "",
                "plazo_pago_honorarios_dias": "30",
                "monto_honorarios_mediador": "",
                "plazo_pago_mediador_dias": "15",
                "banco_actor": "",
                "cbu_actor": "",
                "alias_actor": "",
                "cuit_actor": "",
                "banco_letrado_actor": "",
                "cbu_letrado_actor": "",
                "alias_letrado_actor": "",
                "cuit_letrado_actor": ""
            }

    def _create_agreement_data_variables(self, initial_data):
        """
        Crea las variables de Tkinter para el diálogo.
        
        Args:
            initial_data: Datos iniciales
            
        Returns:
            dict: Diccionario con variables de Tkinter
        """
        try:
            data_vars = {}
            
            for key, value in initial_data.items():
                data_vars[key] = tk.StringVar(value=value)
            
            return data_vars
            
        except Exception as e:
            print(f"[ERROR] Error creando variables de datos: {e}")
            return {}

    def _create_agreement_dialog_interface(self, dialog, data_vars):
        """
        Crea la interfaz del diálogo.
        
        Args:
            dialog: Ventana del diálogo
            data_vars: Variables de datos
        """
        try:
            # Frame principal con scroll
            main_frame = ttk.Frame(dialog)
            main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # Canvas y scrollbar para contenido scrolleable
            canvas = tk.Canvas(main_frame)
            scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )
            
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)
            
            # Configurar grid del frame scrolleable
            scrollable_frame.columnconfigure(1, weight=1)
            
            # Título del diálogo
            title_label = ttk.Label(scrollable_frame, text="Datos del Acuerdo de Mediación", 
                                   font=('TkDefaultFont', 12, 'bold'))
            title_label.grid(row=0, column=0, columnspan=2, pady=(0, 15), sticky=tk.W)
            
            # Crear campos del formulario
            self._create_agreement_form_fields(scrollable_frame, data_vars)
            
            # Crear botones
            self._create_agreement_dialog_buttons(scrollable_frame, dialog, data_vars)
            
            # Configurar scroll
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
            # Configurar scroll con mouse wheel
            def _on_mousewheel(event):
                canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            
            canvas.bind("<MouseWheel>", _on_mousewheel)
            
            # Almacenar referencias para uso posterior
            dialog.canvas = canvas
            dialog.scrollable_frame = scrollable_frame
            
        except Exception as e:
            print(f"[ERROR] Error creando interfaz del diálogo: {e}")

    def _create_agreement_form_fields(self, parent_frame, data_vars):
        """
        Crea los campos del formulario completo para acuerdos de mediación.
        
        Args:
            parent_frame: Frame padre
            data_vars: Variables de datos
        """
        try:
            row = 1
            
            # Sección: Datos Monetarios de Compensación
            monetary_label = ttk.Label(parent_frame, text="💰 Compensación al Actor", 
                                     font=('TkDefaultFont', 11, 'bold'))
            monetary_label.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(10, 5))
            row += 1
            
            # Monto de compensación
            ttk.Label(parent_frame, text="Monto de Compensación ($):*").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            monto_entry = ttk.Entry(parent_frame, textvariable=data_vars["monto_compensacion_numeros"], 
                                   font=('TkDefaultFont', 10))
            monto_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            monto_entry.focus()  # Foco inicial
            
            self._create_enhanced_tooltip(monto_entry, 
                "Monto que el demandado pagará al actor\n" +
                "Formato: 100000 o 100.000,50")
            
            row += 1
            
            # Plazo de pago compensación
            ttk.Label(parent_frame, text="Plazo de Pago (días):*").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            plazo_entry = ttk.Entry(parent_frame, textvariable=data_vars["plazo_pago_dias"],
                                   font=('TkDefaultFont', 10))
            plazo_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(plazo_entry, 
                "Días hábiles para el pago de la compensación\n" +
                "Valores comunes: 30, 60, 90 días")
            
            row += 1
            
            # Separador
            ttk.Separator(parent_frame, orient=tk.HORIZONTAL).grid(
                row=row, column=0, columnspan=2, sticky=tk.EW, pady=15)
            row += 1
            
            # Sección: Honorarios del Letrado del Actor
            honorarios_actor_label = ttk.Label(parent_frame, text="⚖️ Honorarios del Letrado del Actor", 
                                              font=('TkDefaultFont', 11, 'bold'))
            honorarios_actor_label.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(0, 5))
            row += 1
            
            # Monto honorarios letrado actor
            ttk.Label(parent_frame, text="Honorarios Letrado ($):*").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            honorarios_letrado_entry = ttk.Entry(parent_frame, textvariable=data_vars["monto_honorarios_letrado_actor"],
                                                font=('TkDefaultFont', 10))
            honorarios_letrado_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(honorarios_letrado_entry, 
                "Honorarios que el demandado pagará al letrado del actor\n" +
                "Formato: 50000 o 50.000,00")
            
            row += 1
            
            # Plazo honorarios letrado
            ttk.Label(parent_frame, text="Plazo Honorarios (días):*").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            plazo_honorarios_entry = ttk.Entry(parent_frame, textvariable=data_vars["plazo_pago_honorarios_dias"],
                                              font=('TkDefaultFont', 10))
            plazo_honorarios_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(plazo_honorarios_entry, 
                "Días hábiles para el pago de honorarios del letrado\n" +
                "Generalmente igual o mayor al plazo de compensación")
            
            row += 1
            
            # Separador
            ttk.Separator(parent_frame, orient=tk.HORIZONTAL).grid(
                row=row, column=0, columnspan=2, sticky=tk.EW, pady=15)
            row += 1
            
            # Sección: Honorarios del Mediador/Conciliador
            honorarios_mediador_label = ttk.Label(parent_frame, text="🤝 Honorarios del Mediador", 
                                                 font=('TkDefaultFont', 11, 'bold'))
            honorarios_mediador_label.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(0, 5))
            row += 1
            
            # Monto honorarios mediador
            ttk.Label(parent_frame, text="Honorarios Mediador ($):*").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            honorarios_mediador_entry = ttk.Entry(parent_frame, textvariable=data_vars["monto_honorarios_mediador"],
                                                 font=('TkDefaultFont', 10))
            honorarios_mediador_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(honorarios_mediador_entry, 
                "Honorarios del mediador/conciliador\n" +
                "Generalmente se divide entre las partes")
            
            row += 1
            
            # Plazo honorarios mediador
            ttk.Label(parent_frame, text="Plazo Mediador (días):*").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            plazo_mediador_entry = ttk.Entry(parent_frame, textvariable=data_vars["plazo_pago_mediador_dias"],
                                           font=('TkDefaultFont', 10))
            plazo_mediador_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(plazo_mediador_entry, 
                "Días hábiles para el pago de honorarios del mediador")
            
            row += 1
            
            # Separador
            ttk.Separator(parent_frame, orient=tk.HORIZONTAL).grid(
                row=row, column=0, columnspan=2, sticky=tk.EW, pady=15)
            row += 1
            
            # Sección: Datos Bancarios del Actor
            banking_actor_label = ttk.Label(parent_frame, text="🏦 Datos Bancarios del Actor", 
                                          font=('TkDefaultFont', 11, 'bold'))
            banking_actor_label.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(0, 5))
            row += 1
            
            # Banco actor
            ttk.Label(parent_frame, text="Banco Actor:").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            banco_actor_entry = ttk.Entry(parent_frame, textvariable=data_vars["banco_actor"],
                                        font=('TkDefaultFont', 10))
            banco_actor_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(banco_actor_entry, "Banco del actor para recibir la compensación")
            
            row += 1
            
            # CBU actor
            ttk.Label(parent_frame, text="CBU Actor:").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            cbu_actor_entry = ttk.Entry(parent_frame, textvariable=data_vars["cbu_actor"],
                                      font=('TkDefaultFont', 10))
            cbu_actor_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(cbu_actor_entry, "CBU del actor (22 dígitos)")
            
            row += 1
            
            # Alias actor
            ttk.Label(parent_frame, text="Alias Actor:").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            alias_actor_entry = ttk.Entry(parent_frame, textvariable=data_vars["alias_actor"],
                                        font=('TkDefaultFont', 10))
            alias_actor_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(alias_actor_entry, "Alias del CBU del actor")
            
            row += 1
            
            # CUIT actor
            ttk.Label(parent_frame, text="CUIT Actor:").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            cuit_actor_entry = ttk.Entry(parent_frame, textvariable=data_vars["cuit_actor"],
                                       font=('TkDefaultFont', 10))
            cuit_actor_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(cuit_actor_entry, "CUIT/CUIL del actor")
            
            row += 1
            
            # Separador
            ttk.Separator(parent_frame, orient=tk.HORIZONTAL).grid(
                row=row, column=0, columnspan=2, sticky=tk.EW, pady=15)
            row += 1
            
            # Sección: Datos Bancarios del Letrado del Actor
            banking_letrado_label = ttk.Label(parent_frame, text="⚖️ Datos Bancarios del Letrado del Actor", 
                                            font=('TkDefaultFont', 11, 'bold'))
            banking_letrado_label.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(0, 5))
            row += 1
            
            # Banco letrado
            ttk.Label(parent_frame, text="Banco Letrado:").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            banco_letrado_entry = ttk.Entry(parent_frame, textvariable=data_vars["banco_letrado_actor"],
                                          font=('TkDefaultFont', 10))
            banco_letrado_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(banco_letrado_entry, "Banco del letrado del actor para recibir honorarios")
            
            row += 1
            
            # CBU letrado
            ttk.Label(parent_frame, text="CBU Letrado:").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            cbu_letrado_entry = ttk.Entry(parent_frame, textvariable=data_vars["cbu_letrado_actor"],
                                        font=('TkDefaultFont', 10))
            cbu_letrado_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(cbu_letrado_entry, "CBU del letrado del actor (22 dígitos)")
            
            row += 1
            
            # Alias letrado
            ttk.Label(parent_frame, text="Alias Letrado:").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            alias_letrado_entry = ttk.Entry(parent_frame, textvariable=data_vars["alias_letrado_actor"],
                                          font=('TkDefaultFont', 10))
            alias_letrado_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(alias_letrado_entry, "Alias del CBU del letrado")
            
            row += 1
            
            # CUIT letrado
            ttk.Label(parent_frame, text="CUIT Letrado:").grid(
                row=row, column=0, sticky=tk.W, pady=3, padx=(10, 5))
            
            cuit_letrado_entry = ttk.Entry(parent_frame, textvariable=data_vars["cuit_letrado_actor"],
                                         font=('TkDefaultFont', 10))
            cuit_letrado_entry.grid(row=row, column=1, sticky=tk.EW, pady=3, padx=(5, 0))
            
            self._create_enhanced_tooltip(cuit_letrado_entry, "CUIT/CUIL del letrado del actor")
            
            row += 1
            
            # Nota sobre campos obligatorios
            note_label = ttk.Label(parent_frame, text="* Campos obligatorios", 
                                  font=('TkDefaultFont', 8), foreground='red')
            note_label.grid(row=row, column=0, columnspan=2, sticky=tk.W, pady=(15, 5))
            
            # Almacenar referencias de campos para validación
            parent_frame.monto_entry = monto_entry
            parent_frame.plazo_entry = plazo_entry
            parent_frame.honorarios_letrado_entry = honorarios_letrado_entry
            parent_frame.honorarios_mediador_entry = honorarios_mediador_entry
            parent_frame.cbu_actor_entry = cbu_actor_entry
            parent_frame.cbu_letrado_entry = cbu_letrado_entry
            parent_frame.cuit_actor_entry = cuit_actor_entry
            parent_frame.cuit_letrado_entry = cuit_letrado_entry
            
        except Exception as e:
            print(f"[ERROR] Error creando campos del formulario: {e}")

    def _create_agreement_dialog_buttons(self, parent_frame, dialog, data_vars):
        """
        Crea los botones del diálogo.
        
        Args:
            parent_frame: Frame padre
            dialog: Ventana del diálogo
            data_vars: Variables de datos
        """
        try:
            # Frame para botones
            button_frame = ttk.Frame(parent_frame)
            button_frame.grid(row=20, column=0, columnspan=2, pady=20, sticky=tk.EW)
            button_frame.columnconfigure(0, weight=1)
            button_frame.columnconfigure(1, weight=1)
            
            # Resultado del diálogo
            dialog.result = {}
            
            # Función de validación y guardado
            def validate_and_save():
                validation_result = self._validate_agreement_form_data(data_vars, dialog)
                if validation_result['valid']:
                    dialog.result = validation_result['data']
                    print(f"[DEBUG] Datos del acuerdo validados exitosamente")
                    dialog.destroy()
            
            # Función de cancelación
            def cancel():
                dialog.result = {}
                print("[DEBUG] Usuario canceló el diálogo de detalles del acuerdo")
                dialog.destroy()
            
            # Botones
            accept_button = ttk.Button(button_frame, text="Aceptar", command=validate_and_save)
            accept_button.grid(row=0, column=0, padx=(0, 5), sticky=tk.EW)
            
            cancel_button = ttk.Button(button_frame, text="Cancelar", command=cancel)
            cancel_button.grid(row=0, column=1, padx=(5, 0), sticky=tk.EW)
            
            # Configurar teclas de acceso rápido
            dialog.bind('<Return>', lambda e: validate_and_save())
            dialog.bind('<Escape>', lambda e: cancel())
            dialog.protocol("WM_DELETE_WINDOW", cancel)
            
            # Almacenar referencias
            dialog.validate_and_save = validate_and_save
            dialog.cancel = cancel
            
        except Exception as e:
            print(f"[ERROR] Error creando botones del diálogo: {e}")

    def _setup_agreement_dialog_validation(self, dialog, data_vars):
        """
        Configura la validación en tiempo real del diálogo.
        
        Args:
            dialog: Ventana del diálogo
            data_vars: Variables de datos
        """
        try:
            # Validación en tiempo real para campos numéricos
            def validate_numeric_field(value, field_type):
                """Valida campos numéricos en tiempo real"""
                if not value:
                    return True  # Permitir vacío temporalmente
                
                if field_type == 'monetary':
                    # Permitir dígitos, puntos, comas y espacios
                    allowed = set('0123456789., ')
                    return all(c in allowed for c in value)
                elif field_type == 'period':
                    # Solo dígitos para días
                    return value.isdigit()
                elif field_type == 'cbu':
                    # Solo dígitos y guiones para CBU
                    allowed = set('0123456789-')
                    return all(c in allowed for c in value)
                elif field_type == 'cuit':
                    # Solo dígitos y guiones para CUIT
                    allowed = set('0123456789-')
                    return all(c in allowed for c in value)
                
                return True
            
            # Registrar validaciones
            vcmd_monetary = (dialog.register(lambda v: validate_numeric_field(v, 'monetary')), '%P')
            vcmd_period = (dialog.register(lambda v: validate_numeric_field(v, 'period')), '%P')
            vcmd_cbu = (dialog.register(lambda v: validate_numeric_field(v, 'cbu')), '%P')
            vcmd_cuit = (dialog.register(lambda v: validate_numeric_field(v, 'cuit')), '%P')
            
            # Aplicar validaciones a los campos correspondientes
            if hasattr(dialog.scrollable_frame, 'monto_entry'):
                dialog.scrollable_frame.monto_entry.configure(validate='key', validatecommand=vcmd_monetary)
            
            if hasattr(dialog.scrollable_frame, 'plazo_entry'):
                dialog.scrollable_frame.plazo_entry.configure(validate='key', validatecommand=vcmd_period)
            
            if hasattr(dialog.scrollable_frame, 'cbu_entry'):
                dialog.scrollable_frame.cbu_entry.configure(validate='key', validatecommand=vcmd_cbu)
            
            if hasattr(dialog.scrollable_frame, 'cuit_entry'):
                dialog.scrollable_frame.cuit_entry.configure(validate='key', validatecommand=vcmd_cuit)
            
        except Exception as e:
            print(f"[ERROR] Error configurando validación del diálogo: {e}")

    def _show_agreement_dialog_and_wait(self, dialog, data_vars):
        """
        Muestra el diálogo y espera el resultado.
        
        Args:
            dialog: Ventana del diálogo
            data_vars: Variables de datos
            
        Returns:
            dict: Resultado del diálogo o None si se canceló
        """
        try:
            # Mostrar el diálogo y esperar
            self.app_controller.root.wait_window(dialog)
            
            # Retornar resultado
            result = getattr(dialog, 'result', {})
            
            if result:
                print(f"[DEBUG] Diálogo completado con éxito - Monto: ${result.get('monto_compensacion_numeros', '0')}")
                return result
            else:
                print("[DEBUG] Diálogo cancelado por el usuario")
                return None
                
        except Exception as e:
            print(f"[ERROR] Error mostrando diálogo: {e}")
            return None

    def _validate_agreement_form_data(self, data_vars, dialog):
        """
        Valida los datos del formulario de acuerdo.
        
        Args:
            data_vars: Variables de datos
            dialog: Ventana del diálogo
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            # Obtener valores actuales
            current_values = {}
            for key, var in data_vars.items():
                current_values[key] = var.get().strip()
            
            # Realizar validación exhaustiva
            validation_result = self._perform_comprehensive_agreement_validation(current_values)
            
            # Si hay errores, mostrarlos
            if not validation_result['valid']:
                self._show_agreement_validation_errors(validation_result['errors'], dialog)
                return {'valid': False, 'data': {}}
            
            # Si hay advertencias, mostrarlas pero continuar
            if validation_result.get('warnings'):
                self._show_agreement_validation_warnings(validation_result['warnings'], dialog)
            
            # Limpiar y formatear datos finales
            cleaned_data = self._clean_and_format_agreement_data(current_values)
            
            return {'valid': True, 'data': cleaned_data}
            
        except Exception as e:
            print(f"[ERROR] Error validando datos del formulario: {e}")
            messagebox.showerror(
                "Error de Validación",
                f"Error inesperado durante la validación: {str(e)}",
                parent=dialog
            )
            return {'valid': False, 'data': {}}

    def _perform_comprehensive_agreement_validation(self, values):
        """
        Realiza validación exhaustiva de los datos del acuerdo.
        
        Args:
            values: Valores a validar
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            errors = []
            warnings = []
            
            # Validar monto de compensación
            monto_validation = self._validate_compensation_amount(values.get('monto_compensacion_numeros', ''))
            if not monto_validation['valid']:
                errors.extend(monto_validation['errors'])
            if monto_validation.get('warnings'):
                warnings.extend(monto_validation['warnings'])
            
            # Validar plazo de pago
            plazo_validation = self._validate_payment_period(values.get('plazo_pago_dias', ''))
            if not plazo_validation['valid']:
                errors.extend(plazo_validation['errors'])
            if plazo_validation.get('warnings'):
                warnings.extend(plazo_validation['warnings'])
            
            # Validar datos bancarios (opcionales)
            banking_validation = self._validate_banking_data(values)
            if banking_validation.get('warnings'):
                warnings.extend(banking_validation['warnings'])
            
            return {
                'valid': len(errors) == 0,
                'errors': errors,
                'warnings': warnings
            }
            
        except Exception as e:
            print(f"[ERROR] Error en validación exhaustiva: {e}")
            return {
                'valid': False,
                'errors': [f"Error interno de validación: {str(e)}"],
                'warnings': []
            }

    def _validate_compensation_amount(self, monto_str):
        """
        Valida el monto de compensación.
        
        Args:
            monto_str: String del monto
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            errors = []
            warnings = []
            
            if not monto_str:
                errors.append("El monto de compensación es obligatorio")
                return {'valid': False, 'errors': errors}
            
            # Limpiar el monto
            monto_limpio = ''.join(c for c in monto_str if c.isdigit() or c in '.,')
            
            if not monto_limpio:
                errors.append("El monto debe contener números válidos")
                return {'valid': False, 'errors': errors}
            
            # Intentar conversión
            try:
                # Normalizar formato (reemplazar coma por punto para decimales)
                if ',' in monto_limpio and '.' in monto_limpio:
                    # Formato argentino: 1.000,50
                    pos_coma = monto_limpio.rfind(',')
                    pos_punto = monto_limpio.rfind('.')
                    if pos_coma > pos_punto:
                        monto_limpio = monto_limpio.replace('.', '').replace(',', '.')
                    else:
                        monto_limpio = monto_limpio.replace(',', '')
                elif ',' in monto_limpio:
                    # Solo coma - asumir decimal si hay 2 dígitos después
                    partes = monto_limpio.split(',')
                    if len(partes) == 2 and len(partes[1]) <= 2:
                        monto_limpio = monto_limpio.replace(',', '.')
                    else:
                        monto_limpio = monto_limpio.replace(',', '')
                
                monto_float = float(monto_limpio)
                
                # Validar rango
                if monto_float <= 0:
                    errors.append("El monto debe ser mayor a cero")
                elif monto_float > 999999999999:  # 999 mil millones
                    errors.append("El monto es demasiado grande (máximo: 999.999.999.999)")
                elif monto_float < 1:
                    warnings.append("El monto es menor a $1 - verifique que sea correcto")
                elif monto_float > 100000000:  # 100 millones
                    warnings.append("El monto es muy alto - verifique que sea correcto")
                
            except ValueError:
                errors.append("El monto debe ser un número válido")
            
            return {
                'valid': len(errors) == 0,
                'errors': errors,
                'warnings': warnings
            }
            
        except Exception as e:
            print(f"[ERROR] Error validando monto: {e}")
            return {'valid': False, 'errors': [f"Error validando monto: {str(e)}"]}

    def _validate_payment_period(self, plazo_str):
        """
        Valida el plazo de pago.
        
        Args:
            plazo_str: String del plazo
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            errors = []
            warnings = []
            
            if not plazo_str:
                errors.append("El plazo de pago es obligatorio")
                return {'valid': False, 'errors': errors}
            
            # Extraer solo dígitos
            plazo_limpio = ''.join(filter(str.isdigit, plazo_str))
            
            if not plazo_limpio:
                errors.append("El plazo debe contener números válidos")
                return {'valid': False, 'errors': errors}
            
            try:
                plazo_int = int(plazo_limpio)
                
                if plazo_int <= 0:
                    errors.append("El plazo debe ser mayor a cero días")
                elif plazo_int > 365:
                    errors.append("El plazo no puede ser mayor a 365 días")
                elif plazo_int > 180:
                    warnings.append("El plazo es mayor a 6 meses - verifique que sea correcto")
                elif plazo_int < 7:
                    warnings.append("El plazo es menor a una semana - verifique que sea correcto")
                
            except ValueError:
                errors.append("El plazo debe ser un número entero de días")
            
            return {
                'valid': len(errors) == 0,
                'errors': errors,
                'warnings': warnings
            }
            
        except Exception as e:
            print(f"[ERROR] Error validando plazo: {e}")
            return {'valid': False, 'errors': [f"Error validando plazo: {str(e)}"]}

    def _validate_banking_data(self, values):
        """
        Valida los datos bancarios opcionales.
        
        Args:
            values: Valores del formulario
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            warnings = []
            
            cbu = values.get('cbu_actor', '').strip()
            cuit = values.get('cuit_actor', '').strip()
            
            # Validar CBU si se proporciona
            if cbu:
                cbu_limpio = ''.join(filter(str.isdigit, cbu))
                if len(cbu_limpio) != 22:
                    warnings.append("El CBU debe tener exactamente 22 dígitos")
            
            # Validar CUIT si se proporciona
            if cuit:
                cuit_limpio = ''.join(filter(str.isdigit, cuit))
                if len(cuit_limpio) != 11:
                    warnings.append("El CUIT/CUIL debe tener exactamente 11 dígitos")
            
            return {'warnings': warnings}
            
        except Exception as e:
            print(f"[ERROR] Error validando datos bancarios: {e}")
            return {'warnings': [f"Error validando datos bancarios: {str(e)}"]}

    def _show_agreement_validation_errors(self, errors, dialog):
        """
        Muestra errores de validación al usuario.
        
        Args:
            errors: Lista de errores
            dialog: Ventana del diálogo
        """
        try:
            if not errors:
                return
            
            error_message = "Por favor corrija los siguientes errores:\n\n"
            for i, error in enumerate(errors, 1):
                error_message += f"{i}. {error}\n"
            
            messagebox.showerror(
                "Errores de Validación",
                error_message,
                parent=dialog
            )
            
        except Exception as e:
            print(f"[ERROR] Error mostrando errores de validación: {e}")

    def _show_agreement_validation_warnings(self, warnings, dialog):
        """
        Muestra advertencias de validación al usuario.
        
        Args:
            warnings: Lista de advertencias
            dialog: Ventana del diálogo
        """
        try:
            if not warnings:
                return
            
            warning_message = "Advertencias encontradas:\n\n"
            for i, warning in enumerate(warnings, 1):
                warning_message += f"{i}. {warning}\n"
            
            warning_message += "\n¿Desea continuar de todas formas?"
            
            result = messagebox.askyesno(
                "Advertencias de Validación",
                warning_message,
                parent=dialog
            )
            
            return result
            
        except Exception as e:
            print(f"[ERROR] Error mostrando advertencias: {e}")
            return True

    def _clean_and_format_agreement_data(self, values):
        """
        Limpia y formatea los datos finales del acuerdo.
        
        Args:
            values: Valores del formulario
            
        Returns:
            dict: Datos limpios y formateados
        """
        try:
            cleaned_data = {}
            
            # Limpiar monto
            monto = values.get('monto_compensacion_numeros', '').strip()
            cleaned_data['monto_compensacion_numeros'] = monto
            
            # Limpiar plazo
            plazo = values.get('plazo_pago_dias', '').strip()
            cleaned_data['plazo_pago_dias'] = ''.join(filter(str.isdigit, plazo))
            
            # Limpiar datos bancarios
            cleaned_data['banco_actor'] = values.get('banco_actor', '').strip()
            cleaned_data['cbu_actor'] = values.get('cbu_actor', '').strip()
            cleaned_data['alias_actor'] = values.get('alias_actor', '').strip()
            cleaned_data['cuit_actor'] = values.get('cuit_actor', '').strip()
            
            return cleaned_data
            
        except Exception as e:
            print(f"[ERROR] Error limpiando datos: {e}")
            return values

    def _show_agreement_dialog_error(self, error):
        """
        Muestra error crítico del diálogo.
        
        Args:
            error: Excepción capturada
        """
        try:
            messagebox.showerror(
                "Error del Diálogo",
                f"Error crítico en el diálogo de detalles del acuerdo:\n\n{str(error)}\n\nContacte al soporte técnico si el problema persiste.",
                parent=self.app_controller.root
            )
        except Exception as e:
            print(f"[ERROR] Error mostrando error del diálogo: {e}")

    def _create_enhanced_tooltip(self, widget, text):
        """
        Crea un tooltip mejorado para un widget.
        
        Args:
            widget: Widget al que agregar el tooltip
            text: Texto del tooltip
        """
        try:
            def on_enter(event):
                try:
                    tooltip = tk.Toplevel()
                    tooltip.wm_overrideredirect(True)
                    
                    # Posicionar tooltip
                    x = event.x_root + 10
                    y = event.y_root + 10
                    
                    # Ajustar si se sale de la pantalla
                    screen_width = tooltip.winfo_screenwidth()
                    screen_height = tooltip.winfo_screenheight()
                    
                    if x + 300 > screen_width:
                        x = event.x_root - 310
                    if y + 100 > screen_height:
                        y = event.y_root - 110
                    
                    tooltip.wm_geometry(f"+{x}+{y}")
                    
                    # Crear label con texto
                    label = tk.Label(
                        tooltip, 
                        text=text, 
                        background="lightyellow",
                        foreground="black",
                        relief="solid", 
                        borderwidth=1, 
                        font=("Arial", 9),
                        wraplength=300,
                        justify=tk.LEFT,
                        padx=5,
                        pady=3
                    )
                    label.pack()
                    
                    widget.tooltip = tooltip
                    
                except Exception as e:
                    print(f"[ERROR] Error creando tooltip: {e}")
            
            def on_leave(event):
                try:
                    if hasattr(widget, 'tooltip'):
                        widget.tooltip.destroy()
                        del widget.tooltip
                except Exception as e:
                    print(f"[ERROR] Error destruyendo tooltip: {e}")
            
            widget.bind("<Enter>", on_enter)
            widget.bind("<Leave>", on_leave)
            
        except Exception as e:
            print(f"[ERROR] Error configurando tooltip: {e}")

    def _create_tooltip(self, widget, text):
        """
        Crea un tooltip simple para un widget.
        
        Args:
            widget: Widget al que agregar el tooltip
            text: Texto del tooltip
        """
        def on_enter(event):
            tooltip = tk.Toplevel()
            tooltip.wm_overrideredirect(True)
            tooltip.wm_geometry(f"+{event.x_root+10}+{event.y_root+10}")
            
            label = tk.Label(tooltip, text=text, background="lightyellow", 
                           relief="solid", borderwidth=1, font=("Arial", 8))
            label.pack()
            
            widget.tooltip = tooltip
        
        def on_leave(event):
            if hasattr(widget, 'tooltip'):
                widget.tooltip.destroy()
                del widget.tooltip
        
        widget.bind("<Enter>", on_enter)
        widget.bind("<Leave>", on_leave)

    def _ask_representatives_assignment_dialog(self, lista_partes, abogados):
        """
        Abre un diálogo para que el usuario asigne manualmente los abogados a las partes.
        Implementa validación robusta y manejo de casos edge.
        Devuelve un diccionario {parte_id: [lista_de_abogados]} o None si se cancela.
        """
        try:
            # Validación de entrada
            if not self._validate_dialog_inputs(lista_partes, abogados):
                return None
            
            # Validar que hay interfaz disponible
            if not (hasattr(self, 'app_controller') and 
                   self.app_controller and 
                   hasattr(self.app_controller, 'root') and
                   self.app_controller.root):
                print("[ERROR] No hay interfaz disponible para mostrar diálogo")
                return None
            
            dialog = tk.Toplevel(self.app_controller.root)
            dialog.title("Asignar Representantes")
            dialog.transient(self.app_controller.root)
            dialog.grab_set()
            dialog.geometry("600x400")
            dialog.resizable(True, True)

            frame = ttk.Frame(dialog, padding="15")
            frame.pack(expand=True, fill=tk.BOTH)

            # Mensaje informativo
            info_text = f"Asigne los abogados a cada parte:\n{len(lista_partes)} partes, {len(abogados)} abogados disponibles"
            ttk.Label(frame, text=info_text, font=('TkDefaultFont', 10, 'bold')).pack(pady=(0, 10))

            # Crear un frame con scroll para las asignaciones
            canvas = tk.Canvas(frame)
            scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)

            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )

            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)

            # Variables para almacenar las asignaciones
            asignaciones_vars = {}

            # Crear controles para cada parte válida
            partes_procesadas = 0
            for i, parte in enumerate(lista_partes):
                if not isinstance(parte, dict) or not parte.get('nombre_completo') or not parte.get('rol_id'):
                    print(f"[WARNING] Parte inválida omitida en diálogo: {parte}")
                    continue
                
                parte_frame = ttk.LabelFrame(
                    scrollable_frame, 
                    text=f"{parte.get('nombre_completo', 'N/A')} ({parte.get('rol_principal', 'N/A')})", 
                    padding="10"
                )
                parte_frame.pack(fill=tk.X, pady=5, padx=5)

                asignaciones_vars[parte['rol_id']] = []

                # Crear checkboxes para cada abogado válido
                abogados_agregados = 0
                for abogado in abogados:
                    if not isinstance(abogado, dict) or not abogado.get('nombre_completo'):
                        continue
                    
                    var = tk.BooleanVar()
                    checkbox = ttk.Checkbutton(
                        parte_frame, 
                        text=f"{abogado.get('nombre_completo', 'N/A')}",
                        variable=var
                    )
                    checkbox.pack(anchor=tk.W, pady=2)
                    asignaciones_vars[parte['rol_id']].append((abogado, var))
                    abogados_agregados += 1
                
                if abogados_agregados == 0:
                    ttk.Label(parte_frame, text="No hay abogados disponibles", foreground="red").pack(anchor=tk.W)
                
                partes_procesadas += 1

            if partes_procesadas == 0:
                ttk.Label(scrollable_frame, text="No hay partes válidas para asignar", foreground="red").pack(pady=20)

            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

            # Botones
            button_frame = ttk.Frame(frame)
            button_frame.pack(fill=tk.X, pady=(10, 0))

            result = {}
            dialog_completed = False

            def save_assignments():
                nonlocal dialog_completed
                try:
                    for parte_id, abogado_vars in asignaciones_vars.items():
                        abogados_asignados = []
                        for abogado, var in abogado_vars:
                            if var.get():
                                abogados_asignados.append(abogado)
                        result[parte_id] = abogados_asignados
                    
                    dialog_completed = True
                    dialog.destroy()
                    print(f"[DEBUG] Asignaciones guardadas para {len(result)} partes")
                    
                except Exception as e:
                    print(f"[ERROR] Error guardando asignaciones: {e}")
                    messagebox.showerror("Error", f"Error guardando asignaciones: {e}", parent=dialog)

            def cancel():
                nonlocal dialog_completed
                try:
                    result.clear()
                    dialog_completed = True
                    dialog.destroy()
                    print("[DEBUG] Diálogo de asignación cancelado por el usuario")
                    
                except Exception as e:
                    print(f"[ERROR] Error cancelando diálogo: {e}")

            def auto_assign():
                """Asignación automática inteligente en el diálogo"""
                try:
                    # Limpiar asignaciones previas
                    for parte_id, abogado_vars in asignaciones_vars.items():
                        for abogado, var in abogado_vars:
                            var.set(False)
                    
                    # Aplicar lógica de asignación automática
                    self._apply_auto_assignment_to_dialog(lista_partes, abogados, asignaciones_vars)
                    
                except Exception as e:
                    print(f"[ERROR] Error en asignación automática del diálogo: {e}")
                    messagebox.showerror("Error", f"Error en asignación automática: {e}", parent=dialog)

            # Crear botones con validación
            if partes_procesadas > 0 and abogados:
                ttk.Button(button_frame, text="Asignación Automática", command=auto_assign).pack(side=tk.LEFT, padx=5)
            
            ttk.Button(button_frame, text="Aceptar", command=save_assignments).pack(side=tk.RIGHT, padx=5)
            ttk.Button(button_frame, text="Cancelar", command=cancel).pack(side=tk.RIGHT)

            dialog.protocol("WM_DELETE_WINDOW", cancel)
            
            # Esperar a que el diálogo se complete
            self.app_controller.root.wait_window(dialog)
            
            return result if dialog_completed and result else None
            
        except Exception as e:
            print(f"[ERROR] Error crítico en diálogo de asignación: {e}")
            return None

    def _validate_dialog_inputs(self, lista_partes, abogados):
        """
        Valida las entradas para el diálogo de asignación.
        
        Args:
            lista_partes (list): Lista de partes
            abogados (list): Lista de abogados
            
        Returns:
            bool: True si las entradas son válidas
        """
        try:
            if not isinstance(lista_partes, list) or not lista_partes:
                print("[ERROR] Lista de partes vacía o inválida para diálogo")
                return False
            
            if not isinstance(abogados, list) or not abogados:
                print("[ERROR] Lista de abogados vacía o inválida para diálogo")
                return False
            
            # Verificar que hay al menos una parte válida
            partes_validas = [p for p in lista_partes if isinstance(p, dict) and p.get('nombre_completo') and p.get('rol_id')]
            if not partes_validas:
                print("[ERROR] No hay partes válidas para mostrar en diálogo")
                return False
            
            # Verificar que hay al menos un abogado válido
            abogados_validos = [a for a in abogados if isinstance(a, dict) and a.get('nombre_completo')]
            if not abogados_validos:
                print("[ERROR] No hay abogados válidos para mostrar en diálogo")
                return False
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Error validando entradas del diálogo: {e}")
            return False

    def _apply_auto_assignment_to_dialog(self, lista_partes, abogados, asignaciones_vars):
        """
        Aplica asignación automática a los controles del diálogo.
        
        Args:
            lista_partes (list): Lista de partes
            abogados (list): Lista de abogados
            asignaciones_vars (dict): Variables del diálogo
        """
        try:
            # Separar por tipo de parte
            actores = [p for p in lista_partes if isinstance(p, dict) and 'actor' in p.get('rol_principal', '').lower()]
            demandados = [p for p in lista_partes if isinstance(p, dict) and 'demandado' in p.get('rol_principal', '').lower()]
            
            asignaciones_realizadas = 0
            
            # Asignar primer abogado al primer actor
            if actores and abogados and len(abogados) > 0:
                actor_id = actores[0].get('rol_id')
                if actor_id in asignaciones_vars and len(asignaciones_vars[actor_id]) > 0:
                    asignaciones_vars[actor_id][0][1].set(True)
                    asignaciones_realizadas += 1
                    print(f"[DEBUG] Auto-asignado primer abogado al actor: {actores[0].get('nombre_completo', 'N/A')}")
            
            # Asignar resto de abogados a demandados
            for i, demandado in enumerate(demandados):
                if len(abogados) <= 1:
                    break
                    
                abogado_index = (i + 1) % len(abogados)
                demandado_id = demandado.get('rol_id')
                
                if (demandado_id in asignaciones_vars and 
                    abogado_index < len(asignaciones_vars[demandado_id])):
                    asignaciones_vars[demandado_id][abogado_index][1].set(True)
                    asignaciones_realizadas += 1
                    print(f"[DEBUG] Auto-asignado abogado {abogado_index + 1} al demandado: {demandado.get('nombre_completo', 'N/A')}")
            
            print(f"[DEBUG] Asignación automática en diálogo completada: {asignaciones_realizadas} asignaciones")
            
        except Exception as e:
            print(f"[ERROR] Error aplicando asignación automática en diálogo: {e}")


    def _ensamblar_representantes(self, lista_base, todos_los_roles):
        """
        Enriquecer cada parte en lista_base con una lista de representantes.
        Busca en todos_los_roles aquellos con representa_a_id == rol_id de la parte.
        Si no hay relaciones establecidas, usa lógica inteligente para asignar abogados.
        
        Args:
            lista_base (list): Lista de partes (actores o demandados)
            todos_los_roles (list): Lista completa de roles del caso
            
        Returns:
            list: Lista de partes enriquecida con representantes
        """
        try:
            # Validación robusta de entrada
            validated_lista_base = self._validate_party_data_structure(lista_base)
            validated_todos_los_roles = self._validate_roles_data_structure(todos_los_roles)
            
            # Si no hay partes válidas, retornar lista vacía
            if not validated_lista_base:
                print("[DEBUG] No hay partes válidas para procesar")
                return []
            
            # Inicializar representantes vacíos para todas las partes válidas
            for parte in validated_lista_base:
                parte['representantes'] = []
            
            # Si no hay roles válidos, retornar partes sin representantes
            if not validated_todos_los_roles:
                print("[DEBUG] No hay roles válidos, todas las partes quedarán sin representantes")
                return validated_lista_base
            
            # Obtener todos los abogados de manera segura
            abogados = self._extract_valid_lawyers(validated_todos_los_roles)
            
            print(f"[DEBUG] Procesando {len(validated_lista_base)} partes con {len(abogados)} abogados disponibles")
            
            if not abogados:
                print("[DEBUG] No hay abogados disponibles, todas las partes quedarán sin representantes")
                return validated_lista_base
            
            # Verificar si hay relaciones establecidas específicamente para abogados
            relaciones_establecidas = self._check_established_relationships(abogados)
            
            if relaciones_establecidas:
                # Usar relaciones de la base de datos
                print(f"[DEBUG] Usando relaciones establecidas en BD para {len(abogados)} abogados")
                self._aplicar_relaciones_establecidas(validated_lista_base, validated_todos_los_roles)
            else:
                # Fallback inteligente: asignar abogados basado en lógica común
                print(f"[DEBUG] No hay relaciones establecidas, usando asignación inteligente para {len(abogados)} abogados")
                self._asignar_representantes_inteligente(validated_lista_base, abogados)
            
            # Validar y limpiar resultado final
            return self._validate_and_clean_final_result(validated_lista_base)
            
        except Exception as e:
            print(f"[ERROR] Error crítico en _ensamblar_representantes: {e}")
            # Fallback de emergencia: retornar lista segura con representantes vacíos
            return self._create_safe_fallback_result(lista_base)

    def _validate_party_data_structure(self, lista_base):
        """
        Valida y limpia la estructura de datos de las partes.
        
        Args:
            lista_base: Lista de partes a validar
            
        Returns:
            list: Lista de partes válidas
        """
        try:
            if not isinstance(lista_base, list):
                print(f"[ERROR] lista_base debe ser una lista, recibido: {type(lista_base)}")
                return []
            
            if not lista_base:
                print("[DEBUG] Lista de partes está vacía")
                return []
            
            partes_validas = []
            for i, parte in enumerate(lista_base):
                if not isinstance(parte, dict):
                    print(f"[WARNING] Parte en índice {i} no es un diccionario válido: {type(parte)}")
                    continue
                
                # Validar campos requeridos
                if not parte.get('nombre_completo'):
                    print(f"[WARNING] Parte en índice {i} no tiene nombre_completo válido")
                    continue
                
                if not parte.get('rol_id'):
                    print(f"[WARNING] Parte '{parte.get('nombre_completo', 'N/A')}' no tiene rol_id válido")
                    continue
                
                # Asegurar que tiene rol_principal
                if not parte.get('rol_principal'):
                    parte['rol_principal'] = 'Parte'
                    print(f"[DEBUG] Asignado rol_principal por defecto a {parte['nombre_completo']}")
                
                partes_validas.append(parte)
            
            print(f"[DEBUG] Validación de partes: {len(partes_validas)} válidas de {len(lista_base)} totales")
            return partes_validas
            
        except Exception as e:
            print(f"[ERROR] Error validando estructura de partes: {e}")
            return []

    def _validate_roles_data_structure(self, todos_los_roles):
        """
        Valida y limpia la estructura de datos de los roles.
        
        Args:
            todos_los_roles: Lista de roles a validar
            
        Returns:
            list: Lista de roles válidos
        """
        try:
            if not isinstance(todos_los_roles, list):
                print(f"[ERROR] todos_los_roles debe ser una lista, recibido: {type(todos_los_roles)}")
                return []
            
            if not todos_los_roles:
                print("[DEBUG] Lista de roles está vacía")
                return []
            
            roles_validos = []
            for i, rol in enumerate(todos_los_roles):
                if not isinstance(rol, dict):
                    print(f"[WARNING] Rol en índice {i} no es un diccionario válido: {type(rol)}")
                    continue
                
                # Validar campos mínimos requeridos
                if not rol.get('nombre_completo'):
                    print(f"[WARNING] Rol en índice {i} no tiene nombre_completo válido")
                    continue
                
                if not rol.get('rol_id'):
                    print(f"[WARNING] Rol '{rol.get('nombre_completo', 'N/A')}' no tiene rol_id válido")
                    continue
                
                # Asegurar que tiene rol_principal
                if not rol.get('rol_principal'):
                    rol['rol_principal'] = 'Rol'
                    print(f"[DEBUG] Asignado rol_principal por defecto a {rol['nombre_completo']}")
                
                roles_validos.append(rol)
            
            print(f"[DEBUG] Validación de roles: {len(roles_validos)} válidos de {len(todos_los_roles)} totales")
            return roles_validos
            
        except Exception as e:
            print(f"[ERROR] Error validando estructura de roles: {e}")
            return []

    def _extract_valid_lawyers(self, todos_los_roles):
        """
        Extrae y valida los abogados de la lista de roles.
        
        Args:
            todos_los_roles (list): Lista de roles validados
            
        Returns:
            list: Lista de abogados válidos
        """
        try:
            abogados = []
            for rol in todos_los_roles:
                rol_principal = rol.get('rol_principal', '').lower()
                if 'abogado' in rol_principal or 'letrado' in rol_principal or 'representante' in rol_principal:
                    # Validar que el abogado tiene datos mínimos
                    if rol.get('nombre_completo') and rol.get('rol_id'):
                        abogados.append(rol)
                        print(f"[DEBUG] Abogado válido encontrado: {rol['nombre_completo']}")
                    else:
                        print(f"[WARNING] Abogado con datos incompletos descartado: {rol}")
            
            return abogados
            
        except Exception as e:
            print(f"[ERROR] Error extrayendo abogados válidos: {e}")
            return []

    def _check_established_relationships(self, abogados):
        """
        Verifica si existen relaciones establecidas entre abogados y partes.
        
        Args:
            abogados (list): Lista de abogados
            
        Returns:
            bool: True si hay relaciones establecidas
        """
        try:
            if not abogados:
                return False
            
            for abogado in abogados:
                if isinstance(abogado, dict) and abogado.get('representa_a_id') is not None:
                    return True
            
            return False
            
        except Exception as e:
            print(f"[ERROR] Error verificando relaciones establecidas: {e}")
            return False

    def _validate_and_clean_final_result(self, lista_partes):
        """
        Valida y limpia el resultado final de la asignación de representantes.
        
        Args:
            lista_partes (list): Lista de partes con representantes asignados
            
        Returns:
            list: Lista de partes con representantes validados
        """
        try:
            for parte in lista_partes:
                if not isinstance(parte, dict):
                    continue
                
                representantes = parte.get('representantes', [])
                nombre_parte = parte.get('nombre_completo', 'N/A')
                
                # Validar cada representante
                representantes_validos = []
                for rep in representantes:
                    if isinstance(rep, dict) and rep.get('nombre_completo') and rep.get('rol_id'):
                        representantes_validos.append(rep)
                        print(f"[DEBUG]   - {rep.get('nombre_completo')} (ID: {rep.get('rol_id', 'N/A')})")
                    else:
                        print(f"[WARNING] Representante inválido descartado para {nombre_parte}: {rep}")
                
                parte['representantes'] = representantes_validos
                print(f"[DEBUG] Parte {nombre_parte} tiene {len(representantes_validos)} representantes válidos")
            
            return lista_partes
            
        except Exception as e:
            print(f"[ERROR] Error validando resultado final: {e}")
            return self._create_safe_fallback_result(lista_partes)

    def _create_safe_fallback_result(self, lista_base):
        """
        Crea un resultado seguro de fallback cuando hay errores críticos.
        
        Args:
            lista_base: Lista original de partes
            
        Returns:
            list: Lista segura con representantes vacíos
        """
        try:
            if not isinstance(lista_base, list):
                return []
            
            resultado_seguro = []
            for parte in lista_base:
                if isinstance(parte, dict):
                    parte_segura = parte.copy()
                    parte_segura['representantes'] = []
                    resultado_seguro.append(parte_segura)
            
            print(f"[DEBUG] Creado resultado de fallback seguro con {len(resultado_seguro)} partes")
            return resultado_seguro
            
        except Exception as e:
            print(f"[ERROR] Error creando resultado de fallback: {e}")
            return []

    def _aplicar_relaciones_establecidas(self, lista_base, todos_los_roles):
        """
        Aplica las relaciones establecidas en la base de datos entre partes y representantes.
        Implementa validación robusta y manejo de casos edge.
        
        Args:
            lista_base (list): Lista de partes validadas
            todos_los_roles (list): Lista completa de roles validados
        """
        try:
            # Validación de entrada
            if not isinstance(lista_base, list) or not lista_base:
                print("[DEBUG] Lista de partes vacía o inválida para aplicar relaciones establecidas")
                return
            
            if not isinstance(todos_los_roles, list) or not todos_los_roles:
                print("[DEBUG] Lista de roles vacía o inválida para aplicar relaciones establecidas")
                return
            
            # Crear índice de roles por representa_a_id para búsqueda eficiente
            roles_por_representado = {}
            for rol in todos_los_roles:
                if not isinstance(rol, dict):
                    continue
                
                representa_a_id = rol.get('representa_a_id')
                if representa_a_id is not None:
                    rol_principal = rol.get('rol_principal', '').lower()
                    if 'abogado' in rol_principal or 'letrado' in rol_principal or 'representante' in rol_principal:
                        if representa_a_id not in roles_por_representado:
                            roles_por_representado[representa_a_id] = []
                        roles_por_representado[representa_a_id].append(rol)
            
            print(f"[DEBUG] Índice de representación creado para {len(roles_por_representado)} relaciones")
            
            # Aplicar relaciones a cada parte
            partes_con_representantes = 0
            total_representantes_asignados = 0
            
            for parte in lista_base:
                if not isinstance(parte, dict):
                    continue
                
                parte_id = parte.get('rol_id')
                nombre_parte = parte.get('nombre_completo', 'N/A')
                
                if parte_id is None:
                    print(f"[WARNING] Parte sin rol_id: {nombre_parte}")
                    parte['representantes'] = []
                    continue
                
                # Buscar representantes para esta parte
                representantes = roles_por_representado.get(parte_id, [])
                
                # Validar representantes encontrados
                representantes_validos = []
                for rep in representantes:
                    if isinstance(rep, dict) and rep.get('nombre_completo') and rep.get('rol_id'):
                        representantes_validos.append(rep)
                    else:
                        print(f"[WARNING] Representante inválido descartado para {nombre_parte}: {rep}")
                
                parte['representantes'] = representantes_validos
                
                if representantes_validos:
                    partes_con_representantes += 1
                    total_representantes_asignados += len(representantes_validos)
                    print(f"[DEBUG] Asignados {len(representantes_validos)} representantes a {nombre_parte}")
                else:
                    print(f"[DEBUG] No se encontraron representantes para {nombre_parte}")
            
            print(f"[DEBUG] Relaciones aplicadas: {partes_con_representantes}/{len(lista_base)} partes con representantes, {total_representantes_asignados} representantes totales")
                
        except Exception as e:
            print(f"[ERROR] Error crítico en _aplicar_relaciones_establecidas: {e}")
            # Fallback: asegurar que todas las partes tengan lista de representantes
            for parte in lista_base:
                if isinstance(parte, dict) and 'representantes' not in parte:
                    parte['representantes'] = []
    
    def _asignar_representantes_inteligente(self, lista_partes, abogados):
        """
        Asigna representantes de manera inteligente cuando no hay relaciones en BD.
        Implementa múltiples estrategias de asignación con fallbacks robustos.
        
        Args:
            lista_partes (list): Lista de partes que necesitan representantes
            abogados (list): Lista de abogados disponibles
        """
        try:
            # Validación robusta de entrada
            if not self._validate_intelligent_assignment_inputs(lista_partes, abogados):
                return
            
            print(f"[DEBUG] Iniciando asignación inteligente: {len(lista_partes)} partes, {len(abogados)} abogados")
            
            # Inicializar representantes vacíos para todas las partes
            for parte in lista_partes:
                if isinstance(parte, dict):
                    parte['representantes'] = []
            
            # Estrategia 1: Intentar diálogo de asignación manual (solo si hay interfaz disponible)
            if self._try_manual_assignment_dialog(lista_partes, abogados):
                return
            
            # Estrategia 2: Asignación automática inteligente
            self._asignacion_automatica_inteligente(lista_partes, abogados)
            
        except Exception as e:
            print(f"[ERROR] Error crítico en asignación inteligente: {e}")
            # Fallback final: asegurar representantes vacíos
            self._ensure_empty_representatives(lista_partes)

    def _validate_intelligent_assignment_inputs(self, lista_partes, abogados):
        """
        Valida las entradas para la asignación inteligente.
        
        Args:
            lista_partes (list): Lista de partes
            abogados (list): Lista de abogados
            
        Returns:
            bool: True si las entradas son válidas
        """
        try:
            if not isinstance(abogados, list) or not abogados:
                print("[DEBUG] No hay abogados disponibles para asignación inteligente")
                self._ensure_empty_representatives(lista_partes)
                return False
            
            if not isinstance(lista_partes, list) or not lista_partes:
                print("[DEBUG] No hay partes para asignar representantes")
                return False
            
            # Validar que hay al menos una parte válida
            partes_validas = [p for p in lista_partes if isinstance(p, dict) and p.get('nombre_completo')]
            if not partes_validas:
                print("[DEBUG] No hay partes válidas para asignar representantes")
                return False
            
            # Validar que hay al menos un abogado válido
            abogados_validos = [a for a in abogados if isinstance(a, dict) and a.get('nombre_completo')]
            if not abogados_validos:
                print("[DEBUG] No hay abogados válidos para asignación")
                self._ensure_empty_representatives(lista_partes)
                return False
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Error validando entradas para asignación inteligente: {e}")
            return False

    def _try_manual_assignment_dialog(self, lista_partes, abogados):
        """
        Intenta mostrar el diálogo de asignación manual.
        
        Args:
            lista_partes (list): Lista de partes
            abogados (list): Lista de abogados
            
        Returns:
            bool: True si se completó la asignación manual exitosamente
        """
        try:
            # Verificar si hay interfaz disponible
            if not (hasattr(self, 'app_controller') and 
                   self.app_controller and 
                   hasattr(self.app_controller, 'root') and
                   self.app_controller.root):
                print("[DEBUG] No hay interfaz disponible, usando asignación automática")
                return False
            
            # Intentar mostrar diálogo
            asignaciones = self._ask_representatives_assignment_dialog(lista_partes, abogados)
            
            if asignaciones and isinstance(asignaciones, dict):
                print("[DEBUG] Usuario proporcionó asignaciones manuales")
                self._aplicar_asignaciones_manuales(lista_partes, asignaciones)
                return True
            else:
                print("[DEBUG] Usuario canceló o no proporcionó asignaciones, usando fallback automático")
                return False
                
        except Exception as e:
            print(f"[DEBUG] Error en diálogo manual, usando fallback automático: {e}")
            return False

    def _ensure_empty_representatives(self, lista_partes):
        """
        Asegura que todas las partes tengan una lista de representantes vacía.
        
        Args:
            lista_partes (list): Lista de partes
        """
        try:
            if isinstance(lista_partes, list):
                for parte in lista_partes:
                    if isinstance(parte, dict):
                        parte['representantes'] = []
        except Exception as e:
            print(f"[ERROR] Error asegurando representantes vacíos: {e}")

    def _aplicar_asignaciones_manuales(self, lista_partes, asignaciones):
        """
        Aplica las asignaciones manuales proporcionadas por el usuario.
        Implementa validación robusta de las asignaciones.
        
        Args:
            lista_partes (list): Lista de partes
            asignaciones (dict): Diccionario {parte_id: [lista_de_abogados]}
        """
        try:
            # Validar entrada
            if not isinstance(lista_partes, list) or not lista_partes:
                print("[DEBUG] Lista de partes vacía o inválida para asignaciones manuales")
                return
            
            if not isinstance(asignaciones, dict) or not asignaciones:
                print("[DEBUG] Asignaciones vacías o inválidas, aplicando representantes vacíos")
                self._ensure_empty_representatives(lista_partes)
                return
            
            asignaciones_aplicadas = 0
            representantes_totales = 0
            
            for parte in lista_partes:
                if not isinstance(parte, dict):
                    continue
                
                parte_id = parte.get('rol_id')
                nombre_parte = parte.get('nombre_completo', 'N/A')
                
                if parte_id is None:
                    print(f"[WARNING] Parte sin rol_id para asignación manual: {nombre_parte}")
                    parte['representantes'] = []
                    continue
                
                if parte_id in asignaciones:
                    representantes = asignaciones[parte_id]
                    
                    # Validar que representantes es una lista
                    if isinstance(representantes, list):
                        # Validar cada representante
                        representantes_validos = []
                        for rep in representantes:
                            if isinstance(rep, dict) and rep.get('nombre_completo') and rep.get('rol_id'):
                                representantes_validos.append(rep)
                            else:
                                print(f"[WARNING] Representante inválido descartado para {nombre_parte}: {rep}")
                        
                        parte['representantes'] = representantes_validos
                        
                        if representantes_validos:
                            asignaciones_aplicadas += 1
                            representantes_totales += len(representantes_validos)
                            print(f"[DEBUG] Asignados {len(representantes_validos)} representantes a {nombre_parte}")
                        else:
                            print(f"[DEBUG] No se asignaron representantes válidos a {nombre_parte}")
                    else:
                        print(f"[WARNING] Asignación inválida para {nombre_parte}: {type(representantes)}")
                        parte['representantes'] = []
                else:
                    parte['representantes'] = []
                    print(f"[DEBUG] No hay asignación para {nombre_parte}")
            
            print(f"[DEBUG] Asignaciones manuales aplicadas: {asignaciones_aplicadas}/{len(lista_partes)} partes, {representantes_totales} representantes totales")
                    
        except Exception as e:
            print(f"[ERROR] Error crítico aplicando asignaciones manuales: {e}")
            # Fallback: asegurar representantes vacíos
            self._ensure_empty_representatives(lista_partes)

    def _asignacion_automatica_inteligente(self, lista_partes, abogados):
        """
        Implementa asignación automática inteligente de representantes.
        Incluye múltiples estrategias y validación robusta.
        
        Args:
            lista_partes (list): Lista de partes
            abogados (list): Lista de abogados disponibles
        """
        try:
            # Validar entrada
            if not isinstance(lista_partes, list) or not lista_partes:
                print("[DEBUG] No hay partes válidas para asignación automática")
                return
            
            if not isinstance(abogados, list) or not abogados:
                print("[DEBUG] No hay abogados válidos para asignación automática")
                self._ensure_empty_representatives(lista_partes)
                return
            
            # Separar por tipo de parte con validación
            actores = self._filter_parties_by_role(lista_partes, 'actor')
            demandados = self._filter_parties_by_role(lista_partes, 'demandado')
            otras_partes = self._filter_other_parties(lista_partes, actores, demandados)
            
            print(f"[DEBUG] Distribución: {len(actores)} actores, {len(demandados)} demandados, {len(otras_partes)} otras")
            
            # Inicializar representantes vacíos para todas las partes
            for parte in lista_partes:
                if isinstance(parte, dict):
                    parte['representantes'] = []
            
            # Aplicar estrategia de distribución basada en la cantidad de abogados
            if len(abogados) == 1:
                self._asignar_abogado_unico(actores, demandados, lista_partes, abogados[0])
            elif len(abogados) == 2 and actores and demandados:
                self._asignar_dos_abogados(actores, demandados, abogados)
            else:
                self._distribucion_equitativa(lista_partes, abogados)
            
            # Validar resultado final
            self._validate_automatic_assignment_result(lista_partes)
                    
        except Exception as e:
            print(f"[ERROR] Error crítico en asignación automática: {e}")
            # Fallback: asegurar representantes vacíos
            self._ensure_empty_representatives(lista_partes)

    def _filter_parties_by_role(self, lista_partes, role_type):
        """
        Filtra partes por tipo de rol con validación.
        
        Args:
            lista_partes (list): Lista de partes
            role_type (str): Tipo de rol a filtrar ('actor' o 'demandado')
            
        Returns:
            list: Lista de partes filtradas
        """
        try:
            filtered_parties = []
            for parte in lista_partes:
                if (isinstance(parte, dict) and 
                    isinstance(parte.get('rol_principal'), str) and
                    role_type.lower() in parte.get('rol_principal', '').lower() and
                    parte.get('nombre_completo') and
                    parte.get('rol_id')):
                    filtered_parties.append(parte)
            
            return filtered_parties
            
        except Exception as e:
            print(f"[ERROR] Error filtrando partes por rol {role_type}: {e}")
            return []

    def _filter_other_parties(self, lista_partes, actores, demandados):
        """
        Filtra partes que no son actores ni demandados.
        
        Args:
            lista_partes (list): Lista de partes
            actores (list): Lista de actores
            demandados (list): Lista de demandados
            
        Returns:
            list: Lista de otras partes
        """
        try:
            actor_ids = {p.get('rol_id') for p in actores if isinstance(p, dict)}
            demandado_ids = {p.get('rol_id') for p in demandados if isinstance(p, dict)}
            
            otras_partes = []
            for parte in lista_partes:
                if (isinstance(parte, dict) and 
                    parte.get('rol_id') not in actor_ids and
                    parte.get('rol_id') not in demandado_ids and
                    parte.get('nombre_completo') and
                    parte.get('rol_id')):
                    otras_partes.append(parte)
            
            return otras_partes
            
        except Exception as e:
            print(f"[ERROR] Error filtrando otras partes: {e}")
            return []

    def _asignar_abogado_unico(self, actores, demandados, lista_partes, abogado):
        """
        Asigna un único abogado siguiendo prioridades lógicas.
        
        Args:
            actores (list): Lista de actores
            demandados (list): Lista de demandados
            lista_partes (list): Lista completa de partes
            abogado (dict): Abogado a asignar
        """
        try:
            # Prioridad 1: Primer actor
            if actores and isinstance(actores[0], dict):
                actores[0]['representantes'] = [abogado]
                print(f"[DEBUG] Único abogado asignado al actor: {actores[0].get('nombre_completo', 'N/A')}")
                return
            
            # Prioridad 2: Primer demandado
            if demandados and isinstance(demandados[0], dict):
                demandados[0]['representantes'] = [abogado]
                print(f"[DEBUG] Único abogado asignado al demandado: {demandados[0].get('nombre_completo', 'N/A')}")
                return
            
            # Prioridad 3: Primera parte válida
            for parte in lista_partes:
                if isinstance(parte, dict) and parte.get('nombre_completo'):
                    parte['representantes'] = [abogado]
                    print(f"[DEBUG] Único abogado asignado a la primera parte: {parte.get('nombre_completo', 'N/A')}")
                    return
            
            print("[WARNING] No se pudo asignar el único abogado a ninguna parte válida")
            
        except Exception as e:
            print(f"[ERROR] Error asignando abogado único: {e}")

    def _asignar_dos_abogados(self, actores, demandados, abogados):
        """
        Asigna dos abogados: uno para actores, uno para demandados.
        
        Args:
            actores (list): Lista de actores
            demandados (list): Lista de demandados
            abogados (list): Lista de dos abogados
        """
        try:
            if len(abogados) < 2:
                print("[ERROR] Se requieren al menos 2 abogados para esta estrategia")
                return
            
            # Asignar primer abogado al primer actor
            if actores and isinstance(actores[0], dict):
                actores[0]['representantes'] = [abogados[0]]
                print(f"[DEBUG] Primer abogado asignado al actor: {actores[0].get('nombre_completo', 'N/A')}")
            
            # Asignar segundo abogado al primer demandado
            if demandados and isinstance(demandados[0], dict):
                demandados[0]['representantes'] = [abogados[1]]
                print(f"[DEBUG] Segundo abogado asignado al demandado: {demandados[0].get('nombre_completo', 'N/A')}")
            
            print("[DEBUG] Asignación 1:1 completada - un abogado por bando")
            
        except Exception as e:
            print(f"[ERROR] Error asignando dos abogados: {e}")

    def _validate_automatic_assignment_result(self, lista_partes):
        """
        Valida el resultado de la asignación automática.
        
        Args:
            lista_partes (list): Lista de partes con asignaciones
        """
        try:
            partes_con_representantes = 0
            total_representantes = 0
            
            for parte in lista_partes:
                if isinstance(parte, dict):
                    representantes = parte.get('representantes', [])
                    if representantes:
                        partes_con_representantes += 1
                        total_representantes += len(representantes)
            
            print(f"[DEBUG] Resultado asignación automática: {partes_con_representantes}/{len(lista_partes)} partes con representantes, {total_representantes} representantes totales")
            
        except Exception as e:
            print(f"[ERROR] Error validando resultado de asignación automática: {e}")

    def _distribucion_equitativa(self, lista_partes, abogados):
        """
        Distribuye los abogados equitativamente entre las partes.
        Implementa validación robusta y manejo de casos edge.
        
        Args:
            lista_partes (list): Lista de partes
            abogados (list): Lista de abogados disponibles
        """
        try:
            # Validar entrada
            if not isinstance(lista_partes, list) or not lista_partes:
                print("[DEBUG] No hay partes para distribución equitativa")
                return
            
            if not isinstance(abogados, list) or not abogados:
                print("[DEBUG] No hay abogados para distribución equitativa")
                return
            
            # Filtrar partes válidas
            partes_validas = []
            for parte in lista_partes:
                if (isinstance(parte, dict) and 
                    parte.get('nombre_completo') and 
                    parte.get('rol_id')):
                    partes_validas.append(parte)
            
            if not partes_validas:
                print("[DEBUG] No hay partes válidas para distribución equitativa")
                return
            
            # Filtrar abogados válidos
            abogados_validos = []
            for abogado in abogados:
                if (isinstance(abogado, dict) and 
                    abogado.get('nombre_completo') and 
                    abogado.get('rol_id')):
                    abogados_validos.append(abogado)
            
            if not abogados_validos:
                print("[DEBUG] No hay abogados válidos para distribución equitativa")
                return
            
            print(f"[DEBUG] Distribución equitativa: {len(partes_validas)} partes válidas, {len(abogados_validos)} abogados válidos")
            
            # Distribuir abogados de manera circular
            asignaciones_exitosas = 0
            for i, parte in enumerate(partes_validas):
                try:
                    abogado_index = i % len(abogados_validos)
                    abogado_asignado = abogados_validos[abogado_index]
                    
                    parte['representantes'] = [abogado_asignado]
                    asignaciones_exitosas += 1
                    
                    print(f"[DEBUG] Asignado abogado {abogado_asignado.get('nombre_completo', 'N/A')} a {parte.get('nombre_completo', 'N/A')}")
                    
                except Exception as e:
                    print(f"[WARNING] Error asignando abogado a parte {i}: {e}")
                    parte['representantes'] = []
            
            print(f"[DEBUG] Distribución equitativa completada: {asignaciones_exitosas}/{len(partes_validas)} asignaciones exitosas")
                
        except Exception as e:
            print(f"[ERROR] Error crítico en distribución equitativa: {e}")
            # Fallback: asegurar representantes vacíos
            self._ensure_empty_representatives(lista_partes)               
        
    def generar_escrito_mediacion(self, caso_id):
        """
        Genera el acuerdo de mediación usando docxtpl con datos jerárquicos.
        
        Este método actúa como orquestador de UI, manejando la interacción con el usuario
        y delegando la lógica de generación al método _generar_documento_con_datos.
        
        Args:
            caso_id: ID del caso para generar el acuerdo
            
        Returns:
            bool: True si el documento se generó y guardó exitosamente
        """
        operation_start_time = datetime.datetime.now()
        print(f"[INFO] Iniciando generación de acuerdo de mediación (UI) para caso ID: {caso_id} a las {operation_start_time.strftime('%H:%M:%S')}")
        
        try:
            # 0. Validación inicial de parámetros (requiere UI para mostrar errores)
            if not self._validate_initial_parameters(caso_id):
                return False
            
            # 1. Validación de dependencias del sistema (requiere UI para mostrar errores)
            if not self._validate_system_dependencies():
                return False
            
            # 2. Recopilación y validación de datos del caso (requiere UI para mostrar errores)
            caso_data = self._collect_and_validate_case_data(caso_id)
            if not caso_data:
                return False
            
            # 3. Recopilación y validación de partes/roles (requiere UI para mostrar errores)
            parties_data = self._collect_and_validate_parties_data(caso_id)
            if not parties_data:
                return False

            # 3.5. Procesamiento de representantes para datos jerárquicos
            print("[DEBUG] Procesando representantes para estructura jerárquica...")
            processed_parties = self._process_representatives(parties_data, caso_id)
            if not processed_parties:
                return False

            # 4. Obtención de detalles del acuerdo del usuario (UI CRÍTICA)
            print("[DEBUG] Solicitando detalles del acuerdo al usuario...")
            # Identificar actor principal para pre-rellenar datos bancarios
            actor_principal = processed_parties['lista_actores'][0] if processed_parties['lista_actores'] else {}
            agreement_details = self._collect_agreement_details(actor_principal)
            if not agreement_details:
                print("[INFO] Usuario canceló la generación del acuerdo")
                return False  # Usuario canceló
            
            print(f"[DEBUG] Detalles del acuerdo obtenidos: {list(agreement_details.keys())}")
            
            # 5. Generación del documento usando lógica pura (NUEVA IMPLEMENTACIÓN)
            print("[DEBUG] Delegando generación de documento al método de lógica pura...")
            document_result = self._generar_documento_con_datos(caso_id, agreement_details, processed_parties)
            
            if not document_result['success']:
                # Mostrar error específico al usuario
                error_msg = document_result['error_message']
                error_type = document_result['error_type']
                
                print(f"[ERROR] Error en generación de documento: {error_msg}")
                
                # Mostrar diálogo de error apropiado según el tipo
                if error_type in ['missing_case', 'empty_caratula', 'no_parties', 'no_actors', 'no_defendants']:
                    ErrorMessageManager.show_error_dialog(self.app_controller.root, error_type)
                else:
                    ErrorMessageManager.show_error_dialog(
                        self.app_controller.root, 
                        'unexpected_error', 
                        {'error': error_msg}
                    )
                return False
            
            print("[DEBUG] Documento generado exitosamente por método de lógica pura")
            
            # 6. Guardado del documento (requiere UI para diálogo de guardado)
            if not self._save_document(document_result['document'], caso_data):
                print("[INFO] Usuario canceló el guardado o error en guardado")
                return False
            
            # Operación completada exitosamente
            operation_end_time = datetime.datetime.now()
            duration = operation_end_time - operation_start_time
            print(f"[SUCCESS] Acuerdo de mediación generado y guardado exitosamente en {duration.total_seconds():.2f} segundos")
            return True
            
        except Exception as e:
            # Manejo de errores críticos no capturados
            print(f"[ERROR] Error crítico en generación de acuerdo (UI): {e}")
            return self._handle_critical_error(e, caso_id, operation_start_time)

    def _validate_initial_parameters(self, caso_id):
        """
        Valida los parámetros iniciales para la generación del acuerdo.
        
        Args:
            caso_id: ID del caso a validar
            
        Returns:
            bool: True si los parámetros son válidos
        """
        try:
            if caso_id is None:
                ErrorMessageManager.log_error('validation_error', {'field': 'caso_id'})
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'validation_error', 
                    {'field': 'ID del caso'}
                )
                return False
            
            # Validar que caso_id sea un número válido
            try:
                int(caso_id)
            except (ValueError, TypeError):
                ErrorMessageManager.log_error('validation_error', {'field': 'caso_id'})
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'validation_error', 
                    {'field': 'ID del caso (debe ser un número válido)'}
                )
                return False
            
            # Validar que tenemos acceso a la interfaz
            if not (hasattr(self, 'app_controller') and 
                   self.app_controller and 
                   hasattr(self.app_controller, 'root')):
                print("[ERROR] No hay acceso a la interfaz de usuario")
                return False
            
            # Validar que tenemos acceso a la base de datos
            if not hasattr(self, 'db') or not self.db:
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'unexpected_error', 
                    {'error': 'No hay conexión a la base de datos'}
                )
                return False
            
            print(f"[DEBUG] Parámetros iniciales validados correctamente para caso ID: {caso_id}")
            return True
            
        except Exception as e:
            print(f"[ERROR] Error validando parámetros iniciales: {e}")
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root, 
                'unexpected_error', 
                {'error': f'validación de parámetros - {str(e)}'}
            )
            return False

    def _validate_system_dependencies(self):
        """
        Valida las dependencias del sistema necesarias para generar acuerdos.
        
        Returns:
            bool: True si todas las dependencias están disponibles
        """
        try:
            print("[DEBUG] Validando dependencias del sistema...")
            
            validation_result = self._validate_mediation_dependencies()
            if not validation_result['success']:
                self._show_validation_errors(validation_result)
                return False
            
            # Mostrar advertencias si las hay, pero continuar
            if validation_result['warnings']:
                self._show_validation_errors(validation_result)
            
            print("[DEBUG] Dependencias del sistema validadas correctamente")
            return True
            
        except Exception as e:
            print(f"[ERROR] Error validando dependencias del sistema: {e}")
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root, 
                'missing_dependencies', 
                {'error': str(e)}
            )
            return False

    def _collect_and_validate_case_data(self, caso_id):
        """
        Recopila y valida los datos del caso.
        
        Args:
            caso_id: ID del caso
            
        Returns:
            dict: Datos del caso validados o None si hay error
        """
        try:
            print(f"[DEBUG] Recopilando datos del caso ID: {caso_id}")
            
            caso = self.db.get_case_by_id(caso_id)
            if not caso:
                ErrorMessageManager.log_error('missing_case', {'case_id': caso_id})
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'missing_case', 
                    {'case_id': caso_id}
                )
                return None
            
            # Validar datos críticos del caso
            if not isinstance(caso, dict):
                print(f"[ERROR] Datos del caso no son un diccionario válido: {type(caso)}")
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'unexpected_error', 
                    {'error': 'estructura de datos del caso inválida'}
                )
                return None
            
            # Validar carátula
            caratula = caso.get('caratula', '').strip()
            if not caratula:
                ErrorMessageManager.log_error('empty_caratula', {'case_id': caso_id})
                ErrorMessageManager.show_error_dialog(self.app_controller.root, 'empty_caratula')
                return None
            
            # Validar otros campos críticos
            numero_expediente = caso.get('numero_expediente', '').strip()
            if not numero_expediente:
                print(f"[WARNING] Caso sin número de expediente: {caratula}")
            
            print(f"[DEBUG] Caso validado: {caratula} (Exp: {numero_expediente or 'Sin número'})")
            return caso
            
        except Exception as e:
            print(f"[ERROR] Error recopilando datos del caso: {e}")
            ErrorMessageManager.log_error('unexpected_error', {'error': str(e)}, technical_details=str(e))
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root, 
                'unexpected_error', 
                {'error': f'recopilación de datos del caso - {str(e)}'}
            )
            return None

    def _collect_and_validate_parties_data(self, caso_id):
        """
        Recopila y valida los datos de las partes del caso.
        
        Args:
            caso_id: ID del caso
            
        Returns:
            dict: Datos de partes validados o None si hay error
        """
        try:
            print(f"[DEBUG] Recopilando datos de partes para caso ID: {caso_id}")
            
            todos_los_roles = self.db.get_roles_by_case_id(caso_id)
            if not todos_los_roles:
                ErrorMessageManager.log_error('no_parties', {'case_id': caso_id})
                ErrorMessageManager.show_error_dialog(self.app_controller.root, 'no_parties')
                return None
            
            if not isinstance(todos_los_roles, list):
                print(f"[ERROR] Roles no son una lista válida: {type(todos_los_roles)}")
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'unexpected_error', 
                    {'error': 'estructura de datos de roles inválida'}
                )
                return None
            
            print(f"[DEBUG] Roles encontrados: {len(todos_los_roles)}")

            # Separar actores y demandados usando list comprehensions
            lista_actores = [r for r in todos_los_roles if isinstance(r, dict) and 'actor' in r.get('rol_principal', '').lower()]
            lista_demandados = [r for r in todos_los_roles if isinstance(r, dict) and 'demandado' in r.get('rol_principal', '').lower()]
            
            # Validar que existan actores y demandados
            if not lista_actores:
                ErrorMessageManager.log_error('no_actors', {'case_id': caso_id})
                ErrorMessageManager.show_error_dialog(self.app_controller.root, 'no_actors')
                return None
            
            if not lista_demandados:
                ErrorMessageManager.log_error('no_defendants', {'case_id': caso_id})
                ErrorMessageManager.show_error_dialog(self.app_controller.root, 'no_defendants')
                return None
            
            print(f"[DEBUG] Partes validadas - Actores: {len(lista_actores)}, Demandados: {len(lista_demandados)}")
            
            return {
                'todos_los_roles': todos_los_roles,
                'lista_actores': lista_actores,
                'lista_demandados': lista_demandados
            }
            
        except Exception as e:
            print(f"[ERROR] Error recopilando datos de partes: {e}")
            ErrorMessageManager.log_error('unexpected_error', {'error': str(e)}, technical_details=str(e))
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root, 
                'unexpected_error', 
                {'error': f'recopilación de datos de partes - {str(e)}'}
            )
            return None

    def _collect_agreement_details(self, actor_principal=None):
        """
        Recopila los detalles del acuerdo del usuario.

        Args:
            actor_principal: Datos del actor principal para pre-rellenar (opcional)

        Returns:
            dict: Detalles del acuerdo o None si el usuario canceló
        """
        try:
            print(f"[DEBUG] Abriendo diálogo de detalles con actor principal: {actor_principal.get('nombre_completo', 'N/A') if actor_principal else 'Ninguno'}")

            details = self._ask_agreement_details_dialog(actor_principal)
            if not details:
                print("[INFO] Usuario canceló el diálogo de detalles del acuerdo")
                return None

            # Validar datos críticos del acuerdo
            if not isinstance(details, dict):
                print(f"[ERROR] Detalles del acuerdo no son un diccionario válido: {type(details)}")
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root,
                    'unexpected_error',
                    {'error': 'estructura de detalles del acuerdo inválida'}
                )
                return None

            # Validar monto de compensación
            monto = details.get('monto_compensacion_numeros', '').strip()
            if not monto:
                ErrorMessageManager.log_error('validation_error', {'field': 'monto_compensacion'})
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root,
                    'validation_error',
                    {'field': 'monto de compensación'}
                )
                return None

            # Validar plazo de pago
            plazo = details.get('plazo_pago_dias', '').strip()
            if not plazo:
                print("[WARNING] No se especificó plazo de pago, usando valor por defecto")
                details['plazo_pago_dias'] = '0'

            print(f"[DEBUG] Detalles del acuerdo validados - Monto: {monto}, Plazo: {details.get('plazo_pago_dias', '0')} días")
            return details

        except Exception as e:
            print(f"[ERROR] Error recopilando detalles del acuerdo: {e}")
            ErrorMessageManager.log_error('unexpected_error', {'error': str(e)}, technical_details=str(e))
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root,
                'unexpected_error',
                {'error': f'recopilación de detalles del acuerdo - {str(e)}'}
            )
            return None

    def _process_representatives(self, parties_data, caso_id):
        """
        Procesa los representantes para actores y demandados.
        
        Args:
            parties_data: Datos de las partes
            caso_id: ID del caso
            
        Returns:
            dict: Partes con representantes procesados o None si hay error
        """
        try:
            print("[DEBUG] Procesando representantes para actores y demandados")
            
            if not isinstance(parties_data, dict):
                print(f"[ERROR] Datos de partes no son un diccionario válido: {type(parties_data)}")
                return None
            
            lista_actores = parties_data.get('lista_actores', [])
            lista_demandados = parties_data.get('lista_demandados', [])
            todos_los_roles = parties_data.get('todos_los_roles', [])
            
            # Procesar representantes con manejo de errores robusto
            try:
                processed_actores = self._ensamblar_representantes(lista_actores, todos_los_roles)
                processed_demandados = self._ensamblar_representantes(lista_demandados, todos_los_roles)
            except Exception as e:
                print(f"[ERROR] Error procesando representantes: {e}")
                ErrorMessageManager.log_error('unexpected_error', {'error': str(e)}, technical_details=str(e))
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'unexpected_error', 
                    {'error': f'procesamiento de representantes - {str(e)}'}
                )
                return None
            
            # Validar resultado del procesamiento
            if not isinstance(processed_actores, list) or not isinstance(processed_demandados, list):
                print("[ERROR] El procesamiento de representantes no devolvió listas válidas")
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'unexpected_error', 
                    {'error': 'resultado del procesamiento de representantes inválido'}
                )
                return None
            
            print(f"[DEBUG] Representantes procesados - Actores: {len(processed_actores)}, Demandados: {len(processed_demandados)}")
            
            return {
                'lista_actores': processed_actores,
                'lista_demandados': processed_demandados,
                'todos_los_roles': todos_los_roles
            }
            
        except Exception as e:
            print(f"[ERROR] Error crítico procesando representantes: {e}")
            ErrorMessageManager.log_error('unexpected_error', {'error': str(e)}, technical_details=str(e))
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root, 
                'unexpected_error', 
                {'error': f'procesamiento de representantes - {str(e)}'}
            )
            return None

    def _build_document_context(self, caso_data, processed_parties, agreement_details):
        """
        Construye el contexto para el documento.
        
        Args:
            caso_data: Datos del caso
            processed_parties: Partes con representantes procesados
            agreement_details: Detalles del acuerdo
            
        Returns:
            dict: Contexto del documento o None si hay error
        """
        try:
            print("[DEBUG] Construyendo contexto para la plantilla del documento")
            
            # Validar entradas
            if not all([isinstance(caso_data, dict), isinstance(processed_parties, dict), isinstance(agreement_details, dict)]):
                print("[ERROR] Parámetros de entrada no son diccionarios válidos")
                return None
            
            lista_actores = processed_parties.get('lista_actores', [])
            lista_demandados = processed_parties.get('lista_demandados', [])
            
            # Conversión segura de números a letras
            monto_letras = self._safe_number_to_words(agreement_details.get('monto_compensacion_numeros', ''))
            plazo_letras = self._safe_period_to_words(agreement_details.get('plazo_pago_dias', '0'))
            
            # Obtener mes en español de manera segura
            mes_acuerdo = self._get_spanish_month()
            
            # Obtener nombre del actor principal
            nombre_actor = self._get_main_actor_name(lista_actores)
            
            # Construir contexto con valores seguros y completos
            context = {
                # Información básica del caso
                'NUMERO_EXPEDIENTE': caso_data.get('numero_expediente', 'SIN NÚMERO'),
                'CARATULA': caso_data.get('caratula', 'SIN CARÁTULA'),
                
                # Fechas del acuerdo
                'DIA_ACUERDO': datetime.date.today().day,
                'MES_ACUERDO': mes_acuerdo,
                'AÑO_ACUERDO': datetime.date.today().year,
                'FECHA_COMPLETA': f"{datetime.date.today().day} de {mes_acuerdo} de {datetime.date.today().year}",
                'FECHA_CORTA': datetime.date.today().strftime('%d/%m/%Y'),
                
                # Partes del caso
                'ACTORES': lista_actores,
                'DEMANDADOS': lista_demandados,
                'NOMBRE_ACTOR': nombre_actor,
                
                # Montos y plazos de compensación
                'MONTO_COMPENSACION_NUMEROS': agreement_details.get('monto_compensacion_numeros', '0'),
                'MONTO_COMPENSACION_LETRAS': monto_letras,
                'PLAZO_PAGO_DIAS': agreement_details.get('plazo_pago_dias', '0'),
                'PLAZO_PAGO_LETRAS': plazo_letras,
                
                # Datos bancarios del actor
                'BANCO_ACTOR': agreement_details.get('banco_actor', ''),
                'CBU_ACTOR': agreement_details.get('cbu_actor', ''),
                'ALIAS_ACTOR': agreement_details.get('alias_actor', ''),
                'CUIT_ACTOR': agreement_details.get('cuit_actor', ''),
                
                # Variables de honorarios (compatibilidad total con plantilla)
                'MONTO_HONORARIOS_NUMEROS': '0',
                'MONTO_HONORARIOS_LETRAS': 'CERO',
                'HONORARIOS_LETRAS': 'CERO',  # Variable específica de tu plantilla
                'PLAZO_PAGO_HONORARIOS_DIAS': '0',
                
                # Variables de ubicación y mediación
                'LUGAR_ACUERDO': 'Ciudad Autónoma de Buenos Aires',
                'MEDIADOR': 'Mediador designado',
                'CENTRO_MEDIACION': 'Centro de Mediación',
                
                # Variables adicionales opcionales
                'OBSERVACIONES': '',
                'CLAUSULAS_ADICIONALES': '',
                
                # Variable auxiliar para representantes
                'rep': {}
            }
            
            # Validar contexto construido
            if not self._validate_document_context(context):
                return None
            
            print(f"[DEBUG] Contexto construido exitosamente con {len(context)} campos")
            return context
            
        except Exception as e:
            print(f"[ERROR] Error construyendo contexto del documento: {e}")
            ErrorMessageManager.log_error('unexpected_error', {'error': str(e)}, technical_details=str(e))
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root, 
                'unexpected_error', 
                {'error': f'construcción del contexto del documento - {str(e)}'}
            )
            return None

    def _safe_number_to_words(self, number_str):
        """Convierte número a palabras de manera segura."""
        try:
            return convertir_numero_a_letras(number_str) or "ERROR EN CONVERSIÓN"
        except Exception as e:
            print(f"[WARNING] Error convirtiendo monto a letras: {e}")
            return "ERROR EN CONVERSIÓN"

    def _safe_period_to_words(self, period_str):
        """Convierte período a palabras de manera segura."""
        try:
            return convertir_plazo_a_letras(period_str) or "CERO"
        except Exception as e:
            print(f"[WARNING] Error convirtiendo plazo a letras: {e}")
            return "CERO"

    def _get_spanish_month(self):
        """Obtiene el mes actual en español de manera segura."""
        try:
            import locale
            locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
            return datetime.date.today().strftime('%B')
        except:
            # Fallback a nombres de meses en español
            meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                    'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
            return meses[datetime.date.today().month - 1]

    def _get_main_actor_name(self, lista_actores):
        """Obtiene el nombre del actor principal de manera segura."""
        try:
            if lista_actores and isinstance(lista_actores[0], dict):
                return lista_actores[0].get('nombre_completo', 'SIN ACTOR')
            return 'SIN ACTOR'
        except Exception as e:
            print(f"[WARNING] Error obteniendo nombre del actor principal: {e}")
            return 'SIN ACTOR'

    def _format_parties_text(self, parties_list, party_type):
        """
        Formatea una lista de partes como texto para la plantilla.

        Args:
            parties_list: Lista de diccionarios con datos de partes
            party_type: Tipo de parte ('ACTOR' o 'DEMANDADO')

        Returns:
            str: Texto formateado de las partes
        """
        try:
            if not parties_list:
                return f"SIN {party_type}ES"

            formatted_parts = []
            for i, parte in enumerate(parties_list, 1):
                if isinstance(parte, dict) and 'nombre_completo' in parte:
                    nombre = parte['nombre_completo'].strip()
                    if nombre:
                        # Agregar número si hay múltiples partes
                        if len(parties_list) > 1:
                            formatted_parts.append(f"{i}. {nombre}")
                        else:
                            formatted_parts.append(nombre)

            if formatted_parts:
                return "\n".join(formatted_parts)
            else:
                return f"SIN {party_type}ES"

        except Exception as e:
            print(f"[WARNING] Error formateando texto de partes {party_type}: {e}")
            return f"ERROR EN {party_type}ES"

    def _validate_document_context(self, context):
        """Valida el contexto del documento."""
        try:
            required_fields = ['NUMERO_EXPEDIENTE', 'CARATULA', 'ACTORES', 'DEMANDADOS', 'MONTO_COMPENSACION_NUMEROS']
            
            for field in required_fields:
                if field not in context:
                    print(f"[ERROR] Campo requerido faltante en contexto: {field}")
                    return False
            
            return True
        except Exception as e:
            print(f"[ERROR] Error validando contexto del documento: {e}")
            return False

    def _generate_and_validate_document(self, document_context):
        """
        Genera y valida el documento.
        
        Args:
            document_context: Contexto del documento
            
        Returns:
            DocxTemplate: Documento generado o None si hay error
        """
        try:
            # Usar la plantilla original que soporta múltiples partes y representantes
            template_path = 'plantillas/mediacion/acuerdo_base.docx'
            print(f"[DEBUG] Generando documento con plantilla: {template_path}")
            
            # Validar plantilla y contexto
            template_validation = self._validate_template_and_context(template_path, document_context)
            
            if not template_validation['success']:
                first_error = template_validation['errors'][0] if template_validation['errors'] else {}
                error_details = {'error': first_error.get('message', 'Error desconocido')}
                
                ErrorMessageManager.log_error('template_error', error_details)
                ErrorMessageManager.show_error_dialog(
                    self.app_controller.root, 
                    'template_error', 
                    error_details
                )
                return None
            
            # Mostrar advertencias si las hay
            if template_validation['warnings']:
                warning_messages = [f"• {w['message']}" for w in template_validation['warnings']]
                warning_msg = "Advertencias encontradas:\n\n" + "\n".join(warning_messages)
                print(f"[WARNING] {warning_msg}")
            
            # Generar documento
            doc = DocxTemplate(template_path)
            required_vars = doc.get_undeclared_template_variables()
            safe_context = self._prepare_safe_context(document_context, required_vars)
            
            print("[DEBUG] Renderizando documento...")
            doc.render(safe_context)
            print("[DEBUG] Documento renderizado exitosamente")
            
            return doc
            
        except Exception as e:
            print(f"[ERROR] Error generando documento: {e}")
            ErrorMessageManager.log_error('template_error', {'error': str(e)}, technical_details=str(e))
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root, 
                'template_error', 
                {'error': f'generación del documento - {str(e)}'}
            )
            return None

    def _save_document(self, document, caso_data):
        """
        Guarda el documento generado con lógica inteligente de carpetas.

        Args:
            document: Documento a guardar
            caso_data: Datos del caso

        Returns:
            bool: True si se guardó exitosamente
        """
        try:
            case_id = caso_data.get('id')
            ruta_carpeta = caso_data.get('ruta_carpeta') or ''
            ruta_carpeta = ruta_carpeta.strip() if ruta_carpeta else ''

            # CASO A: Carpeta asignada y válida
            if ruta_carpeta and self._is_valid_folder_path(ruta_carpeta):
                print(f"[DEBUG] Usando carpeta asignada del caso: {ruta_carpeta}")

                # Construir ruta completa automáticamente
                filename = self._generate_document_filename(caso_data, "Acuerdo Mediacion")
                ruta_guardado = os.path.join(ruta_carpeta, filename)

                print(f"[DEBUG] Guardando documento automáticamente en: {ruta_guardado}")

                # Guardar documento
                document.save(ruta_guardado)

                # Verificar guardado
                if self._verify_document_saved(ruta_guardado):
                    file_size = os.path.getsize(ruta_guardado)
                    print(f"[DEBUG] Documento guardado exitosamente ({file_size:,} bytes)")

                    # Mostrar mensaje de confirmación
                    messagebox.showinfo(
                        "Documento Guardado",
                        f"El documento '{filename}' ha sido guardado en la carpeta del caso.\n\nUbicación: {ruta_guardado}",
                        parent=self.app_controller.root
                    )

                    # Ofrecer abrir la carpeta
                    self._offer_to_open_folder(ruta_carpeta, filename)
                    return True
                else:
                    raise Exception("No se pudo verificar el guardado del documento")

            # CASO B: Carpeta NO asignada o inválida
            else:
                print("[DEBUG] Carpeta del caso no asignada o inválida, solicitando selección")

                # Mostrar mensaje informativo
                messagebox.showwarning(
                    "Carpeta No Asignada",
                    "Este caso aún no tiene una carpeta de documentos asignada.\nPor favor, seleccione una carpeta a continuación.",
                    parent=self.app_controller.root
                )

                # Solicitar selección de carpeta
                initial_dir = ruta_carpeta if ruta_carpeta else os.path.expanduser("~")
                nueva_carpeta = filedialog.askdirectory(
                    title="Seleccionar Carpeta para Documentos del Caso",
                    initialdir=initial_dir,
                    parent=self.app_controller.root
                )

                if not nueva_carpeta:
                    print("[INFO] Usuario canceló la selección de carpeta")
                    return False

                # Actualizar base de datos con la nueva ruta
                if self.db.update_case_folder(case_id, nueva_carpeta):
                    print(f"[DEBUG] Carpeta actualizada en BD: {nueva_carpeta}")

                    # Actualizar caso_data en memoria
                    caso_data['ruta_carpeta'] = nueva_carpeta

                    # Proceder con el guardado usando la nueva carpeta
                    filename = self._generate_document_filename(caso_data, "Acuerdo Mediacion")
                    ruta_guardado = os.path.join(nueva_carpeta, filename)

                    print(f"[DEBUG] Guardando documento en nueva carpeta: {ruta_guardado}")

                    # Guardar documento
                    document.save(ruta_guardado)

                    # Verificar guardado
                    if self._verify_document_saved(ruta_guardado):
                        file_size = os.path.getsize(ruta_guardado)
                        print(f"[DEBUG] Documento guardado exitosamente ({file_size:,} bytes)")

                        # Mostrar mensaje de confirmación
                        messagebox.showinfo(
                            "Documento Guardado",
                            f"El documento '{filename}' ha sido guardado en la nueva carpeta del caso.\n\nUbicación: {ruta_guardado}",
                            parent=self.app_controller.root
                        )

                        # Ofrecer abrir la carpeta
                        self._offer_to_open_folder(nueva_carpeta, filename)
                        return True
                    else:
                        raise Exception("No se pudo verificar el guardado del documento")
                else:
                    raise Exception("No se pudo actualizar la ruta de carpeta en la base de datos")

        except PermissionError as e:
            print(f"[ERROR] Error de permisos guardando documento: {e}")
            ErrorMessageManager.log_error('file_permission', {'path': ruta_guardado if 'ruta_guardado' in locals() else 'desconocida', 'error': str(e)})
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root,
                'file_permission',
                {'path': ruta_guardado if 'ruta_guardado' in locals() else 'desconocida'}
            )
            return False

        except Exception as e:
            print(f"[ERROR] Error guardando documento: {e}")
            ErrorMessageManager.log_error('unexpected_error', {'error': str(e)}, technical_details=str(e))
            ErrorMessageManager.show_error_dialog(
                self.app_controller.root,
                'unexpected_error',
                {'error': f'guardado del documento - {str(e)}'}
            )
            return False

    def _generar_documento_con_datos(self, caso_id, details_acuerdo, processed_parties_data=None):
        """
        Genera un documento de mediación con datos proporcionados (lógica pura sin UI).

        Este método contiene toda la lógica de generación de documentos extraída del método
        original, pero sin dependencias de diálogos de Tkinter. Puede ser utilizado
        independientemente por herramientas de agente.

        Args:
            caso_id: ID del caso para generar el acuerdo
            details_acuerdo (dict): Diccionario con los detalles del acuerdo que incluye:
                - monto_compensacion_numeros (str): Monto en números
                - plazo_pago_dias (str): Plazo de pago en días
                - banco_actor (str): Banco del actor
                - cbu_actor (str): CBU del actor
                - alias_actor (str): Alias del actor
                - cuit_actor (str): CUIT del actor
            processed_parties_data (dict): Datos de partes procesadas con representantes (opcional)

        Returns:
            dict: Resultado de la operación con estructura:
                {
                    'success': bool,
                    'document': DocxTemplate o None,
                    'document_context': dict o None,
                    'filename_suggestion': str,
                    'error_message': str o None,
                    'error_type': str o None
                }
        """
        # Initialize timing and logging
        operation_start_time = time.time()
        operation_id = f"doc_gen_{caso_id}_{int(operation_start_time)}"
        
        self.logger.info(f"[{operation_id}] Starting document generation for case ID: {caso_id}")
        self.logger.debug(f"[{operation_id}] Input parameters: caso_id={caso_id}, details_keys={list(details_acuerdo.keys()) if isinstance(details_acuerdo, dict) else 'invalid'}")
        
        # Performance tracking
        phase_times = {}
        
        try:
            # 1. Validación inicial de parámetros
            phase_start = time.time()
            self.logger.info(f"[{operation_id}] Phase 1: Parameter validation")
            
            validation_result = self._validate_pure_parameters(caso_id, details_acuerdo)
            phase_times['parameter_validation'] = time.time() - phase_start
            
            if not validation_result['success']:
                self.logger.error(f"[{operation_id}] Parameter validation failed: {validation_result.get('error_message', 'Unknown error')}")
                self.logger.info(f"[{operation_id}] Operation failed after {time.time() - operation_start_time:.3f}s")
                return validation_result
            
            self.logger.debug(f"[{operation_id}] Parameter validation completed in {phase_times['parameter_validation']:.3f}s")
            
            # 2. Validación de dependencias del sistema (sin UI)
            phase_start = time.time()
            self.logger.info(f"[{operation_id}] Phase 2: System dependencies validation")
            
            if not self._validate_system_dependencies_pure():
                phase_times['dependencies_validation'] = time.time() - phase_start
                self.logger.error(f"[{operation_id}] System dependencies validation failed after {phase_times['dependencies_validation']:.3f}s")
                self.logger.info(f"[{operation_id}] Operation failed after {time.time() - operation_start_time:.3f}s")
                return {
                    'success': False,
                    'document': None,
                    'document_context': None,
                    'filename_suggestion': '',
                    'error_message': 'Dependencias del sistema no disponibles',
                    'error_type': 'missing_dependencies'
                }
            
            phase_times['dependencies_validation'] = time.time() - phase_start
            self.logger.debug(f"[{operation_id}] Dependencies validation completed in {phase_times['dependencies_validation']:.3f}s")
            
            # 3. Recopilación y validación de datos del caso
            phase_start = time.time()
            self.logger.info(f"[{operation_id}] Phase 3: Case data collection and validation")
            
            caso_data = self._collect_and_validate_case_data_pure(caso_id)
            phase_times['case_data_collection'] = time.time() - phase_start
            
            if not caso_data['success']:
                self.logger.error(f"[{operation_id}] Case data collection failed: {caso_data.get('error_message', 'Unknown error')} (type: {caso_data.get('error_type', 'unknown')})")
                self.logger.info(f"[{operation_id}] Operation failed after {time.time() - operation_start_time:.3f}s")
                return {
                    'success': False,
                    'document': None,
                    'document_context': None,
                    'filename_suggestion': '',
                    'error_message': caso_data['error_message'],
                    'error_type': caso_data['error_type']
                }
            
            self.logger.debug(f"[{operation_id}] Case data collection completed in {phase_times['case_data_collection']:.3f}s")
            self.logger.debug(f"[{operation_id}] Case info: {caso_data['data'].get('caratula', 'No caratula')}")
            
            # 4. Recopilación y validación de partes/roles
            phase_start = time.time()
            self.logger.info(f"[{operation_id}] Phase 4: Parties data collection and validation")

            if processed_parties_data:
                # Usar datos de partes procesadas proporcionados
                parties_data = {
                    'success': True,
                    'data': processed_parties_data,
                    'error_message': None,
                    'error_type': None
                }
                self.logger.debug(f"[{operation_id}] Using provided processed parties data")
            else:
                # Recopilar datos de partes internamente
                parties_data = self._collect_and_validate_parties_data_pure(caso_id)

            phase_times['parties_data_collection'] = time.time() - phase_start

            if not parties_data['success']:
                self.logger.error(f"[{operation_id}] Parties data collection failed: {parties_data.get('error_message', 'Unknown error')} (type: {parties_data.get('error_type', 'unknown')})")
                self.logger.info(f"[{operation_id}] Operation failed after {time.time() - operation_start_time:.3f}s")
                return {
                    'success': False,
                    'document': None,
                    'document_context': None,
                    'filename_suggestion': '',
                    'error_message': parties_data['error_message'],
                    'error_type': parties_data['error_type']
                }

            self.logger.debug(f"[{operation_id}] Parties data collection completed in {phase_times['parties_data_collection']:.3f}s")
            self.logger.debug(f"[{operation_id}] Parties count: {len(parties_data['data']['lista_actores'])} actors, {len(parties_data['data']['lista_demandados'])} defendants")
            
            # 5. Procesamiento de representantes
            phase_start = time.time()
            self.logger.info(f"[{operation_id}] Phase 5: Representatives processing")

            if processed_parties_data:
                # Usar datos de partes ya procesadas
                processed_parties = {
                    'success': True,
                    'data': processed_parties_data,
                    'error_message': None,
                    'error_type': None
                }
                self.logger.debug(f"[{operation_id}] Using provided processed parties data (representatives already processed)")
            else:
                # Procesar representantes internamente
                processed_parties = self._process_representatives_pure(parties_data['data'], caso_id)

            phase_times['representatives_processing'] = time.time() - phase_start

            if not processed_parties['success']:
                self.logger.error(f"[{operation_id}] Representatives processing failed: {processed_parties.get('error_message', 'Unknown error')} (type: {processed_parties.get('error_type', 'unknown')})")
                self.logger.info(f"[{operation_id}] Operation failed after {time.time() - operation_start_time:.3f}s")
                return {
                    'success': False,
                    'document': None,
                    'document_context': None,
                    'filename_suggestion': '',
                    'error_message': processed_parties['error_message'],
                    'error_type': processed_parties['error_type']
                }

            self.logger.debug(f"[{operation_id}] Representatives processing completed in {phase_times['representatives_processing']:.3f}s")
            
            # 6. Construcción del contexto del documento
            phase_start = time.time()
            self.logger.info(f"[{operation_id}] Phase 6: Document context building")
            
            document_context = self._build_document_context_pure(
                caso_data['data'], 
                processed_parties['data'], 
                details_acuerdo
            )
            phase_times['document_context_building'] = time.time() - phase_start
            
            if not document_context['success']:
                self.logger.error(f"[{operation_id}] Document context building failed: {document_context.get('error_message', 'Unknown error')} (type: {document_context.get('error_type', 'unknown')})")
                self.logger.info(f"[{operation_id}] Operation failed after {time.time() - operation_start_time:.3f}s")
                return {
                    'success': False,
                    'document': None,
                    'document_context': None,
                    'filename_suggestion': '',
                    'error_message': document_context['error_message'],
                    'error_type': document_context['error_type']
                }
            
            self.logger.debug(f"[{operation_id}] Document context building completed in {phase_times['document_context_building']:.3f}s")
            self.logger.debug(f"[{operation_id}] Context keys: {list(document_context['data'].keys()) if isinstance(document_context['data'], dict) else 'invalid'}")
            
            # 7. Generación y validación del documento
            phase_start = time.time()
            self.logger.info(f"[{operation_id}] Phase 7: Document generation and validation")
            
            document_result = self._generate_and_validate_document_pure(document_context['data'])
            phase_times['document_generation'] = time.time() - phase_start
            
            if not document_result['success']:
                self.logger.error(f"[{operation_id}] Document generation failed: {document_result.get('error_message', 'Unknown error')} (type: {document_result.get('error_type', 'unknown')})")
                self.logger.info(f"[{operation_id}] Operation failed after {time.time() - operation_start_time:.3f}s")
                return {
                    'success': False,
                    'document': None,
                    'document_context': document_context['data'],
                    'filename_suggestion': '',
                    'error_message': document_result['error_message'],
                    'error_type': document_result['error_type']
                }
            
            self.logger.debug(f"[{operation_id}] Document generation completed in {phase_times['document_generation']:.3f}s")
            
            # 8. Generar sugerencia de nombre de archivo
            phase_start = time.time()
            self.logger.info(f"[{operation_id}] Phase 8: Filename generation")
            
            filename_suggestion = self._generate_safe_filename(caso_data['data'])
            phase_times['filename_generation'] = time.time() - phase_start
            
            # Operación completada exitosamente
            total_duration = time.time() - operation_start_time
            
            # Log performance summary
            self.logger.info(f"[{operation_id}] Document generation completed successfully in {total_duration:.3f}s")
            self.logger.info(f"[{operation_id}] Performance breakdown:")
            for phase, duration in phase_times.items():
                percentage = (duration / total_duration) * 100
                self.logger.info(f"[{operation_id}]   - {phase}: {duration:.3f}s ({percentage:.1f}%)")
            
            self.logger.debug(f"[{operation_id}] Generated filename: {filename_suggestion}")
            
            return {
                'success': True,
                'document': document_result['document'],
                'document_context': document_context['data'],
                'filename_suggestion': filename_suggestion,
                'error_message': None,
                'error_type': None,
                'performance_metrics': {
                    'total_duration': total_duration,
                    'phase_times': phase_times
                }
            }
            
        except Exception as e:
            # Manejo de errores críticos no capturados
            total_duration = time.time() - operation_start_time
            error_msg = f"Error crítico en generación de documento: {str(e)}"
            
            self.logger.error(f"[{operation_id}] Critical error after {total_duration:.3f}s: {error_msg}")
            self.logger.error(f"[{operation_id}] Stack trace: {traceback.format_exc()}")
            
            # Log partial performance metrics if available
            if phase_times:
                self.logger.debug(f"[{operation_id}] Partial performance before failure:")
                for phase, duration in phase_times.items():
                    self.logger.debug(f"[{operation_id}]   - {phase}: {duration:.3f}s")
            
            return {
                'success': False,
                'document': None,
                'document_context': None,
                'filename_suggestion': '',
                'error_message': error_msg,
                'error_type': 'unexpected_error',
                'performance_metrics': {
                    'total_duration': total_duration,
                    'phase_times': phase_times,
                    'failed': True
                }
            }

    def _validate_pure_parameters(self, caso_id, details_acuerdo):
        """
        Valida los parámetros para la generación pura de documentos.
        
        Args:
            caso_id: ID del caso
            details_acuerdo: Diccionario con detalles del acuerdo
            
        Returns:
            dict: Resultado de la validación
        """
        try:
            # Validar caso_id
            if caso_id is None:
                return {
                    'success': False,
                    'error_message': 'ID del caso es requerido',
                    'error_type': 'validation_error'
                }
            
            try:
                int(caso_id)
            except (ValueError, TypeError):
                return {
                    'success': False,
                    'error_message': 'ID del caso debe ser un número válido',
                    'error_type': 'validation_error'
                }
            
            # Validar details_acuerdo
            if not isinstance(details_acuerdo, dict):
                return {
                    'success': False,
                    'error_message': 'Los detalles del acuerdo deben ser un diccionario',
                    'error_type': 'validation_error'
                }
            
            # Validar campos requeridos en details_acuerdo
            required_fields = [
                'monto_compensacion_numeros',
                'plazo_pago_dias',
                'banco_actor',
                'cbu_actor',
                'alias_actor',
                'cuit_actor'
            ]
            
            for field in required_fields:
                if field not in details_acuerdo:
                    return {
                        'success': False,
                        'error_message': f'Campo requerido faltante: {field}',
                        'error_type': 'validation_error'
                    }
            
            # Validar que tenemos acceso a la base de datos
            if not hasattr(self, 'db') or not self.db:
                return {
                    'success': False,
                    'error_message': 'No hay conexión a la base de datos',
                    'error_type': 'database_error'
                }
            
            self.logger.debug("Parameters validated successfully for pure generation")
            return {'success': True}
            
        except Exception as e:
            return {
                'success': False,
                'error_message': f'Error validando parámetros: {str(e)}',
                'error_type': 'validation_error'
            }

    def _validate_system_dependencies_pure(self):
        """
        Valida las dependencias del sistema sin mostrar diálogos de UI.
        
        Returns:
            bool: True si todas las dependencias están disponibles
        """
        try:
            self.logger.debug("Validating system dependencies (pure mode)")
            
            validation_result = self._validate_mediation_dependencies()
            if not validation_result['success']:
                # Log errors but don't show UI dialogs
                for error in validation_result['errors']:
                    self.logger.error(f"Missing dependency: {error.get('message', 'Unknown error')}")
                return False
            
            # Log warnings but continue
            if validation_result['warnings']:
                for warning in validation_result['warnings']:
                    self.logger.warning(f"Dependency warning: {warning.get('message', 'Unknown warning')}")
            
            self.logger.debug("System dependencies validated successfully (pure mode)")
            return True
            
        except Exception as e:
            self.logger.error(f"Error validating system dependencies (pure mode): {e}")
            return False

    def _collect_and_validate_case_data_pure(self, caso_id):
        """
        Recopila y valida los datos del caso sin mostrar diálogos de UI.
        
        Args:
            caso_id: ID del caso
            
        Returns:
            dict: Resultado con datos del caso o error
        """
        try:
            self.logger.debug(f"Collecting case data for ID: {caso_id} (pure mode)")
            
            caso = self.db.get_case_by_id(caso_id)
            if not caso:
                self.logger.warning(f"Case not found: {caso_id}")
                return {
                    'success': False,
                    'data': None,
                    'error_message': f'No se encontró el caso con ID: {caso_id}',
                    'error_type': 'missing_case'
                }
            
            # Validar datos críticos del caso
            if not isinstance(caso, dict):
                self.logger.error(f"Invalid case data structure for case {caso_id}: {type(caso)}")
                return {
                    'success': False,
                    'data': None,
                    'error_message': 'Estructura de datos del caso inválida',
                    'error_type': 'data_structure_error'
                }
            
            # Validar carátula
            caratula = caso.get('caratula', '').strip()
            if not caratula:
                self.logger.error(f"Case {caso_id} has no caratula defined")
                return {
                    'success': False,
                    'data': None,
                    'error_message': 'El caso no tiene carátula definida',
                    'error_type': 'empty_caratula'
                }
            
            # Validar otros campos críticos
            numero_expediente_raw = caso.get('numero_expediente')
            if numero_expediente_raw is None:
                numero_expediente = ''
            else:
                numero_expediente = str(numero_expediente_raw).strip()
            if not numero_expediente:
                self.logger.warning(f"Case without expediente number: {caratula}")
            
            self.logger.debug(f"Case validated (pure mode): {caratula}")
            return {
                'success': True,
                'data': caso,
                'error_message': None,
                'error_type': None
            }
            
        except Exception as e:
            error_msg = f"Error collecting case data: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            return {
                'success': False,
                'data': None,
                'error_message': error_msg,
                'error_type': 'unexpected_error'
            }

    def _collect_and_validate_parties_data_pure(self, caso_id):
        """
        Recopila y valida los datos de las partes sin mostrar diálogos de UI.
        
        Args:
            caso_id: ID del caso
            
        Returns:
            dict: Resultado con datos de partes o error
        """
        try:
            self.logger.debug(f"Collecting parties data for case ID: {caso_id} (pure mode)")
            
            todos_los_roles = self.db.get_parties_by_case_id(caso_id)
            if not todos_los_roles:
                self.logger.warning(f"No parties found for case {caso_id}")
                return {
                    'success': False,
                    'data': None,
                    'error_message': 'No se encontraron partes asociadas al caso',
                    'error_type': 'no_parties'
                }
            
            if not isinstance(todos_los_roles, list):
                self.logger.error(f"Invalid parties data structure for case {caso_id}: {type(todos_los_roles)}")
                return {
                    'success': False,
                    'data': None,
                    'error_message': 'Estructura de datos de roles inválida',
                    'error_type': 'data_structure_error'
                }
            
            # Separar actores y demandados
            lista_actores = [
                parte for parte in todos_los_roles 
                if parte.get('rol', '').lower() in ['actor', 'actora']
            ]
            
            lista_demandados = [
                parte for parte in todos_los_roles 
                if parte.get('rol', '').lower() in ['demandado', 'demandada']
            ]
            
            # Validar que hay al menos un actor y un demandado
            if not lista_actores:
                self.logger.error(f"No actors found for case {caso_id}")
                return {
                    'success': False,
                    'data': None,
                    'error_message': 'No se encontraron actores en el caso',
                    'error_type': 'no_actors'
                }
            
            if not lista_demandados:
                self.logger.error(f"No defendants found for case {caso_id}")
                return {
                    'success': False,
                    'data': None,
                    'error_message': 'No se encontraron demandados en el caso',
                    'error_type': 'no_defendants'
                }
            
            parties_data = {
                'lista_actores': lista_actores,
                'lista_demandados': lista_demandados,
                'todos_los_roles': todos_los_roles
            }
            
            self.logger.debug(f"Parties validated (pure mode): {len(lista_actores)} actors, {len(lista_demandados)} defendants")
            return {
                'success': True,
                'data': parties_data,
                'error_message': None,
                'error_type': None
            }
            
        except Exception as e:
            error_msg = f"Error collecting parties data: {str(e)}"
            print(f"[ERROR] {error_msg}")
            return {
                'success': False,
                'data': None,
                'error_message': error_msg,
                'error_type': 'unexpected_error'
            }

    def _process_representatives_pure(self, parties_data, caso_id):
        """
        Procesa los representantes sin mostrar diálogos de UI.
        
        Args:
            parties_data: Datos de las partes
            caso_id: ID del caso
            
        Returns:
            dict: Resultado con partes procesadas o error
        """
        try:
            print("[DEBUG] Procesando representantes (modo puro)")
            
            # Para el modo puro, usamos asignación automática de representantes
            # sin mostrar diálogos de UI
            lista_actores = parties_data['lista_actores'].copy()
            lista_demandados = parties_data['lista_demandados'].copy()
            
            # Asignar representantes automáticamente
            self._assign_representatives_automatically(lista_actores, caso_id)
            self._assign_representatives_automatically(lista_demandados, caso_id)
            
            processed_data = {
                'lista_actores': lista_actores,
                'lista_demandados': lista_demandados
            }
            
            print("[DEBUG] Representantes procesados exitosamente (modo puro)")
            return {
                'success': True,
                'data': processed_data,
                'error_message': None,
                'error_type': None
            }
            
        except Exception as e:
            error_msg = f"Error procesando representantes: {str(e)}"
            print(f"[ERROR] {error_msg}")
            return {
                'success': False,
                'data': None,
                'error_message': error_msg,
                'error_type': 'unexpected_error'
            }

    def _assign_representatives_automatically(self, lista_partes, caso_id):
        """
        Asigna representantes automáticamente sin UI.
        
        Args:
            lista_partes: Lista de partes a procesar
            caso_id: ID del caso
        """
        try:
            for parte in lista_partes:
                parte_id = parte.get('id')
                if not parte_id:
                    continue
                
                # Obtener representantes existentes
                representantes = self.db.get_representatives_by_party_id(parte_id)
                
                if representantes:
                    # Usar el primer representante disponible
                    rep = representantes[0]
                    parte['representante'] = {
                        'nombre_completo': rep.get('nombre_completo', ''),
                        'matricula': rep.get('matricula', ''),
                        'domicilio': rep.get('domicilio', ''),
                        'telefono': rep.get('telefono', ''),
                        'email': rep.get('email', '')
                    }
                else:
                    # Sin representante
                    parte['representante'] = None
                    
        except Exception as e:
            print(f"[WARNING] Error en asignación automática de representantes: {e}")

    def _build_document_context_pure(self, caso_data, processed_parties, agreement_details):
        """
        Construye el contexto del documento sin dependencias de UI.

        Args:
            caso_data: Datos del caso
            processed_parties: Partes procesadas
            agreement_details: Detalles del acuerdo

        Returns:
            dict: Resultado con contexto o error
        """
        try:
            print("[DEBUG] Construyendo contexto para la plantilla del documento (modo puro)")
            print(f"[DEBUG] Agreement details keys: {list(agreement_details.keys()) if isinstance(agreement_details, dict) else 'invalid'}")
            print(f"[DEBUG] Agreement details values: {agreement_details}")

            # Validar entradas
            if not all([isinstance(caso_data, dict), isinstance(processed_parties, dict), isinstance(agreement_details, dict)]):
                return {
                    'success': False,
                    'data': None,
                    'error_message': 'Parámetros de entrada no son diccionarios válidos',
                    'error_type': 'validation_error'
                }

            lista_actores = processed_parties.get('lista_actores', [])
            lista_demandados = processed_parties.get('lista_demandados', [])

            # Conversión segura de números a letras
            monto_numeros = agreement_details.get('monto_compensacion_numeros', '')
            print(f"[DEBUG] Monto compensacion numeros: '{monto_numeros}'")
            monto_letras = self._safe_number_to_words(monto_numeros)
            print(f"[DEBUG] Monto letras convertido: '{monto_letras}'")

            plazo_dias = agreement_details.get('plazo_pago_dias', '0')
            print(f"[DEBUG] Plazo pago dias: '{plazo_dias}'")
            plazo_letras = self._safe_period_to_words(plazo_dias)
            print(f"[DEBUG] Plazo letras convertido: '{plazo_letras}'")

            # Obtener mes en español de manera segura
            mes_acuerdo = self._get_spanish_month()

            # Obtener nombre del actor principal
            nombre_actor = self._get_main_actor_name(lista_actores)

            # Formatear texto de actores y demandados
            actores_texto = self._format_parties_text(lista_actores, 'ACTOR')
            demandados_texto = self._format_parties_text(lista_demandados, 'DEMANDADO')

            # Preparar datos estructurados para el nuevo template
            case_id = caso_data.get('id')
            ACTORES = self._prepare_parties_with_representatives(lista_actores, 'ACTOR', case_id)
            DEMANDADOS = self._prepare_parties_with_representatives(lista_demandados, 'DEMANDADO', case_id)

            # Obtener detalles del caso
            detalles_caso = caso_data.get('notas', 'Sin detalles adicionales del caso.')

            # Debug logging para seguimiento de datos del acuerdo
            print("[DEBUG] Construyendo contexto del documento...")
            print(f"[DEBUG] Caso data keys: {list(caso_data.keys()) if isinstance(caso_data, dict) else 'invalid'}")
            print(f"[DEBUG] Agreement details keys: {list(agreement_details.keys()) if isinstance(agreement_details, dict) else 'invalid'}")
            print(f"[DEBUG] Agreement details values: {agreement_details}")
            print(f"[DEBUG] Monto compensacion numeros: '{agreement_details.get('monto_compensacion_numeros', 'NOT_FOUND')}'")
            print(f"[DEBUG] Monto letras convertido: '{monto_letras}'")
            print(f"[DEBUG] Plazo dias: '{agreement_details.get('plazo_pago_dias', 'NOT_FOUND')}'")
            print(f"[DEBUG] Plazo letras convertido: '{plazo_letras}'")
            print(f"[DEBUG] Lista actores: {len(lista_actores) if isinstance(lista_actores, list) else 'invalid'}")
            print(f"[DEBUG] Lista demandados: {len(lista_demandados) if isinstance(lista_demandados, list) else 'invalid'}")
            print(f"[DEBUG] Nombre actor: '{nombre_actor}'")

            # Construir contexto con valores seguros y completos
            context = {
                # Información básica del caso (uppercase para compatibilidad)
                'NUMERO_EXPEDIENTE': caso_data.get('numero_expediente', 'SIN NÚMERO'),
                'CARATULA': caso_data.get('caratula', 'SIN CARÁTULA'),

                # Fechas del acuerdo
                'DIA_ACUERDO': datetime.date.today().day,
                'MES_ACUERDO': mes_acuerdo,
                'AÑO_ACUERDO': datetime.date.today().year,
                'FECHA_COMPLETA': f"{datetime.date.today().day} de {mes_acuerdo} de {datetime.date.today().year}",
                'FECHA_CORTA': datetime.date.today().strftime('%d/%m/%Y'),

                # Partes del caso
                'ACTORES': lista_actores,
                'DEMANDADOS': lista_demandados,
                'NOMBRE_ACTOR': nombre_actor,

                # Montos y plazos de compensación
                'MONTO_COMPENSACION_NUMEROS': agreement_details.get('monto_compensacion_numeros', '0'),
                'MONTO_COMPENSACION_LETRAS': monto_letras,
                'PLAZO_PAGO_DIAS': agreement_details.get('plazo_pago_dias', '0'),
                'PLAZO_PAGO_LETRAS': plazo_letras,

                # Datos bancarios del actor
                'BANCO_ACTOR': agreement_details.get('banco_actor', ''),
                'CBU_ACTOR': agreement_details.get('cbu_actor', ''),
                'ALIAS_ACTOR': agreement_details.get('alias_actor', ''),
                'CUIT_ACTOR': agreement_details.get('cuit_actor', ''),

                # Variables de honorarios (compatibilidad total con plantilla)
                'MONTO_HONORARIOS_NUMEROS': '0',
                'MONTO_HONORARIOS_LETRAS': 'CERO',
                'HONORARIOS_LETRAS': 'CERO',
                'PLAZO_PAGO_HONORARIOS_DIAS': '0',

                # Variables de ubicación y mediación
                'LUGAR_ACUERDO': 'Ciudad Autónoma de Buenos Aires',
                'MEDIADOR': 'Mediador designado',
                'CENTRO_MEDIACION': 'Centro de Mediación',

                # Variables adicionales opcionales
                'OBSERVACIONES': '',
                'CLAUSULAS_ADICIONALES': '',

                # Variable auxiliar para representantes
                'rep': {},

                # Variables en minúsculas para coincidir con la plantilla
                'caratula': caso_data.get('caratula', 'SIN CARÁTULA'),
                'numero_expediente': caso_data.get('numero_expediente', 'SIN NÚMERO'),
                'juzgado': caso_data.get('juzgado', 'SIN JUZGADO'),
                'fecha_actual': datetime.date.today().strftime('%d/%m/%Y'),
                'actores_texto': actores_texto,
                'demandados_texto': demandados_texto,
                # Nuevos campos para template mejorado
                'ACTORES': ACTORES,
                'DEMANDADOS': DEMANDADOS,
                'detalles_caso': detalles_caso,
                'monto_compensacion_letras': monto_letras,
                'monto_compensacion_numeros': agreement_details.get('monto_compensacion_numeros', '0'),
                'plazo_pago_letras': plazo_letras,
                'plazo_pago_dias': agreement_details.get('plazo_pago_dias', '0'),
                'banco_actor': agreement_details.get('banco_actor', ''),
                'cbu_actor': agreement_details.get('cbu_actor', ''),
                'alias_actor': agreement_details.get('alias_actor', ''),
                'cuit_actor': agreement_details.get('cuit_actor', '')
            }

            # Debug logging del contexto construido
            print(f"[DEBUG] Contexto construido con {len(context)} variables")
            print(f"[DEBUG] Variables clave del contexto:")
            print(f"[DEBUG]   NUMERO_EXPEDIENTE: '{context.get('NUMERO_EXPEDIENTE', 'NOT_SET')}'")
            print(f"[DEBUG]   CARATULA: '{context.get('CARATULA', 'NOT_SET')}'")
            print(f"[DEBUG]   MONTO_COMPENSACION_NUMEROS: '{context.get('MONTO_COMPENSACION_NUMEROS', 'NOT_SET')}'")
            print(f"[DEBUG]   MONTO_COMPENSACION_LETRAS: '{context.get('MONTO_COMPENSACION_LETRAS', 'NOT_SET')}'")
            print(f"[DEBUG]   PLAZO_PAGO_DIAS: '{context.get('PLAZO_PAGO_DIAS', 'NOT_SET')}'")
            print(f"[DEBUG]   PLAZO_PAGO_LETRAS: '{context.get('PLAZO_PAGO_LETRAS', 'NOT_SET')}'")
            print(f"[DEBUG]   ACTORES count: {len(context.get('ACTORES', [])) if isinstance(context.get('ACTORES'), list) else 'invalid'}")
            print(f"[DEBUG]   DEMANDADOS count: {len(context.get('DEMANDADOS', [])) if isinstance(context.get('DEMANDADOS'), list) else 'invalid'}")

            # Validar contexto construido
            if not self._validate_document_context(context):
                return {
                    'success': False,
                    'data': None,
                    'error_message': 'El contexto del documento no pasó la validación',
                    'error_type': 'context_validation_error'
                }

            print(f"[DEBUG] Contexto construido exitosamente (modo puro) con {len(context)} campos")
            return {
                'success': True,
                'data': context,
                'error_message': None,
                'error_type': None
            }

        except Exception as e:
            error_msg = f"Error construyendo contexto del documento: {str(e)}"
            print(f"[ERROR] {error_msg}")
            return {
                'success': False,
                'data': None,
                'error_message': error_msg,
                'error_type': 'unexpected_error'
            }

    def _generate_and_validate_document_pure(self, document_context):
        """
        Genera y valida el documento sin mostrar diálogos de UI.

        Args:
            document_context: Contexto del documento

        Returns:
            dict: Resultado con documento o error
        """
        try:
            # Usar la plantilla original que soporta múltiples partes y representantes
            template_path = 'plantillas/mediacion/acuerdo_base.docx'
            print(f"[DEBUG] Generando documento con plantilla original: {template_path} (modo puro)")
            
            # Validar plantilla y contexto
            template_validation = self._validate_template_and_context(template_path, document_context)
            
            if not template_validation['success']:
                first_error = template_validation['errors'][0] if template_validation['errors'] else {}
                error_message = first_error.get('message', 'Error desconocido en plantilla')
                
                print(f"[ERROR] Error de plantilla: {error_message}")
                return {
                    'success': False,
                    'document': None,
                    'error_message': error_message,
                    'error_type': 'template_error'
                }
            
            # Log advertencias si las hay
            if template_validation['warnings']:
                for warning in template_validation['warnings']:
                    print(f"[WARNING] Advertencia de plantilla: {warning.get('message', 'Advertencia desconocida')}")
            
            # Generar documento
            doc = DocxTemplate(template_path)
            required_vars = doc.get_undeclared_template_variables()
            safe_context = self._prepare_safe_context(document_context, required_vars)
            
            print("[DEBUG] Renderizando documento (modo puro)...")
            doc.render(safe_context)
            print("[DEBUG] Documento renderizado exitosamente (modo puro)")
            
            return {
                'success': True,
                'document': doc,
                'error_message': None,
                'error_type': None
            }
            
        except Exception as e:
            error_msg = f"Error generando documento: {str(e)}"
            print(f"[ERROR] {error_msg}")
            return {
                'success': False,
                'document': None,
                'error_message': error_msg,
                'error_type': 'document_generation_error'
            }

    def _is_valid_folder_path(self, folder_path):
        """
        Valida si una ruta de carpeta existe y es accesible.

        Args:
            folder_path (str): Ruta de la carpeta a validar

        Returns:
            bool: True si la carpeta es válida y accesible
        """
        try:
            if not folder_path or not folder_path.strip():
                return False

            # Verificar que existe
            if not os.path.exists(folder_path):
                print(f"[DEBUG] Carpeta no existe: {folder_path}")
                return False

            # Verificar que es un directorio
            if not os.path.isdir(folder_path):
                print(f"[DEBUG] Ruta no es un directorio: {folder_path}")
                return False

            # Verificar permisos de escritura
            if not os.access(folder_path, os.W_OK):
                print(f"[DEBUG] Sin permisos de escritura en: {folder_path}")
                return False

            return True

        except Exception as e:
            print(f"[WARNING] Error validando carpeta {folder_path}: {e}")
            return False

    def _generate_document_filename(self, caso_data, document_type="Documento"):
        """
        Genera un nombre de archivo automático basado en los datos del caso.

        Args:
            caso_data (dict): Datos del caso
            document_type (str): Tipo de documento (ej: "Acuerdo Mediacion", "Escrito Generico")

        Returns:
            str: Nombre de archivo generado
        """
        try:
            # Obtener carátula del caso
            caratula = caso_data.get('caratula', '').strip()

            # Limpiar carátula para usar en nombre de archivo
            if caratula:
                # Remover caracteres problemáticos para nombres de archivo
                caratula_limpia = "".join(c for c in caratula if c.isalnum() or c in (' ', '-', '_')).rstrip()
                caratula_limpia = caratula_limpia.replace(' ', '_')
                if len(caratula_limpia) > 50:  # Limitar longitud
                    caratula_limpia = caratula_limpia[:50]
            else:
                caratula_limpia = "Sin_Caratula"

            # Obtener número de expediente si existe
            numero_exp = caso_data.get('numero_expediente', '').strip()
            if numero_exp:
                numero_exp_limpio = "".join(c for c in numero_exp if c.isalnum() or c in (' ', '-', '_')).rstrip()
                numero_exp_limpio = numero_exp_limpio.replace(' ', '_')
                filename = f"{document_type} - {numero_exp_limpio} - {caratula_limpia}.docx"
            else:
                # Usar fecha como fallback
                fecha = datetime.date.today().strftime('%Y%m%d')
                filename = f"{document_type} - {caratula_limpia} - {fecha}.docx"

            return filename

        except Exception as e:
            print(f"[WARNING] Error generando nombre de archivo: {e}")
            return f"{document_type} - {datetime.date.today().strftime('%Y%m%d')}.docx"

    def _verify_document_saved(self, file_path):
        """
        Verifica que el documento se guardó correctamente.

        Args:
            file_path (str): Ruta del archivo a verificar

        Returns:
            bool: True si el archivo existe y tiene contenido
        """
        try:
            if not os.path.exists(file_path):
                print(f"[ERROR] Archivo no encontrado después del guardado: {file_path}")
                return False

            file_size = os.path.getsize(file_path)
            if file_size == 0:
                print(f"[ERROR] Archivo guardado pero está vacío: {file_path}")
                return False

            print(f"[DEBUG] Archivo verificado correctamente: {file_path} ({file_size:,} bytes)")
            return True

        except Exception as e:
            print(f"[ERROR] Error verificando archivo guardado: {e}")
            return False

    def _offer_to_open_folder(self, folder_path, filename):
        """
        Ofrece al usuario abrir la carpeta donde se guardó el documento.

        Args:
            folder_path (str): Ruta de la carpeta
            filename (str): Nombre del archivo guardado
        """
        try:
            if messagebox.askyesno(
                "Abrir Carpeta",
                f"¿Desea abrir la carpeta donde se guardó el documento?\n\nCarpeta: {folder_path}\nArchivo: {filename}",
                parent=self.app_controller.root
            ):
                self._open_folder(folder_path)
        except Exception as e:
            print(f"[WARNING] Error ofreciendo abrir carpeta: {e}")

    def _open_folder(self, folder_path):
        """
        Abre la carpeta en el explorador de archivos.

        Args:
            folder_path (str): Ruta de la carpeta a abrir
        """
        try:
            import subprocess
            import sys

            print(f"[DEBUG] Intentando abrir carpeta: {folder_path}")

            if sys.platform.startswith('win'):
                os.startfile(folder_path)
            elif sys.platform.startswith('darwin'):
                subprocess.call(['open', folder_path])
            else:
                subprocess.call(['xdg-open', folder_path])

            print("[DEBUG] Carpeta abierta exitosamente")

        except Exception as e:
            print(f"[WARNING] No se pudo abrir la carpeta: {e}")
            messagebox.showwarning(
                "Advertencia",
                f"La carpeta se creó correctamente, pero no se pudo abrir automáticamente:\n{folder_path}\n\nPuede abrirla manualmente desde el explorador de archivos.",
                parent=self.app_controller.root
            )

    def _save_document_intelligent(self, document, caso_data, titulo_escrito):
        """
        Guarda un documento genérico usando lógica inteligente de carpetas.

        Args:
            document: Documento a guardar
            caso_data: Datos del caso
            titulo_escrito: Título del escrito

        Returns:
            str: Ruta donde se guardó el documento, o None si falló
        """
        try:
            case_id = caso_data.get('id')
            ruta_carpeta = caso_data.get('ruta_carpeta') or ''
            ruta_carpeta = ruta_carpeta.strip() if ruta_carpeta else ''

            # CASO A: Carpeta asignada y válida
            if ruta_carpeta and self._is_valid_folder_path(ruta_carpeta):
                print(f"[DEBUG] Usando carpeta asignada del caso para escrito genérico: {ruta_carpeta}")

                # Construir nombre de archivo para escrito genérico
                filename = self._generate_generic_document_filename(caso_data, titulo_escrito)
                ruta_guardado = os.path.join(ruta_carpeta, filename)

                print(f"[DEBUG] Guardando escrito genérico automáticamente en: {ruta_guardado}")

                # Guardar documento
                document.save(ruta_guardado)

                # Verificar guardado
                if self._verify_document_saved(ruta_guardado):
                    file_size = os.path.getsize(ruta_guardado)
                    print(f"[DEBUG] Escrito genérico guardado exitosamente ({file_size:,} bytes)")

                    # Mostrar mensaje de confirmación
                    messagebox.showinfo(
                        "Documento Guardado",
                        f"El escrito '{titulo_escrito}' ha sido guardado en la carpeta del caso.\n\nUbicación: {ruta_guardado}",
                        parent=self.app_controller.root
                    )

                    # Ofrecer abrir la carpeta
                    self._offer_to_open_folder(ruta_carpeta, filename)
                    return ruta_guardado
                else:
                    raise Exception("No se pudo verificar el guardado del escrito genérico")

            # CASO B: Carpeta NO asignada o inválida
            else:
                print("[DEBUG] Carpeta del caso no asignada o inválida para escrito genérico, solicitando selección")

                # Mostrar mensaje informativo
                messagebox.showwarning(
                    "Carpeta No Asignada",
                    "Este caso aún no tiene una carpeta de documentos asignada.\nPor favor, seleccione una carpeta a continuación.",
                    parent=self.app_controller.root
                )

                # Solicitar selección de carpeta
                initial_dir = ruta_carpeta if ruta_carpeta else os.path.expanduser("~")
                nueva_carpeta = filedialog.askdirectory(
                    title="Seleccionar Carpeta para Documentos del Caso",
                    initialdir=initial_dir,
                    parent=self.app_controller.root
                )

                if not nueva_carpeta:
                    print("[INFO] Usuario canceló la selección de carpeta para escrito genérico")
                    return None

                # Validar que la carpeta seleccionada sea válida
                if not self._is_valid_folder_path(nueva_carpeta):
                    messagebox.showerror(
                        "Carpeta Inválida",
                        "La carpeta seleccionada no es válida o no tiene permisos de escritura.",
                        parent=self.app_controller.root
                    )
                    return None

                # Actualizar base de datos con la nueva ruta
                if self.db.update_case_folder(case_id, nueva_carpeta):
                    print(f"[DEBUG] Carpeta actualizada en BD para escrito genérico: {nueva_carpeta}")

                    # Actualizar caso_data en memoria
                    caso_data['ruta_carpeta'] = nueva_carpeta

                    # Proceder con el guardado usando la nueva carpeta
                    filename = self._generate_generic_document_filename(caso_data, titulo_escrito)
                    ruta_guardado = os.path.join(nueva_carpeta, filename)

                    print(f"[DEBUG] Guardando escrito genérico en nueva carpeta: {ruta_guardado}")

                    # Guardar documento
                    document.save(ruta_guardado)

                    # Verificar guardado
                    if self._verify_document_saved(ruta_guardado):
                        file_size = os.path.getsize(ruta_guardado)
                        print(f"[DEBUG] Escrito genérico guardado exitosamente ({file_size:,} bytes)")

                        # Mostrar mensaje de confirmación
                        messagebox.showinfo(
                            "Documento Guardado",
                            f"El escrito '{titulo_escrito}' ha sido guardado en la nueva carpeta del caso.\n\nUbicación: {ruta_guardado}",
                            parent=self.app_controller.root
                        )

                        # Ofrecer abrir la carpeta
                        self._offer_to_open_folder(nueva_carpeta, filename)
                        return ruta_guardado
                    else:
                        raise Exception("No se pudo verificar el guardado del escrito genérico")
                else:
                    raise Exception("No se pudo actualizar la ruta de carpeta en la base de datos")

        except Exception as e:
            print(f"[ERROR] Error guardando escrito genérico: {e}")
            messagebox.showerror(
                "Error",
                f"Error al guardar el escrito genérico: {str(e)}",
                parent=self.app_controller.root
            )
            return None

    def _generate_generic_document_filename(self, caso_data, titulo_escrito):
        """
        Genera un nombre de archivo para documentos genéricos.

        Args:
            caso_data (dict): Datos del caso
            titulo_escrito (str): Título del escrito

        Returns:
            str: Nombre de archivo generado
        """
        try:
            # Obtener carátula del caso
            caratula = caso_data.get('caratula', '').strip()

            # Limpiar título del escrito
            titulo_limpio = "".join(c for c in titulo_escrito if c.isalnum() or c in (' ', '-', '_')).rstrip()
            titulo_limpio = titulo_limpio.replace(' ', '_')
            if len(titulo_limpio) > 30:  # Limitar longitud
                titulo_limpio = titulo_limpio[:30]

            # Limpiar carátula para usar en nombre de archivo
            if caratula:
                caratula_limpia = "".join(c for c in caratula if c.isalnum() or c in (' ', '-', '_')).rstrip()
                caratula_limpia = caratula_limpia.replace(' ', '_')
                if len(caratula_limpia) > 30:  # Limitar longitud
                    caratula_limpia = caratula_limpia[:30]
            else:
                caratula_limpia = "Sin_Caratula"

            # Obtener número de expediente si existe
            numero_exp = caso_data.get('numero_expediente', '').strip()
            if numero_exp:
                numero_exp_limpio = "".join(c for c in numero_exp if c.isalnum() or c in (' ', '-', '_')).rstrip()
                numero_exp_limpio = numero_exp_limpio.replace(' ', '_')
                filename = f"Borrador - {titulo_limpio} - {numero_exp_limpio} - {caratula_limpia}.docx"
            else:
                # Usar fecha como fallback
                fecha = datetime.date.today().strftime('%Y%m%d')
                filename = f"Borrador - {titulo_limpio} - {caratula_limpia} - {fecha}.docx"

            return filename

        except Exception as e:
            print(f"[WARNING] Error generando nombre de archivo para escrito genérico: {e}")
            return f"Borrador - {datetime.date.today().strftime('%Y%m%d')}.docx"

    def _prepare_parties_with_representatives(self, parties_list, party_type, case_id=None):
        """
        Prepara una lista estructurada de partes con sus representantes para el template.

        Args:
            parties_list (list): Lista de partes del caso
            party_type (str): Tipo de parte ('ACTOR' o 'DEMANDADO')
            case_id (int): ID del caso

        Returns:
            list: Lista de diccionarios con información de cada parte y sus representantes
        """
        parties_data = []

        for party in parties_list:
            # Información básica de la parte
            party_info = {
                'nombre_completo': party.get('nombre_completo', 'Sin nombre'),
                'dni': party.get('dni', ''),
                'cuit': party.get('cuit', ''),
                'domicilio_real': party.get('domicilio_real', ''),
                'domicilio_legal': party.get('domicilio_legal', ''),
                'representantes': []
            }

            # Buscar representantes de esta parte
            if case_id:
                representantes = self._get_representatives_for_party(party.get('id'), party.get('id'), case_id)
                party_info['representantes'] = representantes

            parties_data.append(party_info)

        return parties_data

    def _get_representatives_for_party(self, rol_id, party_id, case_id):
        """
        Obtiene los representantes legales de una parte específica.

        Args:
            rol_id (int): ID del rol de la parte
            party_id (int): ID de la parte
            case_id (int): ID del caso

        Returns:
            list: Lista de representantes con su información
        """
        representantes = []

        try:
            # Obtener roles relacionados con esta parte
            roles = db.get_roles_by_case_id(case_id)

            for rol in roles:
                # Verificar si este rol representa a la parte actual
                if rol.get('representa_a_id') == party_id or rol.get('representa_a_id') == rol_id:
                    representante_info = {
                        'nombre_completo': rol.get('nombre_completo', 'Sin nombre'),
                        'cuit': rol.get('cuit', ''),
                        'personeria': rol.get('datos_bancarios', 'Poder General Judicial')  # Usar datos_bancarios como personería por defecto
                    }
                    representantes.append(representante_info)

        except Exception as e:
            print(f"[WARNING] Error obteniendo representantes para parte {party_id}: {e}")

        return representantes

    def _generate_safe_filename(self, caso_data):
        """Genera un nombre de archivo seguro."""
        try:
            numero_exp = caso_data.get('numero_expediente', '').strip()
            if numero_exp:
                numero_exp_limpio = "".join(c for c in numero_exp if c.isalnum() or c in (' ', '-', '_')).rstrip()
                return f"Acuerdo_Mediacion_{numero_exp_limpio}.docx"
            else:
                return f"Acuerdo_Mediacion_{datetime.date.today().strftime('%Y%m%d')}.docx"
        except Exception as e:
            print(f"[WARNING] Error generando nombre de archivo: {e}")
            return f"Acuerdo_Mediacion_{datetime.date.today().strftime('%Y%m%d')}.docx"

    def _offer_to_open_document(self, ruta_guardado, file_size):
        """Ofrece al usuario abrir el documento generado."""
        try:
            if messagebox.askyesno(
                "Documento Generado", 
                f"Acuerdo generado exitosamente en:\n{ruta_guardado}\n\nTamaño: {file_size:,} bytes\n\n¿Desea abrir el documento ahora?",
                parent=self.app_controller.root
            ):
                self._open_document(ruta_guardado)
        except Exception as e:
            print(f"[WARNING] Error ofreciendo abrir documento: {e}")

    def _open_document(self, ruta_guardado):
        """Abre el documento generado."""
        try:
            import subprocess
            import sys
            print("[DEBUG] Intentando abrir el documento...")
            
            if sys.platform.startswith('win'):
                os.startfile(ruta_guardado)
            elif sys.platform.startswith('darwin'):
                subprocess.call(['open', ruta_guardado])
            else:
                subprocess.call(['xdg-open', ruta_guardado])
            
            print("[DEBUG] Documento abierto exitosamente")
            
        except Exception as e:
            print(f"[WARNING] No se pudo abrir el documento: {e}")
            messagebox.showwarning(
                "Advertencia", 
                f"El documento se guardó correctamente, pero no se pudo abrir automáticamente:\n{str(e)}\n\nPuede abrirlo manualmente desde:\n{ruta_guardado}",
                parent=self.app_controller.root
            )

    def _handle_critical_error(self, error, caso_id, operation_start_time):
        """
        Maneja errores críticos no capturados específicamente.
        
        Args:
            error: Excepción capturada
            caso_id: ID del caso
            operation_start_time: Tiempo de inicio de la operación
            
        Returns:
            bool: False (operación falló)
        """
        try:
            operation_end_time = datetime.datetime.now()
            duration = operation_end_time - operation_start_time
            
            error_msg = f"Error crítico al generar el acuerdo para caso ID {caso_id}: {str(error)}"
            print(f"[ERROR] {error_msg}")
            print(f"[ERROR] Operación falló después de {duration.total_seconds():.2f} segundos")
            
            # Mostrar información técnica adicional para debugging
            import traceback
            technical_details = traceback.format_exc()
            print(f"[ERROR] Detalles técnicos:\n{technical_details}")
            
            # Log del error
            ErrorMessageManager.log_error('unexpected_error', {'error': str(error), 'case_id': caso_id}, technical_details)
            
            # Mostrar error al usuario con información útil
            messagebox.showerror(
                "Error Crítico", 
                f"{error_msg}\n\nDuración: {duration.total_seconds():.2f} segundos\n\nSi el problema persiste, contacte al soporte técnico.\n\nDetalles técnicos guardados en el log del sistema.",
                parent=self.app_controller.root
            )
            
            return False
            
        except Exception as e:
            print(f"[ERROR] Error manejando error crítico: {e}")
            return False
        
    def save_case(
        self,
        case_id,
        cliente_id,
        caratula,
        num_exp,
        anio_car,
        juzgado,
        juris,
        notas,
        ruta,
        inact_days,
        inact_enabled,
        etiquetas_caso_str,
        dialog,
    ):
        """Guarda los datos del caso en la base de datos"""
        if not caratula.strip():
            messagebox.showwarning(
                "Advertencia",
                "La carátula del caso no puede estar vacía.",
                parent=dialog,
            )
            return

        success_main_data = False
        saved_case_id = case_id

        if case_id is None:  # Nuevo caso
            new_id = self.db.add_case(
                cliente_id,
                caratula.strip(),
                num_exp.strip(),
                anio_car.strip(),
                juzgado.strip(),
                juris.strip(),
                "",  # etapa_procesal - empty string as default
                notas.strip(),
                ruta.strip(),
                inact_days,
                inact_enabled,
            )
            if new_id:
                success_main_data = True
                saved_case_id = new_id
                msg_op = "agregado"
                self.app_controller.selected_case = self.db.get_case_by_id(saved_case_id)
            else:
                msg_op = "falló al agregar"
        else:  # Editar caso
            if self.db.update_case(
                case_id,
                caratula.strip(),
                num_exp.strip(),
                anio_car.strip(),
                juzgado.strip(),
                juris.strip(),
                "",  # etapa_procesal - empty string as default
                notas.strip(),
                ruta.strip(),
                inact_days,
                inact_enabled,
            ):
                success_main_data = True
                msg_op = "actualizado"
            # Actualizar datos del caso seleccionado si es el mismo
                if (
                    self.app_controller.selected_case
                    and self.app_controller.selected_case["id"] == case_id
                ):
                    self.app_controller.selected_case = self.db.get_case_by_id(case_id)
            else:
                msg_op = "falló al actualizar"

        if success_main_data:
        # Lógica para guardar etiquetas del caso
            if saved_case_id is not None:
                self._save_case_tags(saved_case_id, etiquetas_caso_str)

            messagebox.showinfo(
                "Éxito",
                f"Caso {msg_op} con éxito. Etiquetas actualizadas.",
                parent=self.app_controller.root,
            )   
            dialog.destroy()
            self.app_controller._refresh_open_case_window(saved_case_id)

        # Recargar la lista de casos del cliente actual
            if self.app_controller.selected_client:
                self.load_cases_by_client(self.app_controller.selected_client["id"])
        else:
            messagebox.showerror(
                "Error",
                f"No se pudo guardar la información principal del caso.",
                parent=dialog,
            )

    def _save_case_tags(self, case_id, etiquetas_str):
        """Guarda las etiquetas del caso"""
        # Obtener los nombres de las etiquetas ingresadas por el usuario
        nombres_etiquetas_nuevas = [
            tag.strip().lower() for tag in etiquetas_str.split(",") if tag.strip()
        ]

        # Obtener las etiquetas actualmente asignadas al caso desde la BD
        etiquetas_actuales_obj_db = self.db.get_etiquetas_de_caso(case_id)
        nombres_etiquetas_actuales_db = {
            e["nombre_etiqueta"].lower() for e in etiquetas_actuales_obj_db
        }

        # Determinar qué etiquetas añadir y cuáles quitar
        ids_etiquetas_a_asignar = set()
        for nombre_tag_nuevo in nombres_etiquetas_nuevas:
            tag_id = self.db.add_etiqueta(nombre_tag_nuevo)
            if tag_id:
                ids_etiquetas_a_asignar.add(tag_id)

        etiquetas_ids_actuales_db = {
            e["id_etiqueta"] for e in etiquetas_actuales_obj_db
        }

        # Etiquetas a asignar (nuevas o que ya estaban y deben permanecer)
        for tag_id_to_assign in ids_etiquetas_a_asignar:
            self.db.asignar_etiqueta_a_caso(case_id, tag_id_to_assign)

        # Etiquetas a quitar (estaban en BD pero no en la nueva lista del usuario)
        ids_etiquetas_a_quitar = etiquetas_ids_actuales_db - ids_etiquetas_a_asignar
        for tag_id_to_remove in ids_etiquetas_a_quitar:
            self.db.quitar_etiqueta_de_caso(case_id, tag_id_to_remove)

    def delete_case(self):
        """Elimina el caso seleccionado"""
        if not self.app_controller.selected_case:
            messagebox.showwarning(
                "Advertencia",
                "Selecciona un caso para eliminar.",
                parent=self.app_controller.root,
            )
            return

        case_id = self.app_controller.selected_case["id"]
        case_caratula = self.app_controller.selected_case.get(
            "caratula", f"ID {case_id}"
        )

        if messagebox.askyesno(
            "Confirmar",
            f"¿Eliminar caso '{case_caratula}' y TODAS sus actividades, audiencias y documentos asociados?",
            parent=self.app_controller.root,
            icon="warning",
        ):
            if self.db.delete_case(case_id):
                messagebox.showinfo(
                    "Éxito", "Caso eliminado.", parent=self.app_controller.root
                )
                # Recargar la lista de casos del cliente actual
                if self.app_controller.selected_client:
                    self.load_cases_by_client(self.app_controller.selected_client["id"])
                self.app_controller.actualizar_lista_audiencias()
                self.app_controller.marcar_dias_audiencias_calendario()
            else:
                messagebox.showerror(
                    "Error",
                    "No se pudo eliminar el caso.",
                    parent=self.app_controller.root,
                )

    def open_actividad_dialog(self, caso_id, actividad_id=None):
        """Abre un diálogo simple para agregar o editar una actividad del caso."""
        is_edit = actividad_id is not None
        dialog = tk.Toplevel(self.app_controller.root)
        dialog.title("Editar Actividad" if is_edit else "Agregar Actividad")
        dialog.transient(self.app_controller.root)
        dialog.grab_set()

        frame = ttk.Frame(dialog, padding="15")
        frame.pack(fill=tk.BOTH, expand=True)
        frame.columnconfigure(1, weight=1)

        # ---
        actividad_data = {}
        if is_edit:
            actividad_data = self.db.get_actividad_by_id(actividad_id)
            if not actividad_data:
                messagebox.showerror("Error", "No se encontró la actividad.")
                dialog.destroy()
                return
    
        # ---
        ttk.Label(frame, text="Tipo:").grid(row=0, column=0, sticky=tk.W, pady=3)
        tipo_var = tk.StringVar(value=actividad_data.get("tipo_actividad", ""))
        # Usamos un Combobox como habíamos diseñado
        tipos_comunes = [
            "Movimiento MEV",
            "Presentación de Escrito",
            "Notificación Recibida",
            "Cédula / Oficio",
            "Llamada Telefónica",
            "Reunión con Cliente",
            "Reunión con Contraparte",
            "Estudio de Causa",
            "Otro"
        ]
        tipo_combo = ttk.Combobox(frame, textvariable=tipo_var, values=tipos_comunes, width=38)
        tipo_combo.grid(row=0, column=1, sticky=tk.EW, pady=3)

        ttk.Label(frame, text="Descripción:").grid(row=1, column=0, sticky=tk.NW, pady=3)
        desc_text = tk.Text(frame, height=8, width=40, wrap=tk.WORD)
        desc_text.grid(row=1, column=1, sticky="nsew")
        desc_text.insert("1.0", actividad_data.get("descripcion", ""))
        frame.rowconfigure(1, weight=1)
    
        # ---
        def on_save():
            tipo = tipo_var.get().strip()
            descripcion = desc_text.get("1.0", tk.END).strip()
            if not tipo or not descripcion:
                messagebox.showwarning("Datos Incompletos", "El tipo y la descripción son obligatorios.", parent=dialog)
                return

            fecha_hora = datetime.datetime.now()
        
            if is_edit:
                success = self.db.update_actividad_caso(actividad_id, tipo, descripcion)
            else:
                success = self.db.add_actividad_caso(caso_id, fecha_hora, tipo, descripcion, self.app_controller.current_user)

            if success:
                self.app_controller._refresh_open_case_window(caso_id)
                dialog.destroy()
            else:
                messagebox.showerror("Error", "No se pudo guardar la actividad.", parent=dialog)

        # ---
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=2, column=1, sticky=tk.E, pady=10)
        ttk.Button(button_frame, text="Guardar", command=on_save).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Cancelar", command=dialog.destroy).pack(side=tk.LEFT)


    def load_cases_by_client(self, client_id):
        """Carga la lista de casos de un cliente en el TreeView"""
        try:
            # Validar parámetro de entrada
            if client_id is None or client_id == "":
                print(f"Error: client_id inválido: {client_id}")
                return False

            # Convertir a entero si es necesario
            try:
                client_id = int(client_id)
            except (ValueError, TypeError):
                print(f"Error: client_id debe ser un número entero válido: {client_id}")
                return False

            print(f"Cargando casos para cliente ID: {client_id}")

            # Limpiar estado anterior
            self.clear_case_list()
            self.app_controller.selected_case = None

            # Obtener casos de la base de datos
            try:
                cases = self.db.get_cases_by_client(client_id)
                if cases is None:
                    print(f"Error: No se pudieron obtener casos para cliente {client_id}")
                    return False
            except Exception as db_error:
                print(f"Error de base de datos al obtener casos para cliente {client_id}: {db_error}")
                return False

            print(f"Encontrados {len(cases)} casos para cliente {client_id}")

            # Procesar cada caso
            cases_loaded = 0
            for case in cases:
                try:
                    # Validar que el caso tenga los campos requeridos
                    if not isinstance(case, dict) or "id" not in case or "caratula" not in case:
                        print(f"Advertencia: Caso inválido encontrado: {case}")
                        continue

                    case_id = case["id"]
                    caratula = case.get("caratula", "Sin carátula")

                    # Obtener etiquetas del caso
                    try:
                        etiquetas_obj = self.db.get_etiquetas_de_caso(case_id)
                        if etiquetas_obj:
                            etiquetas_nombres = [e.get("nombre_etiqueta", "") for e in etiquetas_obj if e]
                            etiquetas_str = ", ".join(etiquetas_nombres) if etiquetas_nombres else ""
                        else:
                            etiquetas_str = ""
                    except Exception as tag_error:
                        print(f"Error al obtener etiquetas para caso {case_id}: {tag_error}")
                        etiquetas_str = ""

                    # Formatear número de expediente y año
                    num_exp = case.get("numero_expediente", "")
                    anio_car = case.get("anio_caratula", "")
                    nro_anio = ""
                    if num_exp and anio_car:
                        nro_anio = f"{num_exp}/{anio_car}"
                    elif num_exp:
                        nro_anio = num_exp
                    elif anio_car:
                        nro_anio = f"/{anio_car}"

                    # Insertar en el TreeView con las 3 columnas
                    try:
                        self.app_controller.case_tree.insert(
                            "", tk.END, values=(case_id, nro_anio, caratula), iid=str(case_id)
                        )
                        cases_loaded += 1
                    except Exception as tree_error:
                        print(f"Error al insertar caso {case_id} en TreeView: {tree_error}")
                        continue

                except Exception as case_error:
                    print(f"Error procesando caso {case.get('id', 'desconocido')}: {case_error}")
                    continue

            print(f"Casos cargados exitosamente: {cases_loaded} de {len(cases)}")

            # Limpiar selección y detalles
            self.app_controller.selected_case = None
            self.disable_case_buttons()
            self.app_controller.update_add_audiencia_button_state()

            return cases_loaded > 0

        except Exception as e:
            print(f"Error general al cargar casos para cliente {client_id}: {e}")
            import traceback
            traceback.print_exc()
            return False
        
    def clear_case_list(self):
        """Limpia la lista de casos"""
        for i in self.app_controller.case_tree.get_children():
            self.app_controller.case_tree.delete(i)
        self.app_controller.selected_case = None
        self.disable_case_buttons()

    def on_case_select(self, event):
        """Maneja la selección de un caso en el TreeView"""
        selected_items = self.app_controller.case_tree.selection()
        if selected_items:
            try:
                case_id = int(selected_items[0])
                case_data = self.db.get_case_by_id(case_id)
                if case_data:
                    self.app_controller.selected_case = case_data
                    self.enable_case_buttons()
                else:
                    print(f"Error: No se encontró caso con ID {case_id}")
                    self.app_controller.selected_case = None
                    self.disable_case_buttons()
            except (ValueError, IndexError) as e:
                print(f"Error seleccionando caso: {e}")
                self.app_controller.selected_case = None
                self.disable_case_buttons()
        else:
            self.app_controller.selected_case = None
            self.disable_case_buttons()

        self.app_controller.update_add_audiencia_button_state()

    def enable_case_buttons(self):
        """Habilita los botones de caso"""
        self.app_controller.edit_case_btn.config(state=tk.NORMAL)
        self.app_controller.delete_case_btn.config(state=tk.NORMAL)
        self.app_controller.add_case_btn.config(state=tk.NORMAL)

    def disable_case_buttons(self):
        """Deshabilita los botones de caso"""
        self.app_controller.edit_case_btn.config(state=tk.DISABLED)
        self.app_controller.delete_case_btn.config(state=tk.DISABLED)
        self.app_controller.add_case_btn.config(state=tk.NORMAL)  # Add button should stay enabled

    def abrir_dialogo_escrito_generico(self):
        """Abre el diálogo para confeccionar un escrito genérico"""
        if not self.app_controller.selected_case:
            messagebox.showwarning(
                "Advertencia",
                "No hay un caso seleccionado para generar el escrito.",
                parent=self.app_controller.root
            )
            return
        
        try:
            # Check if python-docx is available
            if not DOCX_AVAILABLE:
                messagebox.showerror(
                    "Error Librería",
                    "Falta 'python-docx'. Instálala con: pip install python-docx",
                    parent=self.app_controller.root
                )
                return
            
            # Open dialog to get user input
            dialog = EscritoGenericoDialog(self.app_controller.root)
            user_input = dialog.get_user_input()
            
            if user_input:
                # Generate document with user input
                self.generar_escrito_generico(
                    self.app_controller.selected_case["id"],
                    user_input["titulo"],
                    user_input["cuerpo"],
                    user_input.get("header_type", "PJN")
                )
                
        except Exception as e:
            messagebox.showerror(
                "Error",
                f"Error inesperado al abrir el diálogo: {e}",
                parent=self.app_controller.root
            )
            import traceback
            traceback.print_exc()

    def generar_escrito_generico(self, caso_id, titulo_escrito, cuerpo_escrito, header_type="PJN"):
        """Genera el documento Word con los datos del caso"""
        try:
            # Collect all necessary case data
            contexto = self._recopilar_datos_caso(caso_id)
            if not contexto:
                return
            
            # Add user input to context
            contexto['titulo_escrito'] = titulo_escrito
            contexto['cuerpo_escrito'] = cuerpo_escrito
            contexto['header_type'] = header_type
            
            # Create Word document
            doc = self._crear_documento_word(contexto)
            if not doc:
                return
            
            # Save document with intelligent folder logic
            caso_data = self.db.get_case_by_id(caso_id)
            if not caso_data:
                messagebox.showerror(
                    "Error",
                    "No se pudieron obtener los datos del caso para guardar el documento.",
                    parent=self.app_controller.root
                )
                return

            # Use intelligent saving logic
            ruta_guardado = self._save_document_intelligent(doc, caso_data, titulo_escrito)

            if ruta_guardado:
                # Log the document generation activity
                self._registrar_actividad_generacion(caso_id, titulo_escrito, ruta_guardado)
            
        except Exception as e:
            messagebox.showerror(
                "Error",
                f"Error al generar el documento: {e}",
                parent=self.app_controller.root
            )
            import traceback
            traceback.print_exc()

    def _recopilar_datos_caso(self, caso_id):
        """Recopila todos los datos necesarios del caso, cliente y usuario"""
        try:
            # Get case data
            caso_data = self.db.get_case_by_id(caso_id)
            if not caso_data:
                messagebox.showerror(
                    "Error",
                    "No se pudieron obtener los datos del caso.",
                    parent=self.app_controller.root
                )
                return None
            
            # Get user/lawyer data
            usuario_data = self.db.get_datos_usuario()
            if not usuario_data:
                messagebox.showwarning(
                    "Advertencia",
                    "No se encontraron datos del usuario. Se usarán valores por defecto.",
                    parent=self.app_controller.root
                )
                usuario_data = {}
            
            # Get client data
            cliente_data = None
            if caso_data.get("cliente_id"):
                cliente_data = self.db.get_client_by_id(caso_data["cliente_id"])
            
            if not cliente_data:
                messagebox.showwarning(
                    "Advertencia",
                    "No se encontraron datos del cliente. Se usarán valores por defecto.",
                    parent=self.app_controller.root
                )
                cliente_data = {}
            
            # Build context dictionary
            contexto = {
                # Case data
                'juzgado': caso_data.get('juzgado', '[JUZGADO PENDIENTE]'),
                'numero_expediente': caso_data.get('numero_expediente', '[NÚMERO]'),
                'anio_caratula': caso_data.get('anio_caratula', '[AÑO]'),
                'caratula': caso_data.get('caratula', '[CARÁTULA PENDIENTE]'),
                
                # Client data
                'cliente_nombre': cliente_data.get('nombre', '[CLIENTE PENDIENTE]'),
                
                # Lawyer data
                'abogado_nombre': usuario_data.get('nombre_abogado', '[ABOGADO PENDIENTE]'),
                'tomo': usuario_data.get('tomo', '[TOMO]'),
                'folio': usuario_data.get('folio', '[FOLIO]'),
                'colegio': usuario_data.get('colegio', '[COLEGIO]'),
            }
            
            # Format case number if both parts are available
            if contexto['numero_expediente'] != '[NÚMERO]' and contexto['anio_caratula'] != '[AÑO]':
                contexto['numero_expediente_completo'] = f"{contexto['numero_expediente']}/{contexto['anio_caratula']}"
            else:
                contexto['numero_expediente_completo'] = f"{contexto['numero_expediente']}/{contexto['anio_caratula']}"
            
            return contexto
            
        except Exception as e:
            messagebox.showerror(
                "Error",
                f"Error al recopilar datos del caso: {e}",
                parent=self.app_controller.root
            )
            import traceback
            traceback.print_exc()
            return None

    def _crear_documento_word(self, contexto):
        """Crea el documento Word usando python-docx"""
        try:
            # Create new document
            doc = Document()
            
            # Set document margins (1 inch on all sides)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add court information (left-aligned)
            court_para = doc.add_paragraph()
            court_para.add_run(contexto['juzgado']).bold = True
            court_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add case number (left-aligned)
            case_num_para = doc.add_paragraph()
            case_num_para.add_run(f"Expte. Nº {contexto['numero_expediente_completo']}")
            case_num_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add some space
            doc.add_paragraph()
            
            # Add document title (centered, uppercase)
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(contexto['titulo_escrito'].upper())
            title_run.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add formal greeting
            greeting_para = doc.add_paragraph()
            greeting_para.add_run("Señor Juez,")
            greeting_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add presentation paragraph based on header type
            header_type = contexto.get('header_type', 'PJN')
            
            if header_type == 'SCBA':
                # SCBA format presentation
                presentation_text = (
                    f"{contexto['cliente_nombre']}, por derecho propio, patrocinado por el Dr. DARIO JAVIER RAMIREZ, abogado, inscripto en el Tomo XV Folio 421 del CASM, manteniendo domicilio constituido en los Autos caratulados: \"{contexto['caratula']}\", a V.S. digo:"
                )
            else:
                # PJN format presentation (default)
                presentation_text = (
                    f"{contexto['cliente_nombre']}, por derecho propio, patrocinado por el Dr. DARIO JAVIER RAMIREZ, abogado, inscripto en el Tomo 91 Folio 49 del CPACF, manteniendo domicilio constituido en los Autos caratulados: \"{contexto['caratula']}\", a V.S. digo:"
                )
            
            presentation_para = doc.add_paragraph()
            presentation_para.add_run(presentation_text)
            presentation_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add some space
            doc.add_paragraph()
            
            # Add section marker
            section_para = doc.add_paragraph()
            section_run = section_para.add_run("I.-")
            section_run.bold = True
            section_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add user-provided body content
            body_para = doc.add_paragraph()
            body_para.add_run(contexto['cuerpo_escrito'])
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Add some space
            
            doc.add_paragraph()
            
            # Add standard petition
            petition_para = doc.add_paragraph()
            petition_para.add_run("Proveer de conformidad, que")
            petition_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add "SERA JUSTICIA.-" (centered)
            justice_para = doc.add_paragraph()
            justice_run = justice_para.add_run("SERA JUSTICIA.-")
            justice_run.bold = True
            justice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return doc
            
        except Exception as e:
            messagebox.showerror(
                "Error",
                f"Error al crear el documento Word: {e}",
                parent=self.app_controller.root
            )
            import traceback
            traceback.print_exc()
            return None

    def _registrar_actividad_generacion(self, caso_id, titulo_escrito, ruta_archivo):
        """Registra la generación del documento como una actividad del caso"""
        try:
            fecha_hora = datetime.datetime.now()
            tipo_actividad = "Generación de Documento"
            descripcion = f"Confeccionando escrito genérico: '{titulo_escrito}'"
            
            # Get current user (if available)
            usuario_actual = getattr(self.app_controller, 'current_user', 'Sistema')
            
            # Add activity to database
            actividad_id = self.db.add_actividad_caso(
                caso_id=caso_id,
                fecha_hora=fecha_hora,
                tipo_actividad=tipo_actividad,
                descripcion=descripcion,
                creado_por=usuario_actual,
                referencia_documento=ruta_archivo
            )
            
            if actividad_id:
                # Refresh the case window if it's open
                self.app_controller._refresh_open_case_window(caso_id)
            else:
                print("Advertencia: No se pudo registrar la actividad de generación de documento")
                
        except Exception as e:
            print(f"Error al registrar actividad de generación: {e}")
            import traceback
            traceback.print_exc()
            # Don't show error to user as this is not critical for document generation

# Funciones de compatibilidad para mantener la interfaz existente
'''def open_case_dialog(app_controller, case_id=None):
    """Función de compatibilidad - usa la nueva clase CaseManager"""
    if not hasattr(app_controller, "case_manager"):
        app_controller.case_manager = CaseManager(app_controller)
    app_controller.case_manager.open_case_dialog(case_id)


def save_case(
    app_controller,
    case_id,
    cliente_id,
    caratula,
    num_exp,
    anio_car,
    juzgado,
    juris,
    notas,
    ruta,
    inact_days,
    inact_enabled,
    etiquetas_caso_str,
    dialog,
):
    """Función de compatibilidad - usa la nueva clase CaseManager"""
    if not hasattr(app_controller, "case_manager"):
        app_controller.case_manager = CaseManager(app_controller)
    app_controller.case_manager.save_case(
        case_id,
        cliente_id,
        caratula,
        num_exp,
        anio_car,
        juzgado,
        juris,
        notas,
        ruta,
        inact_days,
        inact_enabled,
        etiquetas_caso_str,
        dialog,
    )
    '''