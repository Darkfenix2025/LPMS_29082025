from docx import Document
from docxtpl import DocxTemplate
import re
import os

def update_placeholders_in_docx(file_path, replacements):
    """
    Update placeholders in a Word document with correct database field names.

    Args:
        file_path (str): Path to the docx file
        replacements (dict): Dictionary of old_placeholder -> new_placeholder

    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Load the document
        doc = Document(file_path)
        replacements_made = 0

        # Process all paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                original_text = run.text
                updated_text = original_text

                # Apply all replacements
                for old_placeholder, new_placeholder in replacements.items():
                    # More flexible pattern that doesn't require word boundaries
                    pattern = re.escape(old_placeholder)
                    if re.search(pattern, updated_text, flags=re.IGNORECASE):
                        updated_text = re.sub(pattern, new_placeholder, updated_text, flags=re.IGNORECASE)
                        replacements_made += 1
                        print(f"Replaced: {old_placeholder} -> {new_placeholder}")

                # Update the text if changed
                if updated_text != original_text:
                    run.text = updated_text

        # Process tables as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            original_text = run.text
                            updated_text = original_text

                            # Apply all replacements
                            for old_placeholder, new_placeholder in replacements.items():
                                pattern = re.escape(old_placeholder)
                                if re.search(pattern, updated_text, flags=re.IGNORECASE):
                                    updated_text = re.sub(pattern, new_placeholder, updated_text, flags=re.IGNORECASE)
                                    replacements_made += 1
                                    print(f"Replaced in table: {old_placeholder} -> {new_placeholder}")

                            if updated_text != original_text:
                                run.text = updated_text

        # Save the document
        doc.save(file_path)
        print(f"Document updated successfully: {file_path}")
        print(f"Total replacements made: {replacements_made}")
        return True

    except Exception as e:
        print(f"Error updating document: {e}")
        return False

def analyze_template_variables(file_path):
    """
    Analyze the template variables in a DOCX file using docxtpl.

    Args:
        file_path (str): Path to the docx file

    Returns:
        list: List of template variables found
    """
    try:
        doc = DocxTemplate(file_path)
        variables = doc.get_undeclared_template_variables()
        print(f"Template variables found: {variables}")
        return variables
    except Exception as e:
        print(f"Error analyzing template: {e}")
        return []

def main():
    file_path = 'plantillas/mediacion/acuerdo_base.docx'

    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return False

    print(f"Analyzing template: {file_path}")

    # First, analyze what variables are in the template
    template_vars = analyze_template_variables(file_path)
    print(f"Found {len(template_vars)} template variables: {sorted(template_vars)}")

    # Define replacements to match context field names from case_dialog_manager.py
    replacements = {
        # Fix remaining issues from previous replacements
        'AÑO_ACUERDO_acuerdo': 'AÑO_ACUERDO',
        'MONTO_COMPENSACION_NUMEROS_letras': 'MONTO_COMPENSACION_LETRAS',
        'PLAZO_PAGO_DIAS_letras': 'PLAZO_PAGO_LETRAS',
        'PLAZO_PAGO_DIAS_dias': 'PLAZO_PAGO_DIAS',
        'PLAZO_PAGO_DIAS_LETRAS': 'PLAZO_PAGO_LETRAS',
        'PLAZO_PAGO_DIAS_DIAS': 'PLAZO_PAGO_DIAS',
        'monto_honorarios_letrado__letras': 'MONTO_HONORARIOS_LETRAS',
        'monto_honorarios_mediador__letras': 'MONTO_HONORARIOS_MEDIADOR_LETRAS',
        # Fix the remaining unmapped placeholders
        'actor.nombre_completo': 'ACTORES[0].nombre_completo',
        'actor.domicilio_real': 'ACTORES[0].domicilio_real',
        'rep.nombre_completo': 'ACTORES[0].representantes[0].nombre_completo',
        'rep.personeria': 'ACTORES[0].representantes[0].personeria',
        'demandado.nombre_completo': 'DEMANDADOS[0].nombre_completo',
        'demandado.domicilio_real': 'DEMANDADOS[0].domicilio_real',
        'abogado_demandado.nombre_completo': 'DEMANDADOS[0].representantes[0].nombre_completo',
        'abogado_demandado.domicilio_real': 'DEMANDADOS[0].representantes[0].domicilio_real',
        'act.nombre_completo': 'ACTORES[0].nombre_completo',
        'dem.nombre_completo': 'DEMANDADOS[0].nombre_completo',
        'dem.cuit': 'DEMANDADOS[0].cuit',
        'act.cuit': 'ACTORES[0].cuit',
        # Remove empty placeholders that cause errors
        '{{}}': '',
        # Update any remaining lowercase placeholders
        'numero_expediente': 'NUMERO_EXPEDIENTE',
        'anio': 'AÑO_ACUERDO',
        'caratula': 'CARATULA',
        'fecha_acuerdo': 'DIA_ACUERDO',
        'mes_acuerdo': 'MES_ACUERDO',
        'anio_acuerdo': 'AÑO_ACUERDO',
        'monto_acuerdo': 'MONTO_COMPENSACION_NUMEROS',
        'monto_acuerdo_letras': 'MONTO_COMPENSACION_LETRAS',
        'plazo_pago': 'PLAZO_PAGO_DIAS',
        'datos_bancarios': 'BANCO_ACTOR CBU_ACTOR ALIAS_ACTOR CUIT_ACTOR',
        'banco_actor': 'BANCO_ACTOR',
        'cbu_actor': 'CBU_ACTOR',
        'alias_actor': 'ALIAS_ACTOR',
        'cuit_actor': 'CUIT_ACTOR',
        'banco_letrado_actor': 'BANCO_LETRADO_ACTOR',
        'cbu_letrado_actor': 'CBU_LETRADO_ACTOR',
        'alias_letrado_actor': 'ALIAS_LETRADO_ACTOR',
        'rep.cuit': 'ACTORES[0].representantes[0].cuit',
        'monto_honorarios_letrado_numeros': 'MONTO_HONORARIOS_NUMEROS',
        'plazo_pago_letras': 'PLAZO_PAGO_LETRAS',
        'plazo_pago_dias': 'PLAZO_PAGO_DIAS',
        'monto_honorarios_mediador_numeros': 'MONTO_HONORARIOS_MEDIADOR_NUMEROS',
        'detalles_caso': 'DETALLES_CASO',
    }

    print(f"Number of replacements to apply: {len(replacements)}")

    success = update_placeholders_in_docx(file_path, replacements)

    if success:
        print("Placeholder update completed successfully!")
        print("\nReplacements applied:")
        for old, new in replacements.items():
            print(f"  {old} -> {new}")
    else:
        print("Failed to update placeholders.")

    return success

if __name__ == "__main__":
    main()