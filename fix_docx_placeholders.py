#!/usr/bin/env python3
"""
Script to fix placeholders in acuerdo_base.docx file.
This script will update the corrupted docx file with proper placeholders
that correspond to database fields.
"""

import os
import sys
from pathlib import Path
from typing import Dict, List

# Try to import docx, but handle if not available
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")

def create_new_template_document(output_path: str) -> bool:
    """
    Create a new Word document with proper placeholders based on the template.

    Args:
        output_path: Path where to save the new document

    Returns:
        bool: True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        print("Error: python-docx is required for this operation")
        return False

    try:
        # Create a new document
        doc = Document()

        # Set up document formatting
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add title
        title = doc.add_heading('ACTA DE ACUERDO CONCILIATORIO', 0)
        title.alignment = 1  # Center alignment

        # Add case information
        doc.add_paragraph(f"EXPEDIENTE N°: {{numero_expediente}}")
        doc.add_paragraph(f"CARÁTULA: \"{{caratula}}\"")
        doc.add_paragraph()

        # Add date paragraph
        date_para = doc.add_paragraph()
        date_para.add_run("En la Ciudad Autónoma de Buenos Aires, a los ")
        date_para.add_run("{fecha_acuerdo}").bold = True
        date_para.add_run(" días del mes de ")
        date_para.add_run("{mes_acuerdo}").bold = True
        date_para.add_run(" de ")
        date_para.add_run("{anio_acuerdo}").bold = True
        date_para.add_run(", en el marco de la audiencia de conciliación celebrada de manera VIRTUAL correspondiente al expediente de la referencia, ante mí, Dr. DARIO JAVIER RAMIREZ, Conciliador de Consumo Matrícula SCRC N° 470, comparecen:")
        doc.add_paragraph()

        # Add parties section
        doc.add_heading('PARTES INTERVINIENTES', level=1)

        # Actor section
        actor_heading = doc.add_paragraph()
        actor_heading.add_run("1. La PARTE REQUIRENTE:").bold = True

        actor_info = doc.add_paragraph()
        actor_info.add_run("Sr. ")
        actor_info.add_run("{actor.nombre_completo}").bold = True
        actor_info.add_run(", C.U.I.T. N° ")
        actor_info.add_run("{actor.cuit}").bold = True
        actor_info.add_run(", con domicilio constituido en ")
        actor_info.add_run("{actor.domicilio_real}").bold = True
        actor_info.add_run(", Ciudad Autónoma de Buenos Aires, en adelante \"el Requirente\".")
        doc.add_paragraph()

        # Defendant section
        defendant_heading = doc.add_paragraph()
        defendant_heading.add_run("2. La PARTE REQUERIDA:").bold = True

        defendant_info = doc.add_paragraph()
        defendant_info.add_run("{demandado.nombre_completo}").bold = True
        defendant_info.add_run(", C.U.I.T. N° ")
        defendant_info.add_run("{demandado.cuit}").bold = True
        defendant_info.add_run(", con domicilio en ")
        defendant_info.add_run("{demandado.domicilio_real}").bold = True
        defendant_info.add_run(", Ciudad Autónoma de Buenos Aires, representada en este acto por su Apoderado, ")
        defendant_info.add_run("{abogado_demandado.nombre_completo}").bold = True
        defendant_info.add_run(", con domicilio constituido en ")
        defendant_info.add_run("{abogado_demandado.domicilio_real}").bold = True
        defendant_info.add_run(", Ciudad Autónoma de Buenos Aires, y que en adelante será nombrada en el presente acuerdo como \"la Requerida\".")
        doc.add_paragraph()

        # Agreement clauses
        doc.add_heading('CLÁUSULAS DEL ACUERDO', level=1)

        # Clause 1
        clause1 = doc.add_paragraph()
        clause1.add_run("CLÁUSULA PRIMERA: OBJETO Y ANTECEDENTES.").bold = True
        doc.add_paragraph("El presente acuerdo tiene por objeto la total y definitiva extinción de las obligaciones emergentes del reclamo que diera origen al expediente de la referencia.")
        doc.add_paragraph()

        # Clause 2
        clause2 = doc.add_paragraph()
        clause2.add_run("CLÁUSULA SEGUNDA: ACUERDO TRANSACCIONAL.").bold = True

        agreement_para = doc.add_paragraph()
        agreement_para.add_run("LA REQUERIDA, ")
        agreement_para.add_run("{demandado.nombre_completo}").bold = True
        agreement_para.add_run(", sin reconocer hechos ni derecho alguno y al solo efecto conciliatorio, ofrece al Requirente, y éste acepta, las siguientes obligaciones:")
        doc.add_paragraph()

        # Financial obligation
        financial_para = doc.add_paragraph()
        financial_para.add_run("OBLIGACIÓN DE DAR SUMA DE DINERO: La Requerida se compromete a abonar al Requirente, a título de gesto comercial, la suma única y total de PESOS ")
        financial_para.add_run("{monto_acuerdo}").bold = True
        financial_para.add_run(" ($ ")
        financial_para.add_run("{monto_acuerdo}").bold = True
        financial_para.add_run("). Dicho importe se hará efectivo mediante transferencia bancaria en un plazo máximo de ")
        financial_para.add_run("{plazo_pago}").bold = True
        financial_para.add_run(" días hábiles a contar desde la firma del presente. La transferencia se realizará a la cuenta de titularidad del Requirente, cuyos datos son:")
        doc.add_paragraph("{datos_bancarios}")
        doc.add_paragraph()

        # Clause 3
        clause3 = doc.add_paragraph()
        clause3.add_run("CLÁUSULA TERCERA: ACEPTACIÓN Y RENUNCIA DEL REQUIRENTE.").bold = True

        acceptance_para = doc.add_paragraph()
        acceptance_para.add_run("EL REQUIRENTE manifiesta su plena conformidad y acepta el ofrecimiento formulado en la cláusula anterior. En consecuencia, declara que una vez cumplidas en su totalidad las obligaciones asumidas, nada más tendrá que reclamar a ")
        acceptance_para.add_run("{demandado.nombre_completo}").bold = True
        acceptance_para.add_run(" por el asunto que dio origen a este reclamo.")
        doc.add_paragraph()

        # Add remaining standard clauses
        doc.add_paragraph("CLÁUSULA CUARTA: COSTAS Y HONORARIOS DEL CONCILIADOR.").bold = True
        doc.add_paragraph("Las costas del presente procedimiento, incluyendo los honorarios del Conciliador actuante, serán a exclusivo cargo de LA REQUERIDA.")
        doc.add_paragraph()

        doc.add_paragraph("CLÁUSULA QUINTA: EFECTOS Y HOMOLOGACIÓN.").bold = True
        doc.add_paragraph("Una vez cumplidas en su totalidad las obligaciones asumidas, las partes nada más tendrán que reclamarse. El presente acuerdo tendrá para las partes el efecto de cosa juzgada una vez homologado por la autoridad de aplicación.")
        doc.add_paragraph()

        doc.add_paragraph("CLÁUSULA SEXTA: INCUMPLIMIENTO.").bold = True
        doc.add_paragraph("La mora en el cumplimiento de cualquiera de las obligaciones asumidas se producirá de pleno derecho por el mero vencimiento de los plazos estipulados.")
        doc.add_paragraph()

        doc.add_paragraph("CLÁUSULA SÉPTIMA: JURISDICCIÓN.").bold = True
        doc.add_paragraph("Para el caso de incumplimiento, las partes se someten a la jurisdicción de los Tribunales Ordinarios.")
        doc.add_paragraph()

        # Signatures section
        doc.add_heading('FIRMAS', level=1)

        # Create signature table
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'EL REQUIRENTE'
        hdr_cells[1].text = 'LA REQUERIDA'
        hdr_cells[2].text = 'EL CONCILIADOR'

        # Signature row
        sig_cells = table.rows[1].cells
        sig_cells[0].text = '_________________________'
        sig_cells[1].text = '_________________________'
        sig_cells[2].text = '_________________________'

        # Save the document
        doc.save(output_path)
        print(f"New template document created successfully: {output_path}")
        return True

    except Exception as e:
        print(f"Error creating new template document: {e}")
        return False

def add_missing_database_functions():
    """
    Add the missing database functions that are needed for proper placeholder mapping.
    This creates a supplementary script with the required functions.
    """
    functions_code = '''
# Missing database functions for agreement generation
# Add these functions to crm_database.py

def get_partes_by_caso_and_tipo(case_id: int, tipo: str) -> List[Dict]:
    """
    Get parties by case ID and role type.
    This function is called by ai_agreement_generator.py but doesn't exist.

    Args:
        case_id: Case ID
        tipo: Role type ('ACTOR', 'DEMANDADO', 'ABOGADO', etc.)

    Returns:
        List of parties with the specified role type
    """
    try:
        query = """
            SELECT
                r.id as rol_id,
                r.caso_id,
                r.contacto_id,
                r.rol_principal,
                r.rol_secundario,
                r.representa_a_id,
                r.datos_bancarios,
                c.nombre_completo,
                c.dni,
                c.cuit,
                c.domicilio_real,
                c.domicilio_legal,
                c.email,
                c.telefono
            FROM roles_en_caso r
            JOIN contactos c ON r.contacto_id = c.id
            WHERE r.caso_id = %s AND r.rol_principal = %s
            ORDER BY c.nombre_completo
        """

        parties = execute_query(query, (case_id, tipo.upper()), fetch_all=True)
        return parties if parties else []

    except Exception as e:
        print(f"Error getting parties by case and type: {e}")
        return []

def get_representantes_by_party(case_id: int, party_id: int) -> List[Dict]:
    """
    Get representatives (lawyers) for a specific party.

    Args:
        case_id: Case ID
        party_id: Party ID (contacto_id)

    Returns:
        List of representatives for the party
    """
    try:
        query = """
            SELECT
                r.id as rol_id,
                r.contacto_id,
                r.rol_principal,
                r.rol_secundario,
                c.nombre_completo,
                c.dni,
                c.cuit,
                c.domicilio_real,
                c.domicilio_legal,
                c.email,
                c.telefono
            FROM roles_en_caso r
            JOIN contactos c ON r.contacto_id = c.id
            WHERE r.caso_id = %s
              AND (r.representa_a_id = %s OR r.contacto_id = %s)
              AND r.rol_principal IN ('ABOGADO', 'APODERADO')
            ORDER BY r.rol_principal, c.nombre_completo
        """

        representatives = execute_query(query, (case_id, party_id, party_id), fetch_all=True)
        return representatives if representatives else []

    except Exception as e:
        print(f"Error getting representatives for party: {e}")
        return []

def get_abogados_by_caso_and_tipo(case_id: int, tipo_party: str) -> List[Dict]:
    """
    Get lawyers by case and party type.

    Args:
        case_id: Case ID
        tipo_party: Party type ('ACTOR', 'DEMANDADO')

    Returns:
        List of lawyers representing the specified party type
    """
    try:
        # First get the parties of the specified type
        parties = get_partes_by_caso_and_tipo(case_id, tipo_party)

        if not parties:
            return []

        lawyers = []
        for party in parties:
            party_lawyers = get_representantes_by_party(case_id, party['contacto_id'])
            lawyers.extend(party_lawyers)

        return lawyers

    except Exception as e:
        print(f"Error getting lawyers by case and party type: {e}")
        return []
'''

    try:
        with open('missing_db_functions.py', 'w', encoding='utf-8') as f:
            f.write(functions_code)
        print("Missing database functions saved to: missing_db_functions.py")
        return True
    except Exception as e:
        print(f"Error saving missing functions: {e}")
        return False

def main():
    """Main function to fix the docx placeholders."""
    print("=== DOCX PLACEHOLDER FIXER ===")
    print()

    # Check if original file exists
    original_path = "plantillas/mediacion/acuerdo_base.docx"
    if os.path.exists(original_path):
        print(f"Original file found: {original_path}")
        print("Note: The original file appears to be corrupted based on previous attempts.")
        print("This script will create a new template document with proper placeholders.")
    else:
        print(f"Original file not found: {original_path}")
        print("This script will create a new template document.")

    print()

    # Create new template document
    new_template_path = "plantillas/mediacion/acuerdo_base_fixed.docx"

    # Ensure directory exists
    os.makedirs(os.path.dirname(new_template_path), exist_ok=True)

    print("Creating new template document with proper placeholders...")
    if create_new_template_document(new_template_path):
        print("[OK] New template document created successfully!")
        print(f"  Location: {new_template_path}")
    else:
        print("[ERROR] Failed to create new template document")
        return False

    print()

    # Create missing database functions
    print("Creating missing database functions...")
    if add_missing_database_functions():
        print("[OK] Missing database functions created!")
        print("  Location: missing_db_functions.py")
        print("  Note: These functions should be added to crm_database.py")
    else:
        print("[ERROR] Failed to create missing database functions")

    print()
    print("=== SUMMARY ===")
    print("[OK] Created new template document with proper placeholders")
    print("[OK] Identified missing database functions")
    print()
    print("Next steps:")
    print("1. Review the new template document")
    print("2. Add the missing functions to crm_database.py")
    print("3. Test the agreement generation with the new template")
    print()
    print("The new template includes placeholders for:")
    print("- Case information (numero_expediente, caratula)")
    print("- Party information (actor, demandado, abogado)")
    print("- Agreement details (monto, plazo, datos_bancarios)")
    print("- Generated values (dates)")

    return True

if __name__ == "__main__":
    main()