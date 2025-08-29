#!/usr/bin/env python3
"""
Script to create a basic Word template for mediation agreements.
"""

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    def create_mediation_template():
        """Create a basic mediation agreement template."""
        
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('ACUERDO DE MEDIACIÓN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add case information section
        doc.add_heading('INFORMACIÓN DEL CASO', level=1)
        doc.add_paragraph('Carátula: {{caratula}}')
        doc.add_paragraph('Expediente: {{numero_expediente}}')
        doc.add_paragraph('Juzgado: {{juzgado}}')
        doc.add_paragraph('Fecha: {{fecha_actual}}')
        
        # Add parties section
        doc.add_heading('PARTES', level=1)
        
        doc.add_heading('ACTORES:', level=2)
        doc.add_paragraph('{{actores_texto}}')
        
        doc.add_heading('DEMANDADOS:', level=2)
        doc.add_paragraph('{{demandados_texto}}')
        
        # Add agreement section
        doc.add_heading('ACUERDO', level=1)
        
        doc.add_paragraph(
            'Las partes acuerdan que el demandado pagará al actor la suma de '
            'PESOS {{monto_compensacion_letras}} ($ {{monto_compensacion_numeros}}) '
            'en un plazo de {{plazo_pago_letras}} ({{plazo_pago_dias}}) días corridos '
            'a partir de la fecha de homologación del presente acuerdo.'
        )
        
        # Add payment details section
        doc.add_heading('DATOS PARA EL PAGO', level=1)
        doc.add_paragraph('Banco: {{banco_actor}}')
        doc.add_paragraph('CBU: {{cbu_actor}}')
        doc.add_paragraph('Alias: {{alias_actor}}')
        doc.add_paragraph('CUIT/CUIL: {{cuit_actor}}')
        
        # Add signatures section
        doc.add_heading('FIRMAS', level=1)
        doc.add_paragraph('\n\n')
        doc.add_paragraph('_' * 30 + '    ' + '_' * 30)
        doc.add_paragraph('Firma Actor(es)' + ' ' * 20 + 'Firma Demandado(s)')
        
        doc.add_paragraph('\n\n')
        doc.add_paragraph('_' * 60)
        doc.add_paragraph('Mediador')
        
        # Save the template
        template_path = 'plantillas/mediacion/acuerdo_base.docx'
        doc.save(template_path)
        print(f"✓ Template created successfully at: {template_path}")
        
        return template_path
    
    if __name__ == "__main__":
        create_mediation_template()
        
except ImportError as e:
    print(f"❌ Error: python-docx library not available: {e}")
    print("Install it with: pip install python-docx")
except Exception as e:
    print(f"❌ Error creating template: {e}")