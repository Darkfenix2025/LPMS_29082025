from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_new_mediation_template():
    """
    Create a new, valid mediation agreement template to replace the corrupted one.
    """
    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading('ACTA DE ACUERDO CONCILIATORIO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Case information section
    doc.add_paragraph('')
    case_info = doc.add_paragraph()
    case_info.add_run('EXPEDIENTE N°: ').bold = True
    case_info.add_run('{numero_expediente}')
    case_info.add_run('\nCARÁTULA: ').bold = True
    case_info.add_run('{caratula}')
    case_info.add_run('\nEN LA CIUDAD AUTÓNOMA DE BUENOS AIRES, a los {fecha_acuerdo} días del mes de {mes_acuerdo} de {anio_acuerdo}')

    # Parties section
    doc.add_paragraph('')
    parties_title = doc.add_paragraph('PARTES INTERVINIENTES')
    parties_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parties_title.runs[0].bold = True

    # Actor section
    actor_para = doc.add_paragraph()
    actor_para.add_run('1. LA PARTE ACTORA:\n').bold = True
    actor_para.add_run('Sr./Sra. {actor.nombre_completo}, {actor.tipo_documento} N° {actor.dni}, con domicilio en {actor.domicilio_real}, Ciudad Autónoma de Buenos Aires.')

    # Defendant section
    defendant_para = doc.add_paragraph()
    defendant_para.add_run('2. LA PARTE DEMANDADA:\n').bold = True
    defendant_para.add_run('{defendant.nombre_completo}, con domicilio en {defendant.domicilio_real}, representada en este acto por el Dr./Dra. {lawyer.nombre_completo}.')

    # Agreement section
    doc.add_paragraph('')
    agreement_title = doc.add_paragraph('ACUERDO TRANSACCIONAL')
    agreement_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    agreement_title.runs[0].bold = True

    agreement_para = doc.add_paragraph()
    agreement_para.add_run('Las partes, en ejercicio de su autonomía de la voluntad y con el fin de poner término al presente conflicto, manifiestan haber alcanzado el siguiente acuerdo:')
    agreement_para.add_run('\n\n1. COMPENSACIÓN ECONÓMICA: ').bold = True
    agreement_para.add_run('La parte demandada se compromete a pagar a la parte actora la suma de ${monto_compensacion_numeros} ({monto_compensacion_letras} PESOS).')
    agreement_para.add_run('\n\n2. PLAZO DE PAGO: ').bold = True
    agreement_para.add_run('El pago se efectuará dentro del plazo de {plazo_pago_dias} ({plazo_pago_letras}) días hábiles desde la firma del presente acuerdo.')
    agreement_para.add_run('\n\n3. FORMA DE PAGO: ').bold = True
    agreement_para.add_run('El pago se realizará mediante transferencia bancaria a la cuenta del actor: Banco {banco_actor}, CBU {cbu_actor}, Alias {alias_actor}, CUIT {cuit_actor}.')

    # Signatures section
    doc.add_paragraph('')
    signatures_title = doc.add_paragraph('FIRMAS')
    signatures_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signatures_title.runs[0].bold = True

    signatures_para = doc.add_paragraph()
    signatures_para.add_run('Conforme las partes intervinientes:')
    signatures_para.add_run('\n\n_______________________________                    _______________________________')
    signatures_para.add_run('\n      FIRMA ACTOR                                                        FIRMA DEMANDADO')
    signatures_para.add_run('\n\nEn Buenos Aires, a los {fecha_acuerdo} días del mes de {mes_acuerdo} de {anio_acuerdo}.')

    # Save the document
    output_path = 'plantillas/mediacion/acuerdo_base_nuevo.docx'
    doc.save(output_path)
    print(f"New template created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_new_mediation_template()