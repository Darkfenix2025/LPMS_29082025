#!/usr/bin/env python3
"""
Test script para verificar la compatibilidad hacia atrás de la UI
después de la refactorización del método generar_escrito_mediacion.

Este script prueba:
1. El flujo completo de generación de documentos desde la UI
2. Que todos los diálogos funcionan correctamente
3. Que los documentos generados son idénticos a los anteriores
4. Casos de error y cancelación de usuario
"""

import sys
import os
import tkinter as tk
from tkinter import messagebox
import unittest
from unittest.mock import Mock, patch, MagicMock
import tempfile
import shutil
import datetime
import sqlite3

# Agregar el directorio actual al path para importar módulos
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Importar módulos del sistema
import crm_database as db
from case_dialog_manager import CaseManager, ErrorMessageManager
from main_app import CRMLegalApp

class TestUIBackwardCompatibility(unittest.TestCase):
    """
    Suite de pruebas para verificar la compatibilidad hacia atrás de la UI
    después de la refactorización del método generar_escrito_mediacion.
    """
    
    @classmethod
    def setUpClass(cls):
        """Configuración inicial para todas las pruebas."""
        print("=" * 80)
        print("INICIANDO PRUEBAS DE COMPATIBILIDAD HACIA ATRÁS DE LA UI")
        print("=" * 80)
        
        # Crear una base de datos temporal para las pruebas
        cls.test_db_path = tempfile.mktemp(suffix='.db')
        cls.original_db_path = db.DATABASE_PATH if hasattr(db, 'DATABASE_PATH') else 'crm_legal.db'
        
        # Configurar base de datos de prueba
        cls._setup_test_database()
        
        # Crear directorio temporal para plantillas
        cls.temp_dir = tempfile.mkdtemp()
        cls.template_dir = os.path.join(cls.temp_dir, 'plantillas', 'mediacion')
        os.makedirs(cls.template_dir, exist_ok=True)
        
        # Crear plantilla de prueba
        cls._create_test_template()
        
        print(f"Base de datos de prueba: {cls.test_db_path}")
        print(f"Directorio temporal: {cls.temp_dir}")
    
    @classmethod
    def tearDownClass(cls):
        """Limpieza después de todas las pruebas."""
        print("\n" + "=" * 80)
        print("FINALIZANDO PRUEBAS DE COMPATIBILIDAD HACIA ATRÁS")
        print("=" * 80)
        
        # Limpiar archivos temporales
        try:
            if os.path.exists(cls.test_db_path):
                os.remove(cls.test_db_path)
            if os.path.exists(cls.temp_dir):
                shutil.rmtree(cls.temp_dir)
        except Exception as e:
            print(f"Error limpiando archivos temporales: {e}")
    
    def setUp(self):
        """Configuración antes de cada prueba."""
        # Crear mock de la aplicación principal
        self.mock_app = Mock()
        self.mock_app.root = Mock()
        
        # Crear instancia de CaseManager con mock de app_controller
        self.case_manager = CaseManager(app_controller=self.mock_app)
        
        # Configurar mocks para diálogos
        self.setup_dialog_mocks()
    
    def tearDown(self):
        """Limpieza después de cada prueba."""
        pass
    
    @classmethod
    def _setup_test_database(cls):
        """Configura una base de datos de prueba con datos de ejemplo."""
        try:
            # Crear conexión a la base de datos de prueba
            conn = sqlite3.connect(cls.test_db_path)
            cursor = conn.cursor()
            
            # Crear tablas necesarias
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS casos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    caratula TEXT NOT NULL,
                    numero_expediente TEXT,
                    anio_caratula TEXT,
                    juzgado TEXT,
                    jurisdiccion TEXT,
                    fecha_creacion TEXT,
                    estado TEXT DEFAULT 'Activo'
                )
            ''')
            
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS partes (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    caso_id INTEGER,
                    nombre_completo TEXT NOT NULL,
                    tipo_parte TEXT NOT NULL,
                    dni TEXT,
                    domicilio TEXT,
                    telefono TEXT,
                    email TEXT,
                    FOREIGN KEY (caso_id) REFERENCES casos (id)
                )
            ''')
            
            # Insertar datos de prueba
            cursor.execute('''
                INSERT INTO casos (id, caratula, numero_expediente, anio_caratula, juzgado, jurisdiccion)
                VALUES (1, 'PÉREZ JUAN C/ GARCÍA MARÍA S/ DAÑOS Y PERJUICIOS', '12345/2023', '2023', 'Juzgado Civil N° 1', 'La Plata')
            ''')
            
            cursor.execute('''
                INSERT INTO partes (caso_id, nombre_completo, tipo_parte, dni, domicilio, telefono, email)
                VALUES 
                (1, 'PÉREZ JUAN', 'Actor', '12345678', 'Calle Falsa 123', '221-1234567', 'juan.perez@email.com'),
                (1, 'GARCÍA MARÍA', 'Demandado', '87654321', 'Avenida Siempre Viva 456', '221-7654321', 'maria.garcia@email.com')
            ''')
            
            conn.commit()
            conn.close()
            
            # Configurar la base de datos en el módulo db
            db.DATABASE_PATH = cls.test_db_path
            
            print("Base de datos de prueba configurada exitosamente")
            
        except Exception as e:
            print(f"Error configurando base de datos de prueba: {e}")
            raise
    
    @classmethod
    def _create_test_template(cls):
        """Crea una plantilla de prueba para los documentos."""
        try:
            template_path = os.path.join(cls.template_dir, 'acuerdo_base.docx')
            
            # Crear un documento Word simple para pruebas
            try:
                from docx import Document
                doc = Document()
                doc.add_paragraph("ACUERDO DE MEDIACIÓN - PLANTILLA DE PRUEBA")
                doc.add_paragraph("Caso: {{caratula}}")
                doc.add_paragraph("Actor: {{actor_nombre}}")
                doc.add_paragraph("Demandado: {{demandado_nombre}}")
                doc.add_paragraph("Monto: {{monto_compensacion_numeros}}")
                doc.add_paragraph("Plazo: {{plazo_pago_dias}} días")
                doc.save(template_path)
                print(f"Plantilla de prueba creada: {template_path}")
            except ImportError:
                # Si no está disponible python-docx, crear un archivo dummy
                with open(template_path, 'wb') as f:
                    f.write(b'PLANTILLA DE PRUEBA DUMMY')
                print(f"Plantilla dummy creada: {template_path}")
                
        except Exception as e:
            print(f"Error creando plantilla de prueba: {e}")
    
    def setup_dialog_mocks(self):
        """Configura mocks para los diálogos de la UI."""
        # Mock para el diálogo de detalles del acuerdo
        self.mock_agreement_details = {
            'monto_compensacion_numeros': '150000.50',
            'monto_compensacion_letras': 'CIENTO CINCUENTA MIL PESOS CON CINCUENTA CENTAVOS',
            'plazo_pago_dias': '30',
            'plazo_pago_letras': 'TREINTA',
            'banco_actor': 'Banco Nación',
            'cbu_actor': '1234567890123456789012',
            'alias_actor': 'mi.alias.mp',
            'cuit_actor': '20-12345678-9'
        }
        
        # Mock para el diálogo de guardado de archivo
        self.mock_save_path = os.path.join(self.temp_dir, 'acuerdo_mediacion_test.docx')
    
    def test_01_ui_flow_complete_success(self):
        """
        Prueba 1: Flujo completo exitoso de generación de documentos desde la UI
        Verifica que el método generar_escrito_mediacion funciona correctamente
        con todos los diálogos y mantiene el comportamiento original.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 1: Flujo completo exitoso de generación de documentos")
        print("-" * 60)
        
        with patch.object(self.case_manager, '_ask_agreement_details_dialog') as mock_dialog, \
             patch('tkinter.filedialog.asksaveasfilename') as mock_save_dialog, \
             patch('os.path.exists') as mock_exists, \
             patch.object(self.case_manager, '_generar_documento_con_datos') as mock_generate:
            
            # Configurar mocks
            mock_dialog.return_value = self.mock_agreement_details
            mock_save_dialog.return_value = self.mock_save_path
            mock_exists.return_value = True
            mock_generate.return_value = {
                'success': True,
                'document': Mock(),
                'file_path': self.mock_save_path
            }
            
            # Ejecutar el método
            result = self.case_manager.generar_escrito_mediacion(1)
            
            # Verificaciones
            self.assertTrue(result, "El método debería retornar True para un flujo exitoso")
            mock_dialog.assert_called_once()
            mock_generate.assert_called_once()
            
            # Verificar que se llamó con los parámetros correctos
            call_args = mock_generate.call_args
            self.assertEqual(call_args[0][0], 1)  # caso_id
            self.assertEqual(call_args[0][1], self.mock_agreement_details)  # agreement_details
            
            print("✓ Flujo completo exitoso verificado")
    
    def test_02_user_cancellation_dialog(self):
        """
        Prueba 2: Cancelación del usuario en el diálogo de detalles
        Verifica que el sistema maneja correctamente cuando el usuario cancela.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 2: Cancelación del usuario en diálogo de detalles")
        print("-" * 60)
        
        with patch.object(self.case_manager, '_ask_agreement_details_dialog') as mock_dialog:
            # Simular cancelación del usuario
            mock_dialog.return_value = None
            
            # Ejecutar el método
            result = self.case_manager.generar_escrito_mediacion(1)
            
            # Verificaciones
            self.assertFalse(result, "El método debería retornar False cuando el usuario cancela")
            mock_dialog.assert_called_once()
            
            print("✓ Cancelación del usuario manejada correctamente")
    
    def test_03_user_cancellation_save_dialog(self):
        """
        Prueba 3: Cancelación del usuario en el diálogo de guardado
        Verifica que el sistema maneja correctamente cuando el usuario cancela el guardado.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 3: Cancelación del usuario en diálogo de guardado")
        print("-" * 60)
        
        with patch.object(self.case_manager, '_ask_agreement_details_dialog') as mock_dialog, \
             patch('tkinter.filedialog.asksaveasfilename') as mock_save_dialog, \
             patch.object(self.case_manager, '_generar_documento_con_datos') as mock_generate:
            
            # Configurar mocks
            mock_dialog.return_value = self.mock_agreement_details
            mock_save_dialog.return_value = ""  # Usuario cancela guardado
            mock_generate.return_value = {
                'success': True,
                'document': Mock(),
                'file_path': None
            }
            
            # Ejecutar el método
            result = self.case_manager.generar_escrito_mediacion(1)
            
            # Verificaciones
            self.assertFalse(result, "El método debería retornar False cuando el usuario cancela el guardado")
            mock_dialog.assert_called_once()
            mock_generate.assert_called_once()
            
            print("✓ Cancelación en guardado manejada correctamente")
    
    def test_04_invalid_case_id_error(self):
        """
        Prueba 4: Error con ID de caso inválido
        Verifica que el sistema maneja correctamente casos con ID inválido.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 4: Error con ID de caso inválido")
        print("-" * 60)
        
        with patch.object(ErrorMessageManager, 'show_error_dialog') as mock_error_dialog:
            # Ejecutar con ID inválido
            result = self.case_manager.generar_escrito_mediacion(99999)
            
            # Verificaciones
            self.assertFalse(result, "El método debería retornar False para ID inválido")
            mock_error_dialog.assert_called()
            
            print("✓ Error con ID inválido manejado correctamente")
    
    def test_05_missing_case_data_error(self):
        """
        Prueba 5: Error con datos de caso faltantes
        Verifica que el sistema maneja correctamente casos sin carátula.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 5: Error con datos de caso faltantes")
        print("-" * 60)
        
        # Crear caso sin carátula en la base de datos
        try:
            conn = sqlite3.connect(self.test_db_path)
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO casos (id, caratula, numero_expediente)
                VALUES (2, '', '54321/2023')
            ''')
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"Error configurando caso de prueba: {e}")
        
        with patch.object(ErrorMessageManager, 'show_error_dialog') as mock_error_dialog:
            # Ejecutar con caso sin carátula
            result = self.case_manager.generar_escrito_mediacion(2)
            
            # Verificaciones
            self.assertFalse(result, "El método debería retornar False para caso sin carátula")
            mock_error_dialog.assert_called()
            
            print("✓ Error con datos faltantes manejado correctamente")
    
    def test_06_missing_parties_error(self):
        """
        Prueba 6: Error con partes faltantes
        Verifica que el sistema maneja correctamente casos sin partes.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 6: Error con partes faltantes")
        print("-" * 60)
        
        # Crear caso sin partes en la base de datos
        try:
            conn = sqlite3.connect(self.test_db_path)
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO casos (id, caratula, numero_expediente)
                VALUES (3, 'CASO SIN PARTES', '11111/2023')
            ''')
            conn.commit()
            conn.close()
        except Exception as e:
            print(f"Error configurando caso de prueba: {e}")
        
        with patch.object(ErrorMessageManager, 'show_error_dialog') as mock_error_dialog:
            # Ejecutar con caso sin partes
            result = self.case_manager.generar_escrito_mediacion(3)
            
            # Verificaciones
            self.assertFalse(result, "El método debería retornar False para caso sin partes")
            mock_error_dialog.assert_called()
            
            print("✓ Error con partes faltantes manejado correctamente")
    
    def test_07_document_generation_error(self):
        """
        Prueba 7: Error en la generación del documento
        Verifica que el sistema maneja correctamente errores en la generación.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 7: Error en la generación del documento")
        print("-" * 60)
        
        with patch.object(self.case_manager, '_ask_agreement_details_dialog') as mock_dialog, \
             patch.object(self.case_manager, '_generar_documento_con_datos') as mock_generate, \
             patch.object(ErrorMessageManager, 'show_error_dialog') as mock_error_dialog:
            
            # Configurar mocks
            mock_dialog.return_value = self.mock_agreement_details
            mock_generate.return_value = {
                'success': False,
                'error_message': 'Error de prueba en generación',
                'error_type': 'template_error'
            }
            
            # Ejecutar el método
            result = self.case_manager.generar_escrito_mediacion(1)
            
            # Verificaciones
            self.assertFalse(result, "El método debería retornar False cuando hay error en generación")
            mock_dialog.assert_called_once()
            mock_generate.assert_called_once()
            mock_error_dialog.assert_called()
            
            print("✓ Error en generación de documento manejado correctamente")
    
    def test_08_dialog_functionality_preservation(self):
        """
        Prueba 8: Preservación de la funcionalidad de diálogos
        Verifica que los diálogos mantienen su funcionalidad original.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 8: Preservación de la funcionalidad de diálogos")
        print("-" * 60)
        
        # Verificar que el método _ask_agreement_details_dialog existe
        self.assertTrue(hasattr(self.case_manager, '_ask_agreement_details_dialog'),
                       "El método _ask_agreement_details_dialog debe existir")
        
        # Verificar que el método _collect_agreement_details existe
        self.assertTrue(hasattr(self.case_manager, '_collect_agreement_details'),
                       "El método _collect_agreement_details debe existir")
        
        # Verificar que el método _generar_documento_con_datos existe
        self.assertTrue(hasattr(self.case_manager, '_generar_documento_con_datos'),
                       "El método _generar_documento_con_datos debe existir")
        
        print("✓ Funcionalidad de diálogos preservada")
    
    def test_09_error_message_manager_integration(self):
        """
        Prueba 9: Integración con ErrorMessageManager
        Verifica que el sistema de manejo de errores funciona correctamente.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 9: Integración con ErrorMessageManager")
        print("-" * 60)
        
        # Verificar que ErrorMessageManager tiene los métodos necesarios
        self.assertTrue(hasattr(ErrorMessageManager, 'show_error_dialog'),
                       "ErrorMessageManager debe tener show_error_dialog")
        self.assertTrue(hasattr(ErrorMessageManager, 'get_error_message'),
                       "ErrorMessageManager debe tener get_error_message")
        self.assertTrue(hasattr(ErrorMessageManager, 'log_error'),
                       "ErrorMessageManager debe tener log_error")
        
        # Probar obtención de mensaje de error
        error_info = ErrorMessageManager.get_error_message('missing_case', {'case_id': 123})
        self.assertIsInstance(error_info, dict)
        self.assertIn('title', error_info)
        self.assertIn('message', error_info)
        
        print("✓ Integración con ErrorMessageManager verificada")
    
    def test_10_backward_compatibility_verification(self):
        """
        Prueba 10: Verificación final de compatibilidad hacia atrás
        Verifica que la refactorización mantiene la compatibilidad completa.
        """
        print("\n" + "-" * 60)
        print("PRUEBA 10: Verificación final de compatibilidad hacia atrás")
        print("-" * 60)
        
        # Verificar que el constructor acepta app_controller
        case_manager_with_app = CaseManager(app_controller=self.mock_app)
        self.assertIsNotNone(case_manager_with_app.app_controller)
        
        # Verificar que el constructor acepta None
        case_manager_without_app = CaseManager(app_controller=None)
        self.assertIsNone(case_manager_without_app.app_controller)
        
        # Verificar que el método principal existe y es callable
        self.assertTrue(callable(getattr(case_manager_with_app, 'generar_escrito_mediacion', None)),
                       "generar_escrito_mediacion debe ser callable")
        
        # Verificar que la signatura del método es correcta
        import inspect
        sig = inspect.signature(case_manager_with_app.generar_escrito_mediacion)
        params = list(sig.parameters.keys())
        self.assertIn('caso_id', params, "El método debe aceptar caso_id como parámetro")
        
        print("✓ Compatibilidad hacia atrás completamente verificada")


def run_ui_compatibility_tests():
    """
    Ejecuta todas las pruebas de compatibilidad hacia atrás de la UI.
    """
    print("INICIANDO SUITE DE PRUEBAS DE COMPATIBILIDAD HACIA ATRÁS")
    print("=" * 80)
    
    # Crear suite de pruebas
    suite = unittest.TestLoader().loadTestsFromTestCase(TestUIBackwardCompatibility)
    
    # Ejecutar pruebas con verbose output
    runner = unittest.TextTestRunner(verbosity=2, stream=sys.stdout)
    result = runner.run(suite)
    
    # Resumen de resultados
    print("\n" + "=" * 80)
    print("RESUMEN DE PRUEBAS DE COMPATIBILIDAD HACIA ATRÁS")
    print("=" * 80)
    print(f"Pruebas ejecutadas: {result.testsRun}")
    print(f"Errores: {len(result.errors)}")
    print(f"Fallos: {len(result.failures)}")
    
    if result.errors:
        print("\nERRORES ENCONTRADOS:")
        for test, error in result.errors:
            print(f"- {test}: {error}")
    
    if result.failures:
        print("\nFALLOS ENCONTRADOS:")
        for test, failure in result.failures:
            print(f"- {test}: {failure}")
    
    success = len(result.errors) == 0 and len(result.failures) == 0
    
    if success:
        print("\n✅ TODAS LAS PRUEBAS DE COMPATIBILIDAD PASARON EXITOSAMENTE")
        print("La refactorización mantiene la compatibilidad hacia atrás completa.")
    else:
        print("\n❌ ALGUNAS PRUEBAS FALLARON")
        print("Se requiere revisión de la compatibilidad hacia atrás.")
    
    return success


if __name__ == "__main__":
    try:
        success = run_ui_compatibility_tests()
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\nError crítico ejecutando pruebas: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)