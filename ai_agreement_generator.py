"""
AI-Powered Agreement Generator
Sistema avanzado para generar acuerdos de mediación usando IA y documentos de ejemplo.
"""

import os
import re
import json
import logging
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
from pathlib import Path

# AI and document processing imports
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available for document processing")

try:
    from langchain_community.llms import Ollama
    from langchain.prompts import PromptTemplate
    from langchain.chains import LLMChain
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    print("Warning: LangChain not available for AI processing")

# Local imports
import crm_database as db
from case_dialog_manager import CaseManager

class AIAgreementGenerator:
    """
    Generador de acuerdos de mediación usando IA y análisis de documentos de ejemplo.
    """

    def __init__(self, example_document_path: Optional[str] = None):
        """
        Inicializa el generador de acuerdos con IA.

        Args:
            example_document_path: Ruta al documento de ejemplo para análisis
        """
        self.logger = logging.getLogger(__name__)
        if not self.logger.handlers:
            console_handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            console_handler.setFormatter(formatter)
            self.logger.addHandler(console_handler)
            self.logger.setLevel(logging.INFO)

        self.example_document_path = example_document_path
        self.llm = None
        self.document_structure = {}
        self.case_manager = CaseManager(app_controller=None)

        # Inicializar LLM si está disponible
        if LANGCHAIN_AVAILABLE:
            try:
                self.llm = Ollama(model="mistral-small:22b")
                self.logger.info("LLM inicializado correctamente")
            except Exception as e:
                self.logger.error(f"Error inicializando LLM: {e}")
                self.llm = None

        # Cargar y analizar documento de ejemplo si se proporciona
        if example_document_path and os.path.exists(example_document_path):
            self._analyze_example_document(example_document_path)

    def _analyze_example_document(self, document_path: str) -> Dict[str, Any]:
        """
        Analiza un documento de ejemplo para extraer estructura y patrones.

        Args:
            document_path: Ruta al documento de ejemplo

        Returns:
            Dict con la estructura analizada del documento
        """
        try:
            if not DOCX_AVAILABLE:
                self.logger.error("python-docx no disponible para análisis de documentos")
                return {}

            self.logger.info(f"Analizando documento de ejemplo: {document_path}")

            doc = Document(document_path)
            structure = {
                'paragraphs': [],
                'tables': [],
                'placeholders': [],
                'sections': [],
                'metadata': {}
            }

            # Extraer párrafos y buscar placeholders
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    structure['paragraphs'].append({
                        'index': i,
                        'text': text,
                        'style': paragraph.style.name if paragraph.style else None,
                        'alignment': str(paragraph.alignment) if paragraph.alignment else None
                    })

                    # Buscar placeholders comunes en acuerdos
                    placeholders = self._extract_placeholders(text)
                    if placeholders:
                        structure['placeholders'].extend(placeholders)

            # Extraer información de tablas
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)

                structure['tables'].append({
                    'index': i,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })

            # Extraer metadatos del documento
            structure['metadata'] = {
                'total_paragraphs': len(structure['paragraphs']),
                'total_tables': len(structure['tables']),
                'placeholders_found': len(set(structure['placeholders'])),
                'analysis_date': datetime.now().isoformat()
            }

            self.document_structure = structure
            self.logger.info(f"Análisis completado. Estructura extraída: {structure['metadata']}")

            return structure

        except Exception as e:
            self.logger.error(f"Error analizando documento de ejemplo: {e}")
            return {}

    def _extract_placeholders(self, text: str) -> List[str]:
        """
        Extrae placeholders y variables del texto.

        Args:
            text: Texto a analizar

        Returns:
            Lista de placeholders encontrados
        """
        placeholders = []

        # Patrones comunes de placeholders en documentos legales
        patterns = [
            r'\[([A-Z_][A-Z0-9_]*)\]',  # [VARIABLE_NAME]
            r'\{([a-z_][a-z0-9_]*)\}',  # {variable_name}
            r'<<([^>]+)>>',            # <<Variable>>
            r'\{%([^%}]+)%\}',         # {% template_var %}
            r'\$\{([^}]+)\}',           # ${variable}
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            placeholders.extend(matches)

        return placeholders

    def _get_case_data_for_ai(self, case_id: int) -> Dict[str, Any]:
        """
        Obtiene y estructura los datos del caso para procesamiento con IA.

        Args:
            case_id: ID del caso

        Returns:
            Dict con datos estructurados del caso
        """
        try:
            # Obtener datos básicos del caso
            case_data = db.get_case_by_id(case_id)
            if not case_data:
                return {}

            # Obtener partes del caso
            actors = db.get_partes_by_caso_and_tipo(case_id, 'ACTOR')
            defendants = db.get_partes_by_caso_and_tipo(case_id, 'DEMANDADO')

            # Obtener datos del cliente
            client_data = None
            if case_data.get('cliente_id'):
                client_data = db.get_client_by_id(case_data['cliente_id'])

            # Estructurar datos para IA
            ai_data = {
                'case_info': {
                    'id': case_data.get('id'),
                    'caratula': case_data.get('caratula', ''),
                    'numero_expediente': case_data.get('numero_expediente', ''),
                    'anio_caratula': case_data.get('anio_caratula', ''),
                    'juzgado': case_data.get('juzgado', ''),
                    'jurisdiccion': case_data.get('jurisdiccion', ''),
                    'fecha_inicio': case_data.get('fecha_inicio', ''),
                    'estado': case_data.get('estado', ''),
                    'notas': case_data.get('notas', '')
                },
                'client': {
                    'nombre': client_data.get('nombre', '') if client_data else '',
                    'apellido': client_data.get('apellido', '') if client_data else '',
                    'dni': client_data.get('dni', '') if client_data else '',
                    'cuit': client_data.get('cuit', '') if client_data else '',
                    'domicilio': client_data.get('domicilio', '') if client_data else '',
                    'telefono': client_data.get('telefono', '') if client_data else '',
                    'email': client_data.get('email', '') if client_data else ''
                },
                'actors': [],
                'defendants': []
            }

            # Procesar actores
            for actor in actors:
                actor_info = {
                    'id': actor.get('id'),
                    'nombre_completo': actor.get('nombre_completo', ''),
                    'dni': actor.get('dni', ''),
                    'cuit': actor.get('cuit', ''),
                    'domicilio_real': actor.get('domicilio_real', ''),
                    'domicilio_legal': actor.get('domicilio_legal', ''),
                    'telefono': actor.get('telefono', ''),
                    'email': actor.get('email', ''),
                    'datos_bancarios': actor.get('datos_bancarios', ''),
                    'representantes': []
                }

                # Obtener representantes del actor
                representantes = self._get_representatives_for_party(actor.get('id'), case_id)
                actor_info['representantes'] = representantes

                ai_data['actors'].append(actor_info)

            # Procesar demandados
            for defendant in defendants:
                defendant_info = {
                    'id': defendant.get('id'),
                    'nombre_completo': defendant.get('nombre_completo', ''),
                    'dni': defendant.get('dni', ''),
                    'cuit': defendant.get('cuit', ''),
                    'domicilio_real': defendant.get('domicilio_real', ''),
                    'domicilio_legal': defendant.get('domicilio_legal', ''),
                    'telefono': defendant.get('telefono', ''),
                    'email': defendant.get('email', ''),
                    'datos_bancarios': defendant.get('datos_bancarios', ''),
                    'representantes': []
                }

                # Obtener representantes del demandado
                representantes = self._get_representatives_for_party(defendant.get('id'), case_id)
                defendant_info['representantes'] = representantes

                ai_data['defendants'].append(defendant_info)

            return ai_data

        except Exception as e:
            self.logger.error(f"Error obteniendo datos del caso {case_id}: {e}")
            return {}

    def _get_representatives_for_party(self, party_id: int, case_id: int) -> List[Dict[str, Any]]:
        """
        Obtiene los representantes legales de una parte.

        Args:
            party_id: ID de la parte
            case_id: ID del caso

        Returns:
            Lista de representantes
        """
        representantes = []

        try:
            # Obtener roles relacionados con esta parte
            roles = db.get_roles_by_case_id(case_id)

            for rol in roles:
                # Verificar si este rol representa a la parte actual
                if (rol.get('representa_a_id') == party_id or
                    rol.get('id') == party_id):
                    representante_info = {
                        'id': rol.get('id'),
                        'nombre_completo': rol.get('nombre_completo', ''),
                        'cuit': rol.get('cuit', ''),
                        'personeria': rol.get('datos_bancarios', 'Poder General Judicial'),
                        'telefono': rol.get('telefono', ''),
                        'email': rol.get('email', '')
                    }
                    representantes.append(representante_info)

        except Exception as e:
            self.logger.warning(f"Error obteniendo representantes para parte {party_id}: {e}")

        return representantes

    def generate_agreement_with_ai(self, case_id: int, agreement_details: Dict[str, Any],
                                  example_document_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Genera un acuerdo de mediación usando IA y análisis de documento de ejemplo.

        Args:
            case_id: ID del caso
            agreement_details: Detalles del acuerdo (monto, plazo, datos bancarios)
            example_document_path: Ruta opcional a documento de ejemplo

        Returns:
            Dict con resultado de la generación
        """
        try:
            self.logger.info(f"Iniciando generación de acuerdo con IA para caso {case_id}")

            # Obtener datos del caso
            case_data = self._get_case_data_for_ai(case_id)
            if not case_data:
                return {
                    'success': False,
                    'error_message': f'No se pudieron obtener los datos del caso {case_id}',
                    'error_type': 'case_data_error'
                }

            # Analizar documento de ejemplo si se proporciona
            if example_document_path and os.path.exists(example_document_path):
                example_structure = self._analyze_example_document(example_document_path)
            else:
                example_structure = self.document_structure

            # Generar acuerdo usando IA
            if self.llm and LANGCHAIN_AVAILABLE:
                ai_generated_content = self._generate_content_with_ai(case_data, agreement_details, example_structure)
            else:
                # Fallback a generación tradicional
                self.logger.warning("LLM no disponible, usando generación tradicional")
                return self.case_manager._generar_documento_con_datos(case_id, agreement_details)

            # Crear documento Word con el contenido generado
            if ai_generated_content and DOCX_AVAILABLE:
                document_result = self._create_document_from_ai_content(ai_generated_content, case_data, agreement_details)
                return document_result
            else:
                return {
                    'success': False,
                    'error_message': 'No se pudo generar el contenido del documento',
                    'error_type': 'content_generation_error'
                }

        except Exception as e:
            self.logger.error(f"Error generando acuerdo con IA: {e}")
            return {
                'success': False,
                'error_message': f'Error interno: {str(e)}',
                'error_type': 'internal_error'
            }

    def _generate_content_with_ai(self, case_data: Dict[str, Any],
                                 agreement_details: Dict[str, Any],
                                 example_structure: Dict[str, Any]) -> Optional[str]:
        """
        Genera contenido usando IA basado en datos del caso y estructura de ejemplo.

        Args:
            case_data: Datos estructurados del caso
            agreement_details: Detalles del acuerdo
            example_structure: Estructura del documento de ejemplo

        Returns:
            Contenido generado o None si falla
        """
        try:
            # Crear prompt para la IA
            prompt_template = """
Eres un abogado especializado en derecho civil y laboral argentino. Tu tarea es generar un acuerdo de mediación completo y profesional basado en los datos proporcionados.

DATOS DEL CASO:
{case_data}

DETALLES DEL ACUERDO:
{agreement_details}

ESTRUCTURA DE EJEMPLO ANALIZADA:
{example_structure}

INSTRUCCIONES:
1. Genera un acuerdo de mediación completo en español, siguiendo el formato y estructura legal argentina.
2. Utiliza todos los datos proporcionados del caso, actores, demandados y sus representantes.
3. Incluye el monto de compensación, plazo de pago y datos bancarios especificados.
4. Mantén un tono formal y profesional apropiado para un documento legal.
5. Estructura el documento con secciones claras: encabezado, partes involucradas, objeto del acuerdo, condiciones, firmas, etc.
6. Si hay múltiples actores o demandados, enuméralos claramente.
7. Incluye cláusulas estándar de mediación como confidencialidad, cumplimiento, etc.

ACUERDO DE MEDIACIÓN:
"""

            prompt = PromptTemplate(
                template=prompt_template,
                input_variables=["case_data", "agreement_details", "example_structure"]
            )

            # Crear cadena de IA
            chain = LLMChain(llm=self.llm, prompt=prompt)

            # Ejecutar generación
            result = chain.run(
                case_data=json.dumps(case_data, ensure_ascii=False, indent=2),
                agreement_details=json.dumps(agreement_details, ensure_ascii=False, indent=2),
                example_structure=json.dumps(example_structure, ensure_ascii=False, indent=2)
            )

            self.logger.info("Contenido generado exitosamente con IA")
            return result.strip()

        except Exception as e:
            self.logger.error(f"Error generando contenido con IA: {e}")
            return None

    def _create_document_from_ai_content(self, ai_content: str, case_data: Dict[str, Any],
                                       agreement_details: Dict[str, Any]) -> Dict[str, Any]:
        """
        Crea un documento Word a partir del contenido generado por IA.

        Args:
            ai_content: Contenido generado por IA
            case_data: Datos del caso
            agreement_details: Detalles del acuerdo

        Returns:
            Dict con resultado de la creación del documento
        """
        try:
            if not DOCX_AVAILABLE:
                return {
                    'success': False,
                    'error_message': 'python-docx no disponible para crear documentos',
                    'error_type': 'library_error'
                }

            # Crear nuevo documento
            doc = Document()

            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Dividir el contenido en párrafos
            paragraphs = ai_content.split('\n\n')

            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    paragraph = doc.add_paragraph(paragraph_text.strip())

                    # Aplicar formato especial a títulos y secciones
                    if any(keyword in paragraph_text.upper() for keyword in
                           ['ACUERDO', 'CLÁUSULA', 'ARTÍCULO', 'CONSIDERANDO', 'POR LO TANTO']):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run('')
                        run.bold = True

            # Generar nombre de archivo
            case_info = case_data.get('case_info', {})
            filename = f"Acuerdo_Mediacion_AI_{case_info.get('numero_expediente', 'Sin_Numero')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

            # Guardar documento
            output_path = os.path.join('generated_documents', filename)
            os.makedirs('generated_documents', exist_ok=True)

            doc.save(output_path)

            self.logger.info(f"Documento guardado exitosamente: {output_path}")

            return {
                'success': True,
                'document_path': output_path,
                'filename': filename,
                'error_message': None,
                'error_type': None
            }

        except Exception as e:
            self.logger.error(f"Error creando documento: {e}")
            return {
                'success': False,
                'error_message': f'Error creando documento: {str(e)}',
                'error_type': 'document_creation_error'
            }

    def analyze_document_for_patterns(self, document_path: str) -> Dict[str, Any]:
        """
        Analiza un documento para extraer patrones y estructura reutilizable.

        Args:
            document_path: Ruta al documento a analizar

        Returns:
            Dict con patrones y estructura extraída
        """
        try:
            if not os.path.exists(document_path):
                return {'error': f'Documento no encontrado: {document_path}'}

            structure = self._analyze_example_document(document_path)

            # Extraer patrones específicos de acuerdos de mediación
            patterns = {
                'header_patterns': [],
                'party_patterns': [],
                'agreement_patterns': [],
                'signature_patterns': []
            }

            for paragraph in structure.get('paragraphs', []):
                text = paragraph['text'].upper()

                # Patrones de encabezado
                if any(word in text for word in ['JUZGADO', 'EXPEDIENTE', 'EXPTE', 'CARATULA']):
                    patterns['header_patterns'].append(paragraph)

                # Patrones de partes
                elif any(word in text for word in ['ACTOR', 'DEMANDADO', 'REPRESENTANTE', 'APODERADO']):
                    patterns['party_patterns'].append(paragraph)

                # Patrones de acuerdo
                elif any(word in text for word in ['ACUERDO', 'COMPENSACIÓN', 'PAGO', 'MONTO', 'PLAZO']):
                    patterns['agreement_patterns'].append(paragraph)

                # Patrones de firma
                elif any(word in text for word in ['FIRMA', 'CONFORME', 'ACUERDO']):
                    patterns['signature_patterns'].append(paragraph)

            analysis_result = {
                'structure': structure,
                'patterns': patterns,
                'placeholders': list(set(structure.get('placeholders', []))),
                'recommendations': self._generate_pattern_recommendations(patterns)
            }

            return analysis_result

        except Exception as e:
            self.logger.error(f"Error analizando documento para patrones: {e}")
            return {'error': str(e)}

    def _generate_pattern_recommendations(self, patterns: Dict[str, List]) -> List[str]:
        """
        Genera recomendaciones basadas en los patrones encontrados.

        Args:
            patterns: Patrones extraídos del documento

        Returns:
            Lista de recomendaciones
        """
        recommendations = []

        if not patterns['header_patterns']:
            recommendations.append("Considerar agregar sección de encabezado con datos del juzgado")

        if not patterns['party_patterns']:
            recommendations.append("El documento debería incluir identificación clara de las partes")

        if not patterns['agreement_patterns']:
            recommendations.append("Falta sección de términos del acuerdo (monto, plazo, condiciones)")

        if not patterns['signature_patterns']:
            recommendations.append("Agregar sección de firmas y conformidad")

        if len(patterns['party_patterns']) < 2:
            recommendations.append("El documento parece no cubrir múltiples partes - considerar expansión")

        return recommendations

    def get_available_example_documents(self) -> List[str]:
        """
        Obtiene lista de documentos de ejemplo disponibles.

        Returns:
            Lista de rutas de documentos de ejemplo
        """
        example_dirs = [
            'plantillas/mediacion',
            'ejemplos',
            'examples',
            'documentos_ejemplo'
        ]

        available_documents = []

        for directory in example_dirs:
            if os.path.exists(directory):
                for file in os.listdir(directory):
                    if file.endswith(('.docx', '.doc')) and 'ejemplo' in file.lower():
                        available_documents.append(os.path.join(directory, file))

        return available_documents