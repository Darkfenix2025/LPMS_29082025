#!/usr/bin/env python3
"""
Comprehensive UI backward compatibility test that creates a complete test environment.
"""

import sys
import os
import tkinter as tk
import sqlite3
import tempfile
import shutil
from unittest.mock import patch, MagicMock

# Add current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Import modules
import crm_database as db
from case_dialog_manager import CaseManager

class ComprehensiveUITest:
    """Comprehensive UI test with complete environment setup."""
    
    def __init__(self):
        self.test_results = []
        self.setup_complete_test_environment()
    
    def setup_complete_test_environment(self):
        """Setup complete test environment including templates and proper database."""
        # Create temporary directory for all test files
        self.temp_dir = tempfile.mkdtemp()
        
        # Setup test database
        self.test_db_path = os.path.join(self.temp_dir, 'test_crm.db')
        self.setup_complete_database()
        
        # Setup template directory and files
        self.setup_template_files()
        
        # Create Tkinter root
        self.root = tk.Tk()
        self.root.withdraw()
        
        # Create mock app
        self.mock_app = type('MockApp', (), {})()
        self.mock_app.root = self.root
        
        # Create case manager
        self.case_manager = CaseManager(app_controller=self.mock_app)
    
    def setup_complete_database(self):
        """Setup complete database with all necessary tables and data."""
        try:
            conn = sqlite3.connect(self.test_db_path)
            cursor = conn.cursor()
            
            # Create all necessary tables
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS casos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    caratula TEXT NOT NULL,
                    numero_expediente TEXT,
                    anio_caratula TEXT,
                    juzgado TEXT,
                    jurisdiccion TEXT,
                    fecha_creacion TEXT,
                    estado TEXT DEFAULT 'Activo'
                )
            ''')
            
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS partes (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    caso_id INTEGER,
                    nombre_completo TEXT NOT NULL,
                    tipo_parte TEXT NOT NULL,
                    dni TEXT,
                    domicilio TEXT,
                    telefono TEXT,
                    email TEXT,
                    FOREIGN KEY (caso_id) REFERENCES casos (id)
                )
            ''')
            
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS roles_partes (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    parte_id INTEGER,
                    rol TEXT NOT NULL,
                    FOREIGN KEY (parte_id) REFERENCES partes (id)
                )
            ''')
            
            # Insert comprehensive test data
            cursor.execute('''
                INSERT INTO casos (id, caratula, numero_expediente, anio_caratula, juzgado, jurisdiccion, fecha_creacion)
                VALUES (1, 'PÉREZ JUAN C/ GARCÍA MARÍA S/ DAÑOS Y PERJUICIOS', '12345/2023', '2023', 'Juzgado Civil N° 1', 'La Plata', '2023-01-15')
            ''')
            
            # Insert parties with proper roles
            cursor.execute('''
                INSERT INTO partes (id, caso_id, nombre_completo, tipo_parte, dni, domicilio, telefono, email)
                VALUES 
                (1, 1, 'PÉREZ JUAN', 'Actor', '12345678', 'Calle Falsa 123', '221-1234567', 'juan.perez@email.com'),
                (2, 1, 'GARCÍA MARÍA', 'Demandado', '87654321', 'Avenida Siempre Viva 456', '221-7654321', 'maria.garcia@email.com')
            ''')
            
            # Insert roles
            cursor.execute('''
                INSERT INTO roles_partes (parte_id, rol)
                VALUES 
                (1, 'Actor'),
                (2, 'Demandado')
            ''')
            
            # Insert case without caratula for error testing
            cursor.execute('''
                INSERT INTO casos (id, caratula, numero_expediente, anio_caratula, juzgado, jurisdiccion)
                VALUES (2, '', '54321/2023', '2023', 'Juzgado Civil N° 2', 'La Plata')
            ''')
            
            # Insert case without parties for error testing
            cursor.execute('''
                INSERT INTO casos (id, caratula, numero_expediente, anio_caratula, juzgado, jurisdiccion)
                VALUES (3, 'CASO SIN PARTES', '11111/2023', '2023', 'Juzgado Civil N° 3', 'La Plata')
            ''')
            
            conn.commit()
            conn.close()
            
            # Set database path
            db.DATABASE_PATH = self.test_db_path
            
        except Exception as e:
            print(f"Error setting up database: {e}")
            raise
    
    def setup_template_files(self):
        """Setup template files for document generation."""
        try:
            # Create template directory
            template_dir = os.path.join(self.temp_dir, 'plantillas', 'mediacion')
            os.makedirs(template_dir, exist_ok=True)
            
            template_path = os.path.join(template_dir, 'acuerdo_base.docx')
            
            # Create a simple Word document template
            try:
                from docx import Document
                doc = Document()
                doc.add_paragraph("ACUERDO DE MEDIACIÓN - PLANTILLA DE PRUEBA")
                doc.add_paragraph("Caso: {{caratula}}")
                doc.add_paragraph("Actor: {{actor_nombre}}")
                doc.add_paragraph("Demandado: {{demandado_nombre}}")
                doc.add_paragraph("Monto: {{monto_compensacion_numeros}}")
                doc.add_paragraph("Plazo: {{plazo_pago_dias}} días")
                doc.save(template_path)
            except ImportError:
                # Create a dummy file if python-docx is not available
                with open(template_path, 'wb') as f:
                    f.write(b'DUMMY TEMPLATE FILE FOR TESTING')
            
            # Change working directory to temp dir so relative paths work
            self.original_cwd = os.getcwd()
            os.chdir(self.temp_dir)
            
        except Exception as e:
            print(f"Error setting up templates: {e}")
            raise
    
    def test_ui_flow_with_complete_environment(self):
        """Test UI flow with complete environment setup."""
        print("\\n" + "="*60)
        print("PRUEBA: Flujo UI con entorno completo")
        print("="*60)
        
        test_passed = True
        
        try:
            # Mock agreement details
            mock_agreement_details = {
                'monto_compensacion_numeros': '150000.50',
                'monto_compensacion_letras': 'CIENTO CINCUENTA MIL PESOS CON CINCUENTA CENTAVOS',
                'plazo_pago_dias': '30',
                'plazo_pago_letras': 'TREINTA',
                'banco_actor': 'Banco Nación',
                'cbu_actor': '1234567890123456789012',
                'alias_actor': 'mi.alias.mp',
                'cuit_actor': '20-12345678-9'
            }
            
            # Test 1: Complete successful flow
            with patch.object(self.case_manager, '_ask_agreement_details_dialog', return_value=mock_agreement_details) as mock_dialog, \
                 patch('tkinter.filedialog.asksaveasfilename', return_value=os.path.join(self.temp_dir, 'test_document.docx')) as mock_save:
                
                result = self.case_manager.generar_escrito_mediacion(1)
                
                if result == True:
                    self.test_results.append("✅ Flujo completo exitoso con entorno real")
                else:
                    self.test_results.append(f"❌ Flujo completo falló: {result}")
                    test_passed = False
                
                if mock_dialog.called:
                    self.test_results.append("✅ Diálogo de detalles fue llamado")
                else:
                    self.test_results.append("❌ Diálogo de detalles NO fue llamado")
                    test_passed = False
                
                if mock_save.called:
                    self.test_results.append("✅ Diálogo de guardado fue llamado")
                else:
                    self.test_results.append("⚠️  Diálogo de guardado NO fue llamado (puede ser por error anterior)")
            
            # Test 2: User cancellation
            with patch.object(self.case_manager, '_ask_agreement_details_dialog', return_value=None) as mock_dialog:
                
                result = self.case_manager.generar_escrito_mediacion(1)
                
                if result == False:
                    self.test_results.append("✅ Cancelación de usuario manejada correctamente")
                else:
                    self.test_results.append(f"❌ Cancelación retorna {result} (esperado: False)")
                    test_passed = False
            
            # Test 3: Invalid case ID
            result = self.case_manager.generar_escrito_mediacion(99999)
            if result == False:
                self.test_results.append("✅ Caso inválido manejado correctamente")
            else:
                self.test_results.append(f"❌ Caso inválido retorna {result} (esperado: False)")
                test_passed = False
            
            # Test 4: Case without caratula
            result = self.case_manager.generar_escrito_mediacion(2)
            if result == False:
                self.test_results.append("✅ Caso sin carátula manejado correctamente")
            else:
                self.test_results.append(f"❌ Caso sin carátula retorna {result} (esperado: False)")
                test_passed = False
            
            # Test 5: Case without parties
            result = self.case_manager.generar_escrito_mediacion(3)
            if result == False:
                self.test_results.append("✅ Caso sin partes manejado correctamente")
            else:
                self.test_results.append(f"❌ Caso sin partes retorna {result} (esperado: False)")
                test_passed = False
        
        except Exception as e:
            self.test_results.append(f"❌ Error en prueba de flujo UI: {e}")
            test_passed = False
        
        return test_passed
    
    def test_method_separation(self):
        """Test that the refactored methods work correctly."""
        print("\\n" + "="*60)
        print("PRUEBA: Separación de métodos refactorizados")
        print("="*60)
        
        test_passed = True
        
        try:
            # Test that _generar_documento_con_datos exists and is callable
            if hasattr(self.case_manager, '_generar_documento_con_datos'):
                self.test_results.append("✅ Método _generar_documento_con_datos existe")
                
                if callable(getattr(self.case_manager, '_generar_documento_con_datos')):
                    self.test_results.append("✅ Método _generar_documento_con_datos es callable")
                else:
                    self.test_results.append("❌ Método _generar_documento_con_datos NO es callable")
                    test_passed = False
            else:
                self.test_results.append("❌ Método _generar_documento_con_datos NO existe")
                test_passed = False
            
            # Test method signature
            import inspect
            if hasattr(self.case_manager, '_generar_documento_con_datos'):
                sig = inspect.signature(self.case_manager._generar_documento_con_datos)
                params = list(sig.parameters.keys())
                
                expected_params = ['caso_id', 'details_acuerdo']
                if all(param in params for param in expected_params):
                    self.test_results.append("✅ Signatura de _generar_documento_con_datos es correcta")
                else:
                    self.test_results.append(f"❌ Signatura incorrecta. Esperado: {expected_params}, Actual: {params}")
                    test_passed = False
            
            # Test that the UI orchestrator method exists
            if hasattr(self.case_manager, 'generar_escrito_mediacion'):
                self.test_results.append("✅ Método orquestador generar_escrito_mediacion existe")
            else:
                self.test_results.append("❌ Método orquestador generar_escrito_mediacion NO existe")
                test_passed = False
        
        except Exception as e:
            self.test_results.append(f"❌ Error en prueba de separación de métodos: {e}")
            test_passed = False
        
        return test_passed
    
    def test_error_handling_robustness(self):
        """Test error handling robustness."""
        print("\\n" + "="*60)
        print("PRUEBA: Robustez del manejo de errores")
        print("="*60)
        
        test_passed = True
        
        try:
            # Test with None case_id
            with patch('tkinter.messagebox.showerror'):
                result = self.case_manager.generar_escrito_mediacion(None)
                if result == False:
                    self.test_results.append("✅ caso_id None manejado correctamente")
                else:
                    self.test_results.append(f"❌ caso_id None retorna {result} (esperado: False)")
                    test_passed = False
            
            # Test with invalid case_id type
            with patch('tkinter.messagebox.showerror'):
                result = self.case_manager.generar_escrito_mediacion("invalid")
                if result == False:
                    self.test_results.append("✅ caso_id inválido manejado correctamente")
                else:
                    self.test_results.append(f"❌ caso_id inválido retorna {result} (esperado: False)")
                    test_passed = False
            
            # Test ErrorMessageManager integration
            from case_dialog_manager import ErrorMessageManager
            
            error_info = ErrorMessageManager.get_error_message('missing_case', {'case_id': 123})
            if isinstance(error_info, dict) and 'title' in error_info and 'message' in error_info:
                self.test_results.append("✅ ErrorMessageManager funciona correctamente")
            else:
                self.test_results.append("❌ ErrorMessageManager no funciona correctamente")
                test_passed = False
        
        except Exception as e:
            self.test_results.append(f"❌ Error en prueba de manejo de errores: {e}")
            test_passed = False
        
        return test_passed
    
    def run_all_tests(self):
        """Run all comprehensive tests."""
        print("INICIANDO PRUEBAS COMPREHENSIVAS DE COMPATIBILIDAD UI")
        print("="*80)
        
        all_tests_passed = True
        
        # Run tests
        tests = [
            ("Separación de Métodos", self.test_method_separation),
            ("Manejo de Errores", self.test_error_handling_robustness),
            ("Flujo UI Completo", self.test_ui_flow_with_complete_environment)
        ]
        
        for test_name, test_func in tests:
            print(f"\\nEjecutando: {test_name}")
            try:
                test_result = test_func()
                if not test_result:
                    all_tests_passed = False
            except Exception as e:
                print(f"Error ejecutando {test_name}: {e}")
                self.test_results.append(f"❌ Error crítico en {test_name}: {e}")
                all_tests_passed = False
        
        # Print results
        print("\\n" + "="*80)
        print("RESULTADOS DE PRUEBAS COMPREHENSIVAS")
        print("="*80)
        
        for result in self.test_results:
            print(result)
        
        print("\\n" + "="*80)
        if all_tests_passed:
            print("✅ TODAS LAS PRUEBAS COMPREHENSIVAS PASARON")
            print("La compatibilidad hacia atrás de la UI está completamente verificada.")
        else:
            print("❌ ALGUNAS PRUEBAS COMPREHENSIVAS FALLARON")
            print("Se requiere revisión adicional de la implementación.")
        print("="*80)
        
        return all_tests_passed
    
    def cleanup(self):
        """Clean up test resources."""
        try:
            if hasattr(self, 'original_cwd'):
                os.chdir(self.original_cwd)
            if hasattr(self, 'root'):
                self.root.destroy()
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception as e:
            print(f"Error during cleanup: {e}")


def run_comprehensive_ui_tests():
    """Run comprehensive UI tests."""
    tester = None
    try:
        tester = ComprehensiveUITest()
        success = tester.run_all_tests()
        return success
    except Exception as e:
        print(f"Error running comprehensive tests: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        if tester:
            tester.cleanup()


if __name__ == "__main__":
    try:
        success = run_comprehensive_ui_tests()
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"Critical error: {e}")
        sys.exit(1)